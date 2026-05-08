from __future__ import annotations

import csv
import hashlib
import json
import os
import posixpath
import re
import shutil
import unicodedata
import zipfile
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any
from urllib.parse import unquote
from xml.etree import ElementTree as ET

import fitz
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image


REPO_ROOT = Path(__file__).resolve().parents[3]
PROJECT_ROOT = REPO_ROOT / "projects" / "english-lang-arts-10-2-docx-export"
ZIP_PATH = Path(
    r"C:\Users\dean.guedo\Downloads\D2LCCExport_149674_25-26 _ S2 _ English Lang. Arts 10-2 _ Per 1(A) _ _20265624.ZIP"
)

RAW_DIR = PROJECT_ROOT / "raw"
META_DIR = PROJECT_ROOT / "meta"
EXPORT_DIR = PROJECT_ROOT / "exports"
DOCX_DIR = EXPORT_DIR / "docx"
SUPPORT_DIR = EXPORT_DIR / "supporting-files"


@dataclass(frozen=True)
class UnitSpec:
    output_title: str
    source_title: str | None = None
    source_item_title: str | None = None
    note: str | None = None


REQUESTED_UNITS = [
    UnitSpec("Course Information", source_title="Course Information"),
    UnitSpec(
        "Unit 1: Introduction to Interpreting and Creating Texts",
        source_title="Unit 1: Introduction to Interpreting and Creating Texts",
    ),
    UnitSpec(
        "Unit 2: Literary Exploration - Shorter Texts",
        source_title="Unit 2: Analyzing Shorter Texts",
        note="Screenshot title differs from the ZIP manifest title.",
    ),
    UnitSpec(
        "Unit 3: Shakespeare - Romeo and Juliet Animated Comic",
        source_title="Unit 3: Graphic Novels & Shakespeare",
        note="Screenshot title differs from the ZIP manifest title.",
    ),
    UnitSpec(
        "Unit 4: Reading Comprehension",
        source_item_title="Reading Comprehension Review",
        note=(
            "No top-level Unit 4 Reading Comprehension module exists in this ZIP. "
            "This document contains the matching manifest topic found inside Unit 2."
        ),
    ),
    UnitSpec(
        "Unit 5: Extended Text",
        source_title="Unit 4: Novel Study",
        note="Mapped by course order: the ZIP names this source module Unit 4: Novel Study.",
    ),
    UnitSpec(
        "Unit 6: Informative and Persuasive Texts",
        source_title="Unit 5: Informative and Persuasive Texts",
        note="Mapped by course order: the ZIP names this source module Unit 5.",
    ),
    UnitSpec(
        "Unit 7: Film Study",
        source_title="Unit 6: Film Study",
        note="Mapped by course order: the ZIP names this source module Unit 6.",
    ),
    UnitSpec("Video Clips", source_title="Video Clips"),
    UnitSpec("Teacher Materials (Keep Hidden)", source_title="Teacher Materials (Keep Hidden)"),
]


def local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1]


def item_title(item: ET.Element) -> str:
    for child in item:
        if local_name(child.tag) == "title":
            return "".join(child.itertext()).strip()
    return ""


def item_children(item: ET.Element) -> list[ET.Element]:
    return [child for child in item if local_name(child.tag) == "item"]


def walk_items(item: ET.Element) -> list[ET.Element]:
    result = [item]
    for child in item_children(item):
        result.extend(walk_items(child))
    return result


def normalize_key(value: str) -> str:
    return re.sub(r"\s+", " ", value.strip()).casefold()


def safe_name(value: str, max_len: int = 96) -> str:
    normalized = unicodedata.normalize("NFKD", value)
    ascii_value = normalized.encode("ascii", "ignore").decode("ascii")
    ascii_value = re.sub(r"[^\w\s().,&+-]+", "", ascii_value)
    ascii_value = re.sub(r"\s+", " ", ascii_value).strip()
    ascii_value = ascii_value.replace("&", "and")
    ascii_value = ascii_value.strip(" .")
    if not ascii_value:
        ascii_value = "untitled"
    return ascii_value[:max_len].strip(" .")


def rel_posix(path: Path, start: Path) -> str:
    return Path(os.path.relpath(path, start)).as_posix()


def clean_text(value: str) -> str:
    return re.sub(r"\s+", " ", value.replace("\xa0", " ")).strip()


def package_path_exists(zip_file: zipfile.ZipFile, href: str) -> bool:
    try:
        zip_file.getinfo(href)
        return True
    except KeyError:
        return False


def resolve_package_href(base_href: str, relative_href: str, zip_file: zipfile.ZipFile) -> str | None:
    if not relative_href or re.match(r"^[a-zA-Z][a-zA-Z0-9+.-]*:", relative_href):
        return None
    href = relative_href.split("#", 1)[0].split("?", 1)[0].strip()
    if not href:
        return None
    href = unquote(href.replace("\\", "/"))
    candidate = posixpath.normpath(posixpath.join(posixpath.dirname(base_href), href))
    if package_path_exists(zip_file, candidate):
        return candidate
    if package_path_exists(zip_file, href):
        return href
    return None


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(color)
    properties.append(underline)
    run.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


class DocxExporter:
    def __init__(self, zip_file: zipfile.ZipFile, manifest_root: ET.Element):
        self.zip_file = zip_file
        self.manifest_root = manifest_root
        self.zip_sizes = {
            info.filename: info.file_size for info in zip_file.infolist() if not info.is_dir()
        }
        self.resources = self._read_resources()
        self.audit: dict[str, Any] = {
            "generatedAt": datetime.now().isoformat(timespec="seconds"),
            "sourceZip": str(ZIP_PATH),
            "sourceZipBytes": ZIP_PATH.stat().st_size,
            "requestedUnits": [],
            "excludedTopLevelModules": [],
            "missingResources": [],
            "unresolvedHtmlAssets": [],
            "embeddedImages": 0,
            "renderedPdfPages": 0,
            "copiedSupportFiles": [],
            "docxFiles": [],
        }
        self._copied_support: dict[str, Path] = {}
        self.current_output_title = "Unembedded HTML Assets"

    def _read_resources(self) -> dict[str, dict[str, Any]]:
        resources: dict[str, dict[str, Any]] = {}
        for resource in self.manifest_root.iter():
            if local_name(resource.tag) != "resource":
                continue
            identifier = resource.attrib.get("identifier")
            if not identifier:
                continue
            files = [
                file_node.attrib.get("href")
                for file_node in resource
                if local_name(file_node.tag) == "file" and file_node.attrib.get("href")
            ]
            resources[identifier] = {
                "identifier": identifier,
                "type": resource.attrib.get("type"),
                "files": files,
            }
        return resources

    def top_level_modules(self) -> list[ET.Element]:
        organization = next(
            node for node in self.manifest_root.iter() if local_name(node.tag) == "organization"
        )
        root_items = item_children(organization)
        if len(root_items) == 1 and not item_title(root_items[0]):
            return item_children(root_items[0])
        return root_items

    def find_top_level(self, title: str) -> ET.Element | None:
        wanted = normalize_key(title)
        for item in self.top_level_modules():
            if normalize_key(item_title(item)) == wanted:
                return item
        return None

    def find_item_by_title(self, title: str) -> ET.Element | None:
        wanted = normalize_key(title)
        for item in self.manifest_root.iter():
            if local_name(item.tag) == "item" and normalize_key(item_title(item)) == wanted:
                return item
        return None

    def build_all(self):
        self._prepare_outputs()
        self._write_raw_source_notes()
        requested_source_titles = {
            normalize_key(spec.source_title)
            for spec in REQUESTED_UNITS
            if spec.source_title is not None
        }
        for module in self.top_level_modules():
            title = item_title(module)
            if normalize_key(title) not in requested_source_titles:
                self.audit["excludedTopLevelModules"].append(
                    {
                        "title": title,
                        "itemCount": len([item for item in walk_items(module) if item.attrib.get("identifierref")]),
                    }
                )

        for index, spec in enumerate(REQUESTED_UNITS, start=1):
            print(f"[{index:02d}/{len(REQUESTED_UNITS):02d}] {spec.output_title}")
            self.build_unit(index, spec)

        self.write_resource_indexes()
        self.verify_outputs()
        (META_DIR / "docx-export-audit.json").write_text(
            json.dumps(self.audit, indent=2), encoding="utf-8"
        )
        self.write_audit_markdown()

    def _prepare_outputs(self):
        for target in [DOCX_DIR, SUPPORT_DIR]:
            if target.exists():
                shutil.rmtree(target)
            target.mkdir(parents=True, exist_ok=True)
        RAW_DIR.mkdir(parents=True, exist_ok=True)
        META_DIR.mkdir(parents=True, exist_ok=True)

    def _write_raw_source_notes(self):
        (RAW_DIR / "imsmanifest.xml").write_bytes(self.zip_file.read("imsmanifest.xml"))
        source_note = {
            "sourceZip": str(ZIP_PATH),
            "sourceZipBytes": ZIP_PATH.stat().st_size,
            "sourceZipLastWriteTime": datetime.fromtimestamp(ZIP_PATH.stat().st_mtime).isoformat(
                timespec="seconds"
            ),
            "rawStorageNote": "The full ZIP is not duplicated here because it is multi-gigabyte. The copied imsmanifest.xml is the structural source of truth for this export.",
        }
        (RAW_DIR / "source-package.json").write_text(json.dumps(source_note, indent=2), encoding="utf-8")

    def build_unit(self, index: int, spec: UnitSpec):
        if spec.source_title:
            source_item = self.find_top_level(spec.source_title)
            source_kind = "top-level-module"
        else:
            source_item = self.find_item_by_title(spec.source_item_title or "")
            source_kind = "matching-topic"

        unit_record: dict[str, Any] = {
            "outputTitle": spec.output_title,
            "sourceTitle": item_title(source_item) if source_item is not None else None,
            "sourceKind": source_kind,
            "note": spec.note,
            "includedItems": [],
            "sourceMissing": source_item is None,
        }

        document = self.new_document()
        self.current_output_title = spec.output_title
        document.add_heading(spec.output_title, level=0)
        if source_item is None:
            document.add_paragraph("Source content was not found in imsmanifest.xml.")
            if spec.note:
                document.add_paragraph(spec.note)
            output_path = DOCX_DIR / f"{index:02d} - {safe_name(spec.output_title)}.docx"
            document.save(output_path)
            unit_record["outputPath"] = rel_posix(output_path, PROJECT_ROOT)
            self.audit["requestedUnits"].append(unit_record)
            return

        if spec.note:
            note_para = document.add_paragraph()
            note_para.add_run("Source note: ").bold = True
            note_para.add_run(spec.note)
        source_heading = document.add_paragraph()
        source_heading.add_run("Source manifest title: ").bold = True
        source_heading.add_run(item_title(source_item))

        items = [item for item in walk_items(source_item) if item is not source_item]
        if source_kind == "matching-topic":
            items = [source_item]

        self.add_manifest_outline(document, items)
        for item in items:
            title = item_title(item)
            level = self.item_depth(source_item, item) + 1 if source_kind == "top-level-module" else 1
            self.render_item(document, item, level, spec, unit_record)

        output_path = DOCX_DIR / f"{index:02d} - {safe_name(spec.output_title)}.docx"
        document.save(output_path)
        unit_record["outputPath"] = rel_posix(output_path, PROJECT_ROOT)
        unit_record["outputBytes"] = output_path.stat().st_size
        self.audit["requestedUnits"].append(unit_record)
        self.audit["docxFiles"].append(str(output_path))

    def new_document(self) -> Document:
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        styles = document.styles
        styles["Normal"].font.name = "Aptos"
        styles["Normal"].font.size = Pt(10.5)
        for style_name in ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
            styles[style_name].font.name = "Aptos Display"
        return document

    def item_depth(self, root: ET.Element, target: ET.Element, depth: int = 0) -> int:
        if root is target:
            return depth
        for child in item_children(root):
            found = self.item_depth(child, target, depth + 1)
            if found >= 0:
                return found
        return -1

    def add_manifest_outline(self, document: Document, items: list[ET.Element]):
        if not items:
            return
        document.add_heading("Contents", level=1)
        for item in items:
            title = item_title(item) or "Untitled item"
            para = document.add_paragraph(style="List Bullet")
            para.add_run(title)

    def render_item(
        self,
        document: Document,
        item: ET.Element,
        level: int,
        spec: UnitSpec,
        unit_record: dict[str, Any],
    ):
        title = item_title(item) or "Untitled item"
        heading_level = min(max(level, 1), 4)
        document.add_heading(title, level=heading_level)
        ref = item.attrib.get("identifierref")
        record: dict[str, Any] = {
            "title": title,
            "identifier": item.attrib.get("identifier"),
            "identifierref": ref,
            "files": [],
        }
        unit_record["includedItems"].append(record)

        if not ref:
            if not item_children(item):
                document.add_paragraph("No direct resource is attached to this manifest item.")
            return
        resource = self.resources.get(ref)
        if not resource:
            self.audit["missingResources"].append({"title": title, "identifierref": ref})
            document.add_paragraph(f"Missing manifest resource: {ref}")
            return

        files = resource.get("files") or []
        if not files:
            document.add_paragraph("This manifest resource does not list a file.")
            return

        for href in files:
            ext = Path(href).suffix.lower()
            record["files"].append(href)
            if not package_path_exists(self.zip_file, href):
                self.audit["missingResources"].append({"title": title, "href": href})
                document.add_paragraph(f"Missing source file: {href}")
                continue
            if ext in [".html", ".htm"]:
                self.render_html(document, href)
            elif ext == ".pdf":
                support_path = self.copy_support_file(href, spec.output_title, title)
                self.add_support_reference(document, href, support_path, "Original PDF resource")
                self.render_pdf(document, href)
            elif ext == ".docx":
                support_path = self.copy_support_file(href, spec.output_title, title)
                self.add_support_reference(document, href, support_path, "Original DOCX resource")
                self.render_docx_text(document, href)
            elif ext in [".png", ".jpg", ".jpeg", ".gif"]:
                self.add_image_from_href(document, href)
            else:
                support_path = self.copy_support_file(href, spec.output_title, title)
                label = f"Original {ext.lstrip('.').upper() or 'resource'} file"
                self.add_support_reference(document, href, support_path, label)

    def render_html(self, document: Document, href: str):
        html_bytes = self.zip_file.read(href)
        soup = BeautifulSoup(html_bytes, "lxml")
        for tag in soup(["script", "style", "meta", "link", "title"]):
            tag.decompose()
        body = soup.body or soup
        for child in body.children:
            self.render_block(document, child, href)

    def render_block(self, document: Document, node: Any, base_href: str):
        if isinstance(node, NavigableString):
            text = clean_text(str(node))
            if text:
                document.add_paragraph(text)
            return
        if not isinstance(node, Tag):
            return
        name = (node.name or "").lower()
        if name in ["script", "style", "meta", "link", "title"]:
            return
        if name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            text = clean_text(node.get_text(" "))
            if text:
                document.add_heading(text, level=min(int(name[1]), 4))
            return
        if name == "img":
            self.render_html_image(document, node, base_href)
            return
        if name in ["ul", "ol"]:
            style = "List Number" if name == "ol" else "List Bullet"
            for li in node.find_all("li", recursive=False):
                para = document.add_paragraph(style=style)
                self.add_inline(para, li, base_href)
            return
        if name == "table":
            self.render_table(document, node)
            return
        if name in ["iframe", "video", "audio", "embed", "object"]:
            src = node.get("src") or node.get("data") or ""
            para = document.add_paragraph()
            para.add_run("Embedded media reference: ").bold = True
            para.add_run(src or clean_text(node.get_text(" ")) or "No source listed")
            return
        block_children = [
            child
            for child in node.children
            if isinstance(child, Tag)
            and (child.name or "").lower()
            in [
                "p",
                "div",
                "section",
                "article",
                "h1",
                "h2",
                "h3",
                "h4",
                "h5",
                "h6",
                "ul",
                "ol",
                "table",
                "img",
                "iframe",
                "video",
                "audio",
                "embed",
                "object",
            ]
        ]
        if name in ["p", "blockquote"] or not block_children:
            para = document.add_paragraph()
            if name == "blockquote":
                para.style = "Intense Quote"
            self.add_inline(para, node, base_href)
            if not clean_text(para.text) and not para.runs:
                self.remove_paragraph(para)
            return
        for child in node.children:
            self.render_block(document, child, base_href)

    def add_inline(self, paragraph, node: Any, base_href: str, bold=False, italic=False, underline=False):
        if isinstance(node, NavigableString):
            text = str(node)
            if text:
                run = paragraph.add_run(text)
                run.bold = bold
                run.italic = italic
                run.underline = underline
            return
        if not isinstance(node, Tag):
            return
        name = (node.name or "").lower()
        if name == "br":
            paragraph.add_run().add_break()
            return
        if name == "img":
            self.render_inline_html_image(paragraph, node, base_href)
            return
        if name == "a":
            link_text = clean_text(node.get_text(" ")) or node.get("href") or ""
            href = node.get("href") or ""
            if href:
                add_hyperlink(paragraph, link_text, href)
            else:
                paragraph.add_run(link_text)
            return
        if name in ["strong", "b"]:
            bold = True
        if name in ["em", "i"]:
            italic = True
        if name == "u":
            underline = True
        if name in ["iframe", "video", "audio", "embed", "object"]:
            src = node.get("src") or node.get("data") or ""
            paragraph.add_run(f"[Embedded media reference: {src}]")
            return
        for child in node.children:
            self.add_inline(paragraph, child, base_href, bold=bold, italic=italic, underline=underline)

    def render_table(self, document: Document, table_node: Tag):
        rows = table_node.find_all("tr")
        if not rows:
            return
        max_cols = max(len(row.find_all(["th", "td"], recursive=False)) for row in rows)
        if max_cols == 0:
            return
        table = document.add_table(rows=len(rows), cols=max_cols)
        table.style = "Table Grid"
        for row_index, row_node in enumerate(rows):
            cells = row_node.find_all(["th", "td"], recursive=False)
            for col_index, cell_node in enumerate(cells):
                cell = table.cell(row_index, col_index)
                cell.text = clean_text(cell_node.get_text(" "))

    def render_html_image(self, document: Document, img: Tag, base_href: str):
        src = img.get("src")
        if not src:
            return
        package_href = resolve_package_href(base_href, src, self.zip_file)
        if not package_href:
            self.audit["unresolvedHtmlAssets"].append({"base": base_href, "src": src})
            para = document.add_paragraph()
            para.add_run("Unresolved image reference: ").bold = True
            para.add_run(src)
            return
        self.add_image_from_href(document, package_href)

    def render_inline_html_image(self, paragraph, img: Tag, base_href: str):
        src = img.get("src")
        if not src:
            return
        package_href = resolve_package_href(base_href, src, self.zip_file)
        if not package_href:
            self.audit["unresolvedHtmlAssets"].append({"base": base_href, "src": src})
            paragraph.add_run(f"[Unresolved image: {src}]")
            return
        try:
            image_bytes = self.zip_file.read(package_href)
            normalized_bytes, width_px, height_px = self.normalized_image_bytes(image_bytes)
            if width_px <= 0 or height_px <= 0:
                raise ValueError("invalid image size")
            width_inches = min(3.5, width_px / 140)
            paragraph.add_run().add_picture(BytesIO(normalized_bytes), width=Inches(width_inches))
            self.audit["embeddedImages"] += 1
        except Exception as exc:
            paragraph.add_run(f"[Image could not be embedded: {package_href} ({exc})]")
            if package_path_exists(self.zip_file, package_href):
                self.copy_support_file(package_href, self.current_output_title, "Unembedded image")
            self.audit["unresolvedHtmlAssets"].append({"href": package_href, "error": str(exc)})

    def add_image_from_href(self, document: Document, href: str):
        try:
            image_bytes = self.zip_file.read(href)
            normalized_bytes, width_px, height_px = self.normalized_image_bytes(image_bytes)
            if width_px <= 0 or height_px <= 0:
                raise ValueError("invalid image size")
            max_width = 6.3
            width_inches = min(max_width, width_px / 120)
            document.add_picture(BytesIO(normalized_bytes), width=Inches(width_inches))
            self.audit["embeddedImages"] += 1
        except Exception as exc:
            para = document.add_paragraph()
            para.add_run("Image could not be embedded: ").bold = True
            para.add_run(f"{href} ({exc})")
            if package_path_exists(self.zip_file, href):
                support_path = self.copy_support_file(href, self.current_output_title, "Unembedded image")
                para.add_run(" Preserved original at: ")
                para.add_run(rel_posix(support_path, DOCX_DIR))
            self.audit["unresolvedHtmlAssets"].append({"href": href, "error": str(exc)})

    def normalized_image_bytes(self, image_bytes: bytes) -> tuple[bytes, int, int]:
        image = Image.open(BytesIO(image_bytes))
        image.load()
        if image.mode not in ("RGB", "RGBA"):
            image = image.convert("RGBA" if "A" in image.getbands() else "RGB")
        output = BytesIO()
        image.save(output, format="PNG", dpi=(144, 144))
        return output.getvalue(), image.size[0], image.size[1]

    def render_pdf(self, document: Document, href: str):
        try:
            pdf = fitz.open(stream=self.zip_file.read(href), filetype="pdf")
        except Exception as exc:
            document.add_paragraph(f"PDF could not be opened for rendering: {exc}")
            return
        for page_index in range(pdf.page_count):
            page = pdf.load_page(page_index)
            pixmap = page.get_pixmap(matrix=fitz.Matrix(1.25, 1.25), alpha=False)
            png_bytes = pixmap.tobytes("png")
            document.add_paragraph(f"PDF page {page_index + 1} of {pdf.page_count}")
            document.add_picture(BytesIO(png_bytes), width=Inches(6.4))
            self.audit["renderedPdfPages"] += 1
        pdf.close()

    def render_docx_text(self, document: Document, href: str):
        try:
            source = Document(BytesIO(self.zip_file.read(href)))
        except Exception as exc:
            document.add_paragraph(f"DOCX could not be opened for inline extraction: {exc}")
            return
        document.add_paragraph("Readable text extracted from the source DOCX follows.")
        for paragraph in source.paragraphs:
            text = clean_text(paragraph.text)
            if text:
                document.add_paragraph(text)
        for table in source.tables:
            output_table = document.add_table(rows=len(table.rows), cols=len(table.columns))
            output_table.style = "Table Grid"
            for r_index, row in enumerate(table.rows):
                for c_index, cell in enumerate(row.cells):
                    output_table.cell(r_index, c_index).text = clean_text(cell.text)

    def copy_support_file(self, href: str, unit_title: str, item_name: str) -> Path:
        if href in self._copied_support:
            return self._copied_support[href]
        ext = Path(href).suffix or ".resource"
        unit_dir = SUPPORT_DIR / safe_name(unit_title, 48)
        unit_dir.mkdir(parents=True, exist_ok=True)
        digest = hashlib.sha1(href.encode("utf-8")).hexdigest()[:8]
        basename = f"{safe_name(item_name, 58)}-{digest}{ext}"
        destination = unit_dir / basename
        counter = 2
        while destination.exists():
            destination = unit_dir / f"{Path(basename).stem}-{counter}{ext}"
            counter += 1
        with self.zip_file.open(href) as src, destination.open("wb") as dest:
            shutil.copyfileobj(src, dest, length=1024 * 1024)
        self._copied_support[href] = destination
        self.audit["copiedSupportFiles"].append(
            {
                "sourceHref": href,
                "outputPath": rel_posix(destination, PROJECT_ROOT),
                "bytes": destination.stat().st_size,
            }
        )
        return destination

    def add_support_reference(self, document: Document, href: str, support_path: Path, label: str):
        para = document.add_paragraph()
        para.add_run(f"{label}: ").bold = True
        relative = rel_posix(support_path, DOCX_DIR)
        add_hyperlink(para, relative, relative)
        details = document.add_paragraph()
        details.add_run("Original package path: ").bold = True
        details.add_run(href)

    def remove_paragraph(self, paragraph):
        element = paragraph._element
        element.getparent().remove(element)
        paragraph._p = paragraph._element = None

    def write_resource_indexes(self):
        rows = []
        for entry in self.audit["copiedSupportFiles"]:
            rows.append(entry)
        csv_path = EXPORT_DIR / "supporting-files-index.csv"
        with csv_path.open("w", encoding="utf-8", newline="") as handle:
            writer = csv.DictWriter(handle, fieldnames=["sourceHref", "outputPath", "bytes"])
            writer.writeheader()
            writer.writerows(rows)

        map_payload = {
            "requestedUnits": self.audit["requestedUnits"],
            "excludedTopLevelModules": self.audit["excludedTopLevelModules"],
            "sourceTopLevelModules": [
                {
                    "title": item_title(module),
                    "resourceItemCount": len(
                        [item for item in walk_items(module) if item.attrib.get("identifierref")]
                    ),
                }
                for module in self.top_level_modules()
            ],
        }
        (META_DIR / "docx-conversion-map.json").write_text(
            json.dumps(map_payload, indent=2), encoding="utf-8"
        )

    def verify_outputs(self):
        verification = {
            "expectedDocxCount": len(REQUESTED_UNITS),
            "actualDocxCount": 0,
            "invalidDocx": [],
            "docx": [],
        }
        for path in sorted(DOCX_DIR.glob("*.docx")):
            verification["actualDocxCount"] += 1
            record = {"path": rel_posix(path, PROJECT_ROOT), "bytes": path.stat().st_size}
            try:
                with zipfile.ZipFile(path) as docx_zip:
                    docx_zip.getinfo("word/document.xml")
            except Exception as exc:
                record["error"] = str(exc)
                verification["invalidDocx"].append(record)
            verification["docx"].append(record)
        verification["passed"] = (
            verification["expectedDocxCount"] == verification["actualDocxCount"]
            and not verification["invalidDocx"]
        )
        self.audit["verification"] = verification
        (META_DIR / "docx-export-verification.json").write_text(
            json.dumps(verification, indent=2), encoding="utf-8"
        )
        if not verification["passed"]:
            raise SystemExit("DOCX verification failed. See meta/docx-export-verification.json")

    def write_audit_markdown(self):
        support_bytes = sum(entry["bytes"] for entry in self.audit["copiedSupportFiles"])
        lines = [
            "# English Language Arts 10-2 DOCX Export Audit",
            "",
            f"- Generated: {self.audit['generatedAt']}",
            f"- Source ZIP: `{self.audit['sourceZip']}`",
            f"- Source ZIP size: {self.audit['sourceZipBytes'] / 1024 / 1024:.2f} MB",
            f"- DOCX files generated: {self.audit['verification']['actualDocxCount']}",
            f"- Supporting files copied: {len(self.audit['copiedSupportFiles'])}",
            f"- Supporting file size: {support_bytes / 1024 / 1024:.2f} MB",
            f"- Embedded HTML images: {self.audit['embeddedImages']}",
            f"- Rendered PDF pages: {self.audit['renderedPdfPages']}",
            "",
            "## Generated DOCX Files",
            "",
        ]
        for docx_record in self.audit["verification"]["docx"]:
            lines.append(f"- `{docx_record['path']}` ({docx_record['bytes'] / 1024 / 1024:.2f} MB)")
        lines.extend(["", "## Screenshot / Manifest Title Notes", ""])
        for unit in self.audit["requestedUnits"]:
            if unit.get("note"):
                lines.append(
                    f"- `{unit['outputTitle']}` uses source `{unit.get('sourceTitle')}`: {unit['note']}"
                )
        lines.extend(["", "## Source Top-Level Modules Not Exported As Requested Units", ""])
        for module in self.audit["excludedTopLevelModules"]:
            lines.append(f"- `{module['title']}` ({module['itemCount']} resource item(s))")
        if self.audit["unresolvedHtmlAssets"]:
            lines.extend(["", "## Unresolved HTML Assets", ""])
            for asset in self.audit["unresolvedHtmlAssets"]:
                lines.append(f"- `{asset}`")
        if self.audit["missingResources"]:
            lines.extend(["", "## Missing Manifest Resources", ""])
            for missing in self.audit["missingResources"]:
                lines.append(f"- `{missing}`")
        lines.append("")
        (META_DIR / "docx-export-audit.md").write_text("\n".join(lines), encoding="utf-8")


def main():
    if not ZIP_PATH.exists():
        raise SystemExit(f"Source ZIP not found: {ZIP_PATH}")
    with zipfile.ZipFile(ZIP_PATH) as zip_file:
        manifest_root = ET.fromstring(zip_file.read("imsmanifest.xml"))
        exporter = DocxExporter(zip_file, manifest_root)
        exporter.build_all()


if __name__ == "__main__":
    main()
