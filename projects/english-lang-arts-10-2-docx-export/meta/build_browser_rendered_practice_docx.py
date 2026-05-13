from __future__ import annotations

import csv
import hashlib
import http.server
import importlib.util
import json
import os
import posixpath
import re
import shutil
import socketserver
import subprocess
import sys
import threading
import urllib.parse
import zipfile
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any
from urllib.parse import parse_qs, urlparse
from xml.etree import ElementTree as ET

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image


REPO_ROOT = Path(__file__).resolve().parents[3]
PROJECT_ROOT = REPO_ROOT / "projects" / "english-lang-arts-10-2-docx-export"
ENGLISH_BUILDER_PATH = PROJECT_ROOT / "meta" / "build_docx_export.py"

SOURCE_ZIP_NAME = "D2LCCExport_149674_25-26 _ S2 _ English Lang. Arts 10-2 _ Per 1(A) _ _202651230.zip"
PRACTICE_UNIT_TITLE = "Unit 1: Introduction to Interpreting and Creating Texts"
OUTPUT_TITLE = "Unit 1 Introduction to Interpreting and Creating Texts - browser rendered practice"

META_DIR = PROJECT_ROOT / "meta"
EXPORT_DIR = PROJECT_ROOT / "exports"
DOCX_DIR = EXPORT_DIR / "browser-practice-docx"
WORK_DIR = EXPORT_DIR / "browser-practice-work"
SCREENSHOT_DIR = EXPORT_DIR / "browser-practice-screenshots"
SUPPORT_DIR = EXPORT_DIR / "browser-practice-supporting-files"

VIEWPORT_WIDTH = 1280
VIEWPORT_HEIGHT = 900
CONTENT_WIDTH = 920
DOCX_IMAGE_WIDTH_IN = 7.2
DOCX_IMAGE_HEIGHT_IN = 9.35


def first_existing_path(env_var: str, candidates: list[Path]) -> Path:
    override = os.environ.get(env_var)
    paths = ([Path(override)] if override else []) + candidates
    for path in paths:
        if path.exists():
            return path
    return paths[0]


ZIP_PATH = first_existing_path(
    "ELA102_SOURCE_ZIP",
    [
        Path.home() / "Downloads" / SOURCE_ZIP_NAME,
        Path("/Users/deanguedo/Downloads") / SOURCE_ZIP_NAME,
    ],
)


def load_module(name: str, path: Path) -> Any:
    spec = importlib.util.spec_from_file_location(name, path)
    if spec is None or spec.loader is None:
        raise SystemExit(f"Unable to load builder module: {path}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


english_builder = load_module("english_docx_builder_browser_practice", ENGLISH_BUILDER_PATH)


@dataclass
class RenderItem:
    index: int
    title: str
    identifier: str | None
    identifierref: str | None
    html_hrefs: list[str]
    support_hrefs: list[str]


def local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1]


def item_title(item: ET.Element) -> str:
    for child in item:
        if local_name(child.tag) == "title":
            return "".join(child.itertext()).strip()
    return ""


def item_children(item: ET.Element) -> list[ET.Element]:
    return [child for child in item if local_name(child.tag) == "item"]


def walk_items(item: ET.Element) -> list[ET.Element]:
    result = [item]
    for child in item_children(item):
        result.extend(walk_items(child))
    return result


def normalize_key(value: str) -> str:
    return re.sub(r"\s+", " ", value.strip()).casefold()


def safe_name(value: str, max_len: int = 96) -> str:
    return english_builder.safe_name(value, max_len=max_len)


def rel_posix(path: Path, start: Path) -> str:
    return Path(os.path.relpath(path, start)).as_posix()


def clean_text(value: str) -> str:
    return re.sub(r"\s+", " ", value.replace("\xa0", " ")).strip()


def package_path_exists(zip_file: zipfile.ZipFile, href: str) -> bool:
    try:
        zip_file.getinfo(href)
        return True
    except KeyError:
        return False


def add_hyperlink(paragraph: Any, text: str, url: str) -> None:
    english_builder.add_hyperlink(paragraph, text, url)


def public_video_url(src: str) -> str:
    parsed = urlparse(src)
    host = parsed.netloc.casefold()
    if "youtube" in host:
        query = parse_qs(parsed.query)
        video_id = (query.get("v") or [""])[0]
        if not video_id and "/embed/" in parsed.path:
            video_id = parsed.path.split("/embed/", 1)[1].split("/", 1)[0]
        if video_id:
            return f"https://www.youtube.com/watch?v={video_id}"
    if "youtu.be" in host:
        video_id = parsed.path.strip("/").split("/", 1)[0]
        if video_id:
            return f"https://www.youtube.com/watch?v={video_id}"
    if "ted.com" in host:
        return src
    return src


def is_video_url(url: str) -> bool:
    if not url:
        return False
    parsed = urlparse(url)
    host = parsed.netloc.casefold()
    path = parsed.path.casefold()
    return (
        "youtube" in host
        or "youtu.be" in host
        or "ted.com" in host
        or path.endswith((".mp4", ".mov", ".m4v", ".webm"))
    )


def youtube_thumbnail_url(url: str) -> str | None:
    parsed = urlparse(url)
    host = parsed.netloc.casefold()
    video_id = ""
    if "youtube" in host:
        query = parse_qs(parsed.query)
        video_id = (query.get("v") or [""])[0]
        if not video_id and "/embed/" in parsed.path:
            video_id = parsed.path.split("/embed/", 1)[1].split("/", 1)[0]
    elif "youtu.be" in host:
        video_id = parsed.path.strip("/").split("/", 1)[0]
    if not video_id:
        return None
    return f"https://img.youtube.com/vi/{video_id}/hqdefault.jpg"


def clean_lms_url(href: str) -> str:
    if not href:
        return ""
    if "d2lsessionval=" in href.casefold():
        return href.split("?", 1)[0]
    return href


def is_noise_text(text: str) -> bool:
    return normalize_key(text) in {
        "",
        "image source",
        "image sources",
        "source",
        "template javascript",
    }


def ensure_clean_dirs() -> None:
    for target in [DOCX_DIR, WORK_DIR, SCREENSHOT_DIR, SUPPORT_DIR]:
        if target.exists():
            shutil.rmtree(target)
        target.mkdir(parents=True, exist_ok=True)
    META_DIR.mkdir(parents=True, exist_ok=True)


def read_resources(manifest_root: ET.Element) -> dict[str, dict[str, Any]]:
    resources: dict[str, dict[str, Any]] = {}
    for resource in manifest_root.iter():
        if local_name(resource.tag) != "resource":
            continue
        identifier = resource.attrib.get("identifier")
        if not identifier:
            continue
        files = [
            file_node.attrib.get("href")
            for file_node in resource
            if local_name(file_node.tag) == "file" and file_node.attrib.get("href")
        ]
        resources[identifier] = {
            "identifier": identifier,
            "type": resource.attrib.get("type"),
            "files": files,
        }
    return resources


def top_level_modules(manifest_root: ET.Element) -> list[ET.Element]:
    organization = next(
        node for node in manifest_root.iter() if local_name(node.tag) == "organization"
    )
    root_items = item_children(organization)
    if len(root_items) == 1 and not item_title(root_items[0]):
        return item_children(root_items[0])
    return root_items


def find_top_level(manifest_root: ET.Element, title: str) -> ET.Element | None:
    wanted = normalize_key(title)
    for item in top_level_modules(manifest_root):
        if normalize_key(item_title(item)) == wanted:
            return item
    return None


def collect_unit_items(manifest_root: ET.Element, resources: dict[str, dict[str, Any]]) -> list[RenderItem]:
    source_item = find_top_level(manifest_root, PRACTICE_UNIT_TITLE)
    if source_item is None:
        raise SystemExit(f"Unit not found in manifest: {PRACTICE_UNIT_TITLE}")
    render_items: list[RenderItem] = []
    for item in walk_items(source_item):
        if item is source_item:
            continue
        ref = item.attrib.get("identifierref")
        if not ref:
            continue
        files = list(resources.get(ref, {}).get("files") or [])
        html_hrefs = [href for href in files if Path(href).suffix.lower() in [".html", ".htm"]]
        support_hrefs = [href for href in files if href not in html_hrefs]
        render_items.append(
            RenderItem(
                index=len(render_items) + 1,
                title=item_title(item) or "Untitled item",
                identifier=item.attrib.get("identifier"),
                identifierref=ref,
                html_hrefs=html_hrefs,
                support_hrefs=support_hrefs,
            )
        )
    return render_items


def extract_related_files(zip_file: zipfile.ZipFile, href: str) -> list[str]:
    prefix = posixpath.dirname(href).rstrip("/") + "/"
    extracted: list[str] = []
    for name in zip_file.namelist():
        if name.startswith(prefix) and not name.endswith("/"):
            destination = WORK_DIR / Path(name)
            destination.parent.mkdir(parents=True, exist_ok=True)
            destination.write_bytes(zip_file.read(name))
            extracted.append(name)
    return extracted


def remove_noise_nodes(soup: BeautifulSoup) -> None:
    for tag in soup(["script", "noscript"]):
        tag.decompose()
    for node in list(soup.find_all(string=True)):
        text = clean_text(str(node))
        if is_noise_text(text):
            node.extract()
    for anchor in list(soup.find_all("a")):
        if is_noise_text(anchor.get_text(" ")):
            anchor.decompose()
    for tag in soup.find_all(True):
        for attr in list(tag.attrs):
            if attr.casefold().startswith("data-d2l"):
                del tag.attrs[attr]


def normalize_links_and_media(soup: BeautifulSoup, media_records: list[dict[str, Any]], source_href: str) -> None:
    for media in list(soup.find_all(["iframe", "video", "audio", "embed", "object"])):
        src = (media.get("src") or media.get("data") or "").strip()
        if not src:
            media.decompose()
            continue
        handoff = public_video_url(src)
        thumb = youtube_thumbnail_url(handoff)
        card = soup.new_tag("div")
        card["class"] = "docx-video-card"
        link = soup.new_tag("a", href=handoff)
        if thumb:
            image = soup.new_tag("img", src=thumb, alt="Video preview")
            image["class"] = "docx-video-thumb"
            link.append(image)
        else:
            play = soup.new_tag("span")
            play["class"] = "docx-video-fallback"
            play.string = "Video preview"
            link.append(play)
        url_line = soup.new_tag("p")
        url_link = soup.new_tag("a", href=handoff)
        url_link.string = handoff
        url_line.append(url_link)
        card.append(link)
        card.append(url_line)
        media.replace_with(card)
        media_records.append({"sourceHtml": source_href, "src": src, "handoffUrl": handoff, "thumbnailUrl": thumb})

    for anchor in soup.find_all("a"):
        href = (anchor.get("href") or "").strip()
        if not href:
            continue
        if is_video_url(href):
            handoff = public_video_url(href)
            anchor["href"] = handoff
            anchor.string = handoff
            media_records.append({"sourceHtml": source_href, "src": href, "handoffUrl": handoff, "thumbnailUrl": youtube_thumbnail_url(handoff)})
        else:
            anchor["href"] = clean_lms_url(href)


def wrap_body(soup: BeautifulSoup, title: str) -> None:
    body = soup.body
    if body is None:
        body = soup.new_tag("body")
        for child in list(soup.contents):
            body.append(child.extract())
        soup.append(body)
    wrapper = soup.new_tag("main")
    wrapper["class"] = "docx-content-root"
    for child in list(body.contents):
        wrapper.append(child.extract())
    body.append(wrapper)
    head = soup.head
    if head is None:
        head = soup.new_tag("head")
        if soup.html:
            soup.html.insert(0, head)
        else:
            soup.insert(0, head)
    title_tag = soup.new_tag("title")
    title_tag.string = title
    head.append(title_tag)
    css = soup.new_tag("link", rel="stylesheet", href="/__docx_brightspace_shim.css")
    head.append(css)


def prepare_html_page(zip_file: zipfile.ZipFile, item: RenderItem, href: str) -> dict[str, Any]:
    extract_related_files(zip_file, href)
    raw = zip_file.read(href).decode("utf-8", errors="replace")
    soup = BeautifulSoup(raw, "lxml")
    media_records: list[dict[str, Any]] = []
    remove_noise_nodes(soup)
    normalize_links_and_media(soup, media_records, href)
    wrap_body(soup, item.title)
    output_path = WORK_DIR / Path(href)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(str(soup), encoding="utf-8")
    return {
        "itemIndex": item.index,
        "itemTitle": item.title,
        "sourceHref": href,
        "renderHref": Path(href).as_posix(),
        "screenshot": f"{item.index:03d}-{safe_name(item.title, 52)}-{hashlib.sha1(href.encode('utf-8')).hexdigest()[:8]}.png",
        "media": media_records,
    }


def write_css_shim() -> None:
    css = f"""
@import url('https://s.brightspace.com/lib/fonts/0.6.1/fonts.css');
html, body {{ margin: 0; padding: 0; background: #fff; color: #202122; }}
body {{ font-family: Lato, Arial, sans-serif; font-size: 19px; line-height: 1.42; }}
.docx-content-root {{ width: {CONTENT_WIDTH}px; max-width: {CONTENT_WIDTH}px; margin: 0 auto; padding: 46px 44px 58px; box-sizing: border-box; background: #fff; }}
.container-fluid {{ width: 100%; margin: 0 auto; padding: 0; }}
.row {{ display: flex; flex-wrap: wrap; gap: 30px; align-items: stretch; }}
.col-sm-10, .offset-sm-1, .col-12 {{ width: 100%; max-width: 100%; }}
h1, h2, h3, h4 {{ font-family: Lato, Arial, sans-serif; color: #333c48; line-height: 1.18; margin: 0.7em 0 0.55em; font-weight: 700; }}
h1 {{ font-size: 36px; text-align: center; }}
h2 {{ font-size: 27px; text-align: center; color: #00847f; }}
h3 {{ font-size: 23px; }}
h4 {{ font-size: 21px; }}
p {{ margin: 0 0 1.05em; }}
ul, ol {{ margin-top: 0.4em; margin-bottom: 1.05em; }}
li {{ margin: 0.42em 0; }}
img {{ max-width: 100%; height: auto; }}
hr {{ border: 0; border-top: 1px solid #d7dce2; margin: 30px 0; }}
a {{ color: #006fbf; text-decoration: underline; }}
.two-col-panels {{ margin: 20px 0; }}
.two-col-panels .row {{ display: grid; grid-template-columns: 1fr 1fr; gap: 30px; }}
.card {{ border: 1px solid #d6dbe2; border-radius: 4px; background: #f8f9fa; box-sizing: border-box; }}
.card.bg-light {{ min-height: 205px; display: flex; align-items: center; justify-content: center; text-align: center; padding: 26px 24px; }}
.card.bg-light h4 {{ margin-top: 0; margin-bottom: 8px; font-size: 20px; text-align: center; }}
.card.bg-light p {{ margin-bottom: 0; font-size: 19px; }}
.card-graphic {{ border: 3px solid #00847f; background: #fff; margin: 26px auto; padding: 30px; text-align: center; }}
.card-body {{ padding: 0; }}
.card-text p:last-child {{ margin-bottom: 0; }}
.small {{ font-size: 80%; }}
.docx-video-card {{ margin: 22px auto; padding: 12px; border: 2px solid #222; background: #eef8ee; text-align: left; }}
.docx-video-card a {{ word-break: break-all; }}
.docx-video-thumb {{ display: block; width: 100%; max-width: 760px; margin: 0 auto 10px; }}
.docx-video-fallback {{ display: block; min-height: 260px; background: linear-gradient(140deg, #183240, #386f80); color: #fff; text-align: center; padding-top: 130px; box-sizing: border-box; }}
table {{ border-collapse: collapse; width: 100%; }}
""".strip()
    (WORK_DIR / "__docx_brightspace_shim.css").write_text(css, encoding="utf-8")


def write_render_helper() -> Path:
    helper = WORK_DIR / "__render_pages.mjs"
    helper.write_text(
        r"""
import { chromium } from 'playwright';
import fs from 'node:fs/promises';

const [, , pagesPath, baseUrl, outDir] = process.argv;
const pages = JSON.parse(await fs.readFile(pagesPath, 'utf8'));
const browser = await chromium.launch({ headless: true });
const context = await browser.newContext({
  viewport: { width: 1280, height: 900 },
  deviceScaleFactor: 1
});
const page = await context.newPage();
const results = [];

for (const record of pages) {
  const url = `${baseUrl}/${record.renderHref.split('/').map(encodeURIComponent).join('/')}`;
  const output = `${outDir}/${record.screenshot}`;
  await page.goto(url, { waitUntil: 'networkidle', timeout: 45000 });
  await page.evaluate(() => document.fonts && document.fonts.ready);
  const locator = page.locator('.docx-content-root').first();
  await locator.screenshot({ path: output, timeout: 45000 });
  const box = await locator.boundingBox();
  results.push({ ...record, screenshotPath: output, width: Math.round(box?.width || 0), height: Math.round(box?.height || 0) });
}

await browser.close();
console.log(JSON.stringify(results, null, 2));
""".strip(),
        encoding="utf-8",
    )
    return helper


class QuietHandler(http.server.SimpleHTTPRequestHandler):
    def log_message(self, format: str, *args: Any) -> None:
        return


def serve_workdir() -> tuple[socketserver.TCPServer, str]:
    handler = lambda *args, **kwargs: QuietHandler(*args, directory=str(WORK_DIR), **kwargs)
    httpd = socketserver.TCPServer(("127.0.0.1", 0), handler)
    port = httpd.server_address[1]
    thread = threading.Thread(target=httpd.serve_forever, daemon=True)
    thread.start()
    return httpd, f"http://127.0.0.1:{port}"


def render_pages_with_playwright(pages: list[dict[str, Any]]) -> list[dict[str, Any]]:
    pages_json = WORK_DIR / "__pages.json"
    pages_json.write_text(json.dumps(pages, indent=2), encoding="utf-8")
    helper = write_render_helper()
    httpd, base_url = serve_workdir()
    try:
        command = ["node", str(helper), str(pages_json), base_url, str(SCREENSHOT_DIR)]
        result = subprocess.run(command, cwd=REPO_ROOT, text=True, capture_output=True, check=False)
    finally:
        httpd.shutdown()
        httpd.server_close()
    if result.returncode != 0:
        raise SystemExit(f"Playwright render failed.\nSTDOUT:\n{result.stdout}\nSTDERR:\n{result.stderr}")
    return json.loads(result.stdout)


def copy_support_file(zip_file: zipfile.ZipFile, href: str, item_title_value: str) -> dict[str, Any]:
    digest = hashlib.sha1(href.encode("utf-8")).hexdigest()[:8]
    ext = Path(href).suffix or ".resource"
    basename = f"{safe_name(item_title_value, 58)}-{digest}{ext}"
    destination = SUPPORT_DIR / basename
    destination.parent.mkdir(parents=True, exist_ok=True)
    destination.write_bytes(zip_file.read(href))
    return {"sourceHref": href, "outputPath": rel_posix(destination, PROJECT_ROOT), "bytes": destination.stat().st_size}


def split_image_for_docx(image_path: Path, stem: str) -> list[Path]:
    image = Image.open(image_path)
    image.load()
    max_chunk_height = int(image.width * (DOCX_IMAGE_HEIGHT_IN / DOCX_IMAGE_WIDTH_IN))
    if image.height <= max_chunk_height:
        return [image_path]
    chunks: list[Path] = []
    y = 0
    index = 1
    while y < image.height:
        bottom = min(y + max_chunk_height, image.height)
        chunk = image.crop((0, y, image.width, bottom))
        chunk_path = image_path.with_name(f"{stem}-part-{index:02d}.png")
        chunk.save(chunk_path)
        chunks.append(chunk_path)
        y = bottom
        index += 1
    return chunks


def build_docx(rendered_pages: list[dict[str, Any]], support_records: list[dict[str, Any]]) -> Path:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    styles = document.styles
    styles["Normal"].font.name = "Lato"
    styles["Normal"].font.size = Pt(10)

    for page_index, record in enumerate(rendered_pages):
        if page_index:
            document.add_page_break()
        screenshot_path = Path(record["screenshotPath"])
        chunks = split_image_for_docx(screenshot_path, screenshot_path.stem)
        for chunk_index, chunk_path in enumerate(chunks):
            if chunk_index:
                document.add_page_break()
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run().add_picture(str(chunk_path), width=Inches(DOCX_IMAGE_WIDTH_IN))
        media = record.get("media") or []
        if media:
            document.add_paragraph()
            heading = document.add_paragraph()
            heading.add_run("Video links for Google Docs:").bold = True
            for media_record in media:
                url = media_record.get("handoffUrl")
                if not url:
                    continue
                para = document.add_paragraph()
                add_hyperlink(para, url, url)

    if support_records:
        document.add_section(WD_SECTION_START.NEW_PAGE)
        document.add_heading("Supporting files copied from the Brightspace package", level=1)
        for record in support_records:
            para = document.add_paragraph()
            para.add_run(Path(record["outputPath"]).name + ": ").bold = True
            add_hyperlink(para, record["outputPath"], record["outputPath"])

    output_path = DOCX_DIR / f"01 - {safe_name(OUTPUT_TITLE)}.docx"
    document.save(output_path)
    return output_path


def main() -> None:
    if not ZIP_PATH.exists():
        raise SystemExit(f"Source ZIP not found: {ZIP_PATH}")
    ensure_clean_dirs()
    write_css_shim()

    audit: dict[str, Any] = {
        "generatedAt": datetime.now().isoformat(timespec="seconds"),
        "sourceZip": str(ZIP_PATH),
        "practiceMode": "browser-rendered-fixed-width",
        "targetUnit": PRACTICE_UNIT_TITLE,
        "viewportWidth": VIEWPORT_WIDTH,
        "contentWidth": CONTENT_WIDTH,
        "items": [],
        "renderedPages": [],
        "supportFiles": [],
        "mediaReferences": [],
        "coverageFailures": [],
    }

    with zipfile.ZipFile(ZIP_PATH) as zip_file:
        manifest_root = ET.fromstring(zip_file.read("imsmanifest.xml"))
        resources = read_resources(manifest_root)
        items = collect_unit_items(manifest_root, resources)
        pages: list[dict[str, Any]] = []
        support_records: list[dict[str, Any]] = []

        for item in items:
            item_record = {
                "index": item.index,
                "title": item.title,
                "identifier": item.identifier,
                "identifierref": item.identifierref,
                "htmlHrefs": item.html_hrefs,
                "supportHrefs": item.support_hrefs,
                "status": "pending",
            }
            missing_hrefs = [
                href for href in item.html_hrefs + item.support_hrefs if not package_path_exists(zip_file, href)
            ]
            if missing_hrefs:
                item_record["status"] = "blocked-missing-source"
                item_record["missingHrefs"] = missing_hrefs
                audit["coverageFailures"].append(item_record)
                audit["items"].append(item_record)
                continue
            for href in item.html_hrefs:
                page_record = prepare_html_page(zip_file, item, href)
                pages.append(page_record)
                audit["mediaReferences"].extend(page_record["media"])
            for href in item.support_hrefs:
                support_records.append(copy_support_file(zip_file, href, item.title))
            item_record["status"] = "rendered" if item.html_hrefs else "support-file-only"
            audit["items"].append(item_record)

        if audit["coverageFailures"]:
            (META_DIR / "browser-practice-docx-export-audit.json").write_text(json.dumps(audit, indent=2), encoding="utf-8")
            raise SystemExit("Coverage failures detected before rendering. See browser-practice-docx-export-audit.json")

        rendered_pages = render_pages_with_playwright(pages)
        output_path = build_docx(rendered_pages, support_records)

    audit["renderedPages"] = rendered_pages
    audit["supportFiles"] = support_records
    audit["outputPath"] = rel_posix(output_path, PROJECT_ROOT)
    audit["outputBytes"] = output_path.stat().st_size
    audit["passed"] = True
    (META_DIR / "browser-practice-docx-export-audit.json").write_text(json.dumps(audit, indent=2), encoding="utf-8")
    with (META_DIR / "browser-practice-supporting-files-index.csv").open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=["sourceHref", "outputPath", "bytes"])
        writer.writeheader()
        writer.writerows(support_records)
    print(f"Wrote {output_path}")
    print(f"Rendered pages: {len(rendered_pages)}")
    print(f"Support files: {len(support_records)}")
    print(f"Media references: {len(audit['mediaReferences'])}")


if __name__ == "__main__":
    main()
