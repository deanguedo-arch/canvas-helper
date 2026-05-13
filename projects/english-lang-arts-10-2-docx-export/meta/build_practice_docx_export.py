from __future__ import annotations

import importlib.util
import json
import os
import sys
import zipfile
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any
from urllib.parse import urlparse
from xml.etree import ElementTree as ET

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[3]
PROJECT_ROOT = REPO_ROOT / "projects" / "english-lang-arts-10-2-docx-export"
ENGLISH_BUILDER_PATH = PROJECT_ROOT / "meta" / "build_docx_export.py"
SOCIAL_BUILDER_PATH = (
    REPO_ROOT / "projects" / "social-studies-10-1-docx-export" / "meta" / "build_unit1_docx_export.py"
)
SOURCE_ZIP_NAME = "D2LCCExport_149674_25-26 _ S2 _ English Lang. Arts 10-2 _ Per 1(A) _ _202651230.zip"
PRACTICE_UNIT_TITLE = "Unit 1: Introduction to Interpreting and Creating Texts"
PRACTICE_OUTPUT_TITLE = "Unit 1 Introduction to Interpreting and Creating Texts - styled practice"


def first_existing_path(env_var: str, candidates: list[Path]) -> Path:
    override = os.environ.get(env_var)
    paths = ([Path(override)] if override else []) + candidates
    for path in paths:
        if path.exists():
            return path
    return paths[0]


ZIP_PATH = first_existing_path(
    "ELA102_SOURCE_ZIP",
    [
        Path.home() / "Downloads" / SOURCE_ZIP_NAME,
        Path("/Users/deanguedo/Downloads") / SOURCE_ZIP_NAME,
    ],
)
META_DIR = PROJECT_ROOT / "meta"
EXPORT_DIR = PROJECT_ROOT / "exports"
DOCX_DIR = EXPORT_DIR / "styled-practice-docx-v2"
SUPPORT_DIR = EXPORT_DIR / "styled-practice-supporting-files-v2"


def load_module(name: str, path: Path) -> Any:
    spec = importlib.util.spec_from_file_location(name, path)
    if spec is None or spec.loader is None:
        raise SystemExit(f"Unable to load builder module: {path}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


english_builder = load_module("english_docx_builder_native_practice", ENGLISH_BUILDER_PATH)
social_builder = load_module("social_docx_media_helpers_for_ela102", SOCIAL_BUILDER_PATH)

english_builder.ZIP_PATH = ZIP_PATH
english_builder.PROJECT_ROOT = PROJECT_ROOT
english_builder.META_DIR = META_DIR
english_builder.EXPORT_DIR = EXPORT_DIR
english_builder.DOCX_DIR = DOCX_DIR
english_builder.SUPPORT_DIR = SUPPORT_DIR
english_builder.REQUESTED_UNITS = [
    english_builder.UnitSpec(PRACTICE_OUTPUT_TITLE, source_title=PRACTICE_UNIT_TITLE)
]


class NativePracticeDocxExporter(english_builder.DocxExporter):
    def __init__(self, zip_file: zipfile.ZipFile, manifest_root: ET.Element):
        super().__init__(zip_file, manifest_root)
        self._media_metadata_cache: dict[str, dict[str, Any] | None] = {}
        self._last_media_preview_used_remote = False
        self.audit.update(
            {
                "practiceMode": "native-renderer-with-brightspace-template-styles",
                "targetUnit": PRACTICE_UNIT_TITLE,
                "mediaReferences": [],
                "mediaRawUrlLines": 0,
                "videoLinkReferences": [],
                "cleanedLinks": [],
                "templateStylesheetsSeen": [],
                "notes": [
                    "Practice export uses the existing English renderer as the base.",
                    "Brightspace template classes and inline styles are preserved where DOCX supports them.",
                    "Media normalization is layered on top: clickable preview image plus raw public URL.",
                    "Image source links and unresolved image text are removed from learner-facing body text.",
                ],
            }
        )

    def _prepare_outputs(self):
        for target in [DOCX_DIR, SUPPORT_DIR]:
            if target.exists():
                try:
                    english_builder.shutil.rmtree(target)
                except PermissionError:
                    pass
            target.mkdir(parents=True, exist_ok=True)
        META_DIR.mkdir(parents=True, exist_ok=True)

    def _write_raw_source_notes(self):
        source_note = {
            "sourceZip": str(ZIP_PATH),
            "sourceZipBytes": ZIP_PATH.stat().st_size,
            "sourceZipLastWriteTime": datetime.fromtimestamp(ZIP_PATH.stat().st_mtime).isoformat(
                timespec="seconds"
            ),
            "rawStorageNote": "Practice export does not copy raw files; the source ZIP remains in Downloads.",
        }
        (META_DIR / "practice-source-package.json").write_text(
            json.dumps(source_note, indent=2), encoding="utf-8"
        )

    def new_document(self):
        document = super().new_document()
        section = document.sections[0]
        section.top_margin = english_builder.Inches(0.55)
        section.bottom_margin = english_builder.Inches(0.55)
        section.left_margin = english_builder.Inches(0.55)
        section.right_margin = english_builder.Inches(0.55)
        styles = document.styles
        styles["Normal"].font.name = "Lato"
        styles["Normal"].font.size = Pt(12)
        for style_name in ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
            styles[style_name].font.name = "Lato"
        return document

    def build_unit(self, index: int, spec: Any):
        if spec.source_title:
            source_item = self.find_top_level(spec.source_title)
            source_kind = "top-level-module"
        else:
            source_item = self.find_item_by_title(spec.source_item_title or "")
            source_kind = "matching-topic"

        unit_record: dict[str, Any] = {
            "outputTitle": spec.output_title,
            "sourceTitle": english_builder.item_title(source_item) if source_item is not None else None,
            "sourceKind": source_kind,
            "note": spec.note,
            "includedItems": [],
            "sourceMissing": source_item is None,
        }

        document = self.new_document()
        self.current_output_title = spec.output_title
        if source_item is None:
            document.add_paragraph("Source content was not found in imsmanifest.xml.")
            output_path = DOCX_DIR / f"{index:02d} - {english_builder.safe_name(spec.output_title)}.docx"
            document.save(output_path)
            unit_record["outputPath"] = english_builder.rel_posix(output_path, PROJECT_ROOT)
            self.audit["requestedUnits"].append(unit_record)
            return

        if source_kind == "matching-topic":
            items = [source_item]
        else:
            items = [
                item
                for item in english_builder.walk_items(source_item)
                if item is not source_item and item.attrib.get("identifierref")
            ]

        first_rendered = False
        for item in items:
            if first_rendered:
                document.add_page_break()
            self.render_item(document, item, 1, spec, unit_record)
            first_rendered = True

        output_path = DOCX_DIR / f"{index:02d} - {english_builder.safe_name(spec.output_title)}.docx"
        document.save(output_path)
        unit_record["outputPath"] = english_builder.rel_posix(output_path, PROJECT_ROOT)
        unit_record["outputBytes"] = output_path.stat().st_size
        self.audit["requestedUnits"].append(unit_record)
        self.audit["docxFiles"].append(str(output_path))

    def render_item(self, document: Any, item: Any, level: int, spec: Any, unit_record: dict[str, Any]):
        title = english_builder.item_title(item) or "Untitled item"
        ref = item.attrib.get("identifierref")
        record: dict[str, Any] = {
            "title": title,
            "identifier": item.attrib.get("identifier"),
            "identifierref": ref,
            "files": [],
        }
        unit_record["includedItems"].append(record)

        if not ref:
            return
        resource = self.resources.get(ref)
        if not resource:
            self.audit["missingResources"].append({"title": title, "identifierref": ref})
            return

        files = resource.get("files") or []
        if not files:
            return

        html_files = [href for href in files if Path(href).suffix.lower() in [".html", ".htm"]]
        if title and not html_files:
            self.add_styled_heading(document, title, min(max(level, 1), 3))

        for href in files:
            ext = Path(href).suffix.lower()
            record["files"].append(href)
            if not english_builder.package_path_exists(self.zip_file, href):
                self.audit["missingResources"].append({"title": title, "href": href})
                continue
            if ext in [".html", ".htm"]:
                self.render_html(document, href)
            elif ext == ".pdf":
                support_path = self.copy_support_file(href, spec.output_title, title)
                self.add_support_reference(document, href, support_path, "Original PDF resource")
                self.render_pdf(document, href)
            elif ext == ".docx":
                support_path = self.copy_support_file(href, spec.output_title, title)
                self.add_support_reference(document, href, support_path, "Original DOCX resource")
                self.render_docx_text(document, href)
            elif ext in [".png", ".jpg", ".jpeg", ".gif"]:
                self.add_image_from_href(document, href)
            else:
                support_path = self.copy_support_file(href, spec.output_title, title)
                label = f"Original {ext.lstrip('.').upper() or 'resource'} file"
                self.add_support_reference(document, href, support_path, label)

    def html_has_visible_heading(self, href: str | None, title: str) -> bool:
        if not href or not english_builder.package_path_exists(self.zip_file, href):
            return False
        try:
            soup = english_builder.BeautifulSoup(self.zip_file.read(href), "lxml")
        except Exception:
            return False
        wanted = english_builder.normalize_key(title)
        for heading in soup.find_all(["h1", "h2", "h3"]):
            if english_builder.normalize_key(heading.get_text(" ")) == wanted:
                return True
        return False

    def render_html(self, document: Any, href: str):
        html_bytes = self.zip_file.read(href)
        soup = english_builder.BeautifulSoup(html_bytes, "lxml")
        for link in soup.find_all("link", href=True):
            self.audit["templateStylesheetsSeen"].append({"base": href, "href": link.get("href")})
        for tag in soup(["script", "style", "meta", "link", "title"]):
            tag.decompose()
        body = soup.body or soup
        for child in body.children:
            self.render_block(document, child, href)

    def render_block(self, document, node: Any, base_href: str):
        if isinstance(node, english_builder.NavigableString):
            text = english_builder.clean_text(str(node))
            if text and not self.is_noise_text(text):
                para = document.add_paragraph(text)
                self.apply_paragraph_styles(para)
            return
        if not isinstance(node, english_builder.Tag):
            return

        name = (node.name or "").lower()
        classes = self.node_classes(node)
        if name in ["script", "style", "meta", "link", "title"]:
            return
        if self.should_skip_noise_node(node):
            return
        if name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            text = english_builder.clean_text(node.get_text(" "))
            if text:
                self.add_styled_heading(document, text, min(int(name[1]), 4), source_node=node)
            return
        if name == "hr":
            self.add_horizontal_rule(document)
            return
        if name in ["iframe", "video", "audio", "embed", "object"]:
            self.add_media_card(document, node, base_href)
            return

        media_children = node.find_all(["iframe", "video", "audio", "embed", "object"], recursive=True)
        if media_children and self.is_media_only_node(node):
            for media_node in media_children:
                self.add_media_card(document, media_node, base_href)
            return
        if "two-col-panels" in classes:
            self.render_two_col_panels(document, node, base_href)
            return
        row_cards = self.direct_card_children(node) if "row" in classes else []
        if row_cards:
            self.render_card_row(document, row_cards, base_href)
            return
        if "card" in classes:
            self.render_card(document, node, base_href)
            return
        if name == "a" and self.is_image_source_link(node):
            self.audit["cleanedLinks"].append({"base": base_href, "text": node.get_text(" "), "href": node.get("href")})
            return
        if name == "img":
            self.render_html_image(document, node, base_href)
            return
        if name in ["ul", "ol"]:
            style = "List Number" if name == "ol" else "List Bullet"
            for li in node.find_all("li", recursive=False):
                para = document.add_paragraph(style=style)
                self.apply_paragraph_styles(para, li)
                self.add_inline(para, li, base_href, font_size=self.inline_font_size(li))
            return
        if name == "table":
            self.render_table(document, node)
            return

        block_children = [
            child
            for child in node.children
            if isinstance(child, english_builder.Tag)
            and (child.name or "").lower()
            in [
                "p",
                "div",
                "section",
                "article",
                "h1",
                "h2",
                "h3",
                "h4",
                "h5",
                "h6",
                "ul",
                "ol",
                "table",
                "img",
                "iframe",
                "video",
                "audio",
                "embed",
                "object",
                "hr",
            ]
        ]
        if name in ["p", "blockquote"] or not block_children:
            para = document.add_paragraph()
            if name == "blockquote":
                para.style = "Intense Quote"
            self.apply_paragraph_styles(para, node)
            self.add_inline(para, node, base_href, font_size=self.inline_font_size(node), color=self.inline_color(node))
            if not english_builder.clean_text(para.text) and not para.runs:
                self.remove_paragraph(para)
            return

        for child in node.children:
            self.render_block(document, child, base_href)

    def add_inline(
        self,
        paragraph,
        node: Any,
        base_href: str,
        bold=False,
        italic=False,
        underline=False,
        font_size: float | None = None,
        color: RGBColor | None = None,
    ):
        if isinstance(node, english_builder.NavigableString):
            raw_text = str(node).replace("\xa0", " ")
            text = english_builder.clean_text(raw_text)
            if text and not self.is_noise_text(text):
                if paragraph.text and not paragraph.text.endswith((" ", "\n")) and text[0] not in ".,;:!?)]}":
                    if raw_text[:1].isspace() or text[:1].isalnum():
                        text = " " + text
                if raw_text[-1:].isspace() and not text.endswith(" "):
                    text += " "
                run = paragraph.add_run(text)
                run.bold = bold
                run.italic = italic
                run.underline = underline
                if font_size:
                    run.font.size = Pt(font_size)
                if color:
                    run.font.color.rgb = color
            return
        if not isinstance(node, english_builder.Tag):
            return
        name = (node.name or "").lower()
        if name == "br":
            paragraph.add_run().add_break()
            return
        if name == "img":
            self.render_inline_html_image(paragraph, node, base_href)
            return
        if name == "a":
            href = (node.get("href") or "").strip()
            if self.is_image_source_link(node):
                self.audit["cleanedLinks"].append({"base": base_href, "text": node.get_text(" "), "href": href})
                for img in node.find_all("img"):
                    self.render_inline_html_image(paragraph, img, base_href)
                return
            if self.is_video_url(href):
                handoff_url = self.media_handoff_url(href)
                english_builder.add_hyperlink(paragraph, handoff_url, handoff_url)
                self.audit["videoLinkReferences"].append(
                    {
                        "base": base_href,
                        "text": english_builder.clean_text(node.get_text(" ")),
                        "href": href,
                        "handoffUrl": handoff_url,
                    }
                )
                return
            link_text = english_builder.clean_text(node.get_text(" ")) or href
            cleaned_href = self.clean_non_video_href(href)
            if cleaned_href and self.is_public_href(cleaned_href):
                english_builder.add_hyperlink(paragraph, link_text, cleaned_href)
            else:
                paragraph.add_run(link_text)
            return
        if name in ["iframe", "video", "audio", "embed", "object"]:
            src = node.get("src") or node.get("data") or ""
            if src:
                handoff_url = self.media_handoff_url(src)
                english_builder.add_hyperlink(paragraph, handoff_url, handoff_url)
                self.audit["videoLinkReferences"].append({"base": base_href, "href": src, "handoffUrl": handoff_url})
            return
        if name in ["strong", "b"]:
            bold = True
        if name in ["em", "i"]:
            italic = True
        if name == "u":
            underline = True
        node_font_size = self.inline_font_size(node) or font_size
        node_color = self.inline_color(node) or color
        for child in node.children:
            self.add_inline(
                paragraph,
                child,
                base_href,
                bold=bold,
                italic=italic,
                underline=underline,
                font_size=node_font_size,
                color=node_color,
            )

    def render_html_image(self, document, img: Any, base_href: str):
        src = img.get("src")
        if not src:
            return
        package_href = english_builder.resolve_package_href(base_href, src, self.zip_file)
        if not package_href:
            self.audit["unresolvedHtmlAssets"].append({"base": base_href, "src": src})
            return
        self.add_image_from_href(document, package_href, source_node=img)

    def render_inline_html_image(self, paragraph, img: Any, base_href: str):
        src = img.get("src")
        if not src:
            return
        package_href = english_builder.resolve_package_href(base_href, src, self.zip_file)
        if not package_href:
            self.audit["unresolvedHtmlAssets"].append({"base": base_href, "src": src})
            return
        try:
            image_bytes = self.zip_file.read(package_href)
            normalized_bytes, width_px, height_px = self.normalized_image_bytes(image_bytes)
            if width_px <= 0 or height_px <= 0:
                raise ValueError("invalid image size")
            declared_width = self.declared_image_width_inches(img)
            width_inches = min(6.85, declared_width or width_px / 110)
            paragraph.add_run().add_picture(BytesIO(normalized_bytes), width=english_builder.Inches(width_inches))
            self.audit["embeddedImages"] += 1
        except Exception as exc:
            if english_builder.package_path_exists(self.zip_file, package_href):
                self.copy_support_file(package_href, self.current_output_title, "Unembedded image")
            self.audit["unresolvedHtmlAssets"].append({"href": package_href, "error": str(exc)})

    def add_image_from_href(self, document: Any, href: str, source_node: Any | None = None):
        try:
            image_bytes = self.zip_file.read(href)
            normalized_bytes, width_px, height_px = self.normalized_image_bytes(image_bytes)
            if width_px <= 0 or height_px <= 0:
                raise ValueError("invalid image size")
            declared_width = self.declared_image_width_inches(source_node)
            width_inches = min(6.85, declared_width or width_px / 110)
            paragraph = document.add_paragraph()
            paragraph.alignment = self.image_alignment(source_node)
            paragraph.add_run().add_picture(BytesIO(normalized_bytes), width=english_builder.Inches(width_inches))
            self.audit["embeddedImages"] += 1
        except Exception as exc:
            if english_builder.package_path_exists(self.zip_file, href):
                self.copy_support_file(href, self.current_output_title, "Unembedded image")
            self.audit["unresolvedHtmlAssets"].append({"href": href, "error": str(exc)})

    def is_image_source_link(self, node: Any) -> bool:
        text = english_builder.normalize_key(node.get_text(" ") if hasattr(node, "get_text") else "")
        return text in {"image source", "image sources", "source"}

    def is_video_url(self, url: str) -> bool:
        host = urlparse(url).netloc.casefold()
        path = urlparse(url).path.casefold()
        return any(token in host for token in ["youtube", "youtu.be", "ted.com"]) or path.endswith(
            (".mp4", ".mov", ".m4v", ".webm")
        )

    def add_media_card(self, document, node: Any, base_href: str) -> None:
        src = (node.get("src") or node.get("data") or "").strip()
        if not src:
            return
        raw_title = node.get("title") or english_builder.clean_text(node.get_text(" ")) or "Embedded media"
        title = self.media_display_title(raw_title, src)
        handoff_url = self.media_handoff_url(src)
        link_url = handoff_url or src
        preview = document.add_paragraph()
        preview.alignment = WD_ALIGN_PARAGRAPH.CENTER
        preview_bytes = self.media_preview_image_bytes(title, src)
        used_remote_thumbnail = bool(getattr(self, "_last_media_preview_used_remote", False))
        self.add_hyperlinked_picture(preview, preview_bytes, width=6.35, url=link_url)
        if handoff_url:
            handoff = document.add_paragraph()
            handoff.alignment = WD_ALIGN_PARAGRAPH.CENTER
            english_builder.add_hyperlink(handoff, handoff_url, handoff_url)
            self.audit["mediaRawUrlLines"] += 1
        self.audit["mediaReferences"].append(
            {
                "sourceHtml": base_href,
                "sourceTitle": self.current_output_title,
                "type": (node.name or "media").lower(),
                "title": title,
                "src": src,
                "linkUrl": link_url,
                "handoffUrl": handoff_url,
                "thumbnailUrl": self.media_thumbnail_url(src),
                "usedRemoteThumbnail": used_remote_thumbnail,
            }
        )

    def node_classes(self, node: Any) -> set[str]:
        if not isinstance(node, english_builder.Tag):
            return set()
        return {str(value).casefold() for value in (node.get("class") or [])}

    def style_map(self, node: Any) -> dict[str, str]:
        if not isinstance(node, english_builder.Tag):
            return {}
        styles: dict[str, str] = {}
        for part in (node.get("style") or "").split(";"):
            if ":" not in part:
                continue
            key, value = part.split(":", 1)
            styles[key.strip().casefold()] = value.strip()
        return styles

    def inline_font_size(self, node: Any) -> float | None:
        value = self.style_map(node).get("font-size")
        if not value:
            return None
        match = english_builder.re.search(r"([0-9.]+)\s*(px|pt)?", value)
        if not match:
            return None
        amount = float(match.group(1))
        unit = (match.group(2) or "px").casefold()
        return amount if unit == "pt" else amount * 0.75

    def inline_color(self, node: Any) -> RGBColor | None:
        value = self.style_map(node).get("color")
        if not value:
            return None
        match = english_builder.re.match(r"rgb\((\d+),\s*(\d+),\s*(\d+)\)", value)
        if match:
            return RGBColor(*(int(part) for part in match.groups()))
        match = english_builder.re.match(r"#?([0-9a-fA-F]{6})", value)
        if match:
            raw = match.group(1)
            return RGBColor(int(raw[0:2], 16), int(raw[2:4], 16), int(raw[4:6], 16))
        return None

    def apply_paragraph_styles(self, paragraph: Any, node: Any | None = None, default_align: Any | None = None):
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.line_spacing = 1.12
        alignment = default_align
        if isinstance(node, english_builder.Tag):
            text_align = self.style_map(node).get("text-align")
            if text_align == "center":
                alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif text_align == "right":
                alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if alignment is not None:
            paragraph.alignment = alignment

    def add_styled_heading(self, document: Any, text: str, level: int, source_node: Any | None = None):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(6 if level <= 2 else 4)
        paragraph.paragraph_format.space_after = Pt(10 if level <= 2 else 6)
        if level in [1, 2]:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(english_builder.clean_text(text))
        run.bold = True
        if level == 1:
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(53, 62, 72)
        elif level == 2:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0, 128, 128)
        elif level == 3:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(47, 87, 139)
        else:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(53, 62, 72)
        node_size = self.inline_font_size(source_node)
        node_color = self.inline_color(source_node)
        if node_size:
            run.font.size = Pt(node_size)
        if node_color:
            run.font.color.rgb = node_color
        return paragraph

    def add_horizontal_rule(self, document: Any):
        paragraph = document.add_paragraph()
        p_pr = paragraph._p.get_or_add_pPr()
        border = english_builder.OxmlElement("w:pBdr")
        bottom = english_builder.OxmlElement("w:bottom")
        bottom.set(english_builder.qn("w:val"), "single")
        bottom.set(english_builder.qn("w:sz"), "6")
        bottom.set(english_builder.qn("w:space"), "1")
        bottom.set(english_builder.qn("w:color"), "BFC7D1")
        border.append(bottom)
        p_pr.append(border)

    def direct_card_children(self, node: Any) -> list[Any]:
        cards: list[Any] = []
        if not isinstance(node, english_builder.Tag):
            return cards
        for child in node.find_all(recursive=False):
            child_classes = self.node_classes(child)
            if "card" in child_classes:
                cards.append(child)
            else:
                for grand in child.find_all(recursive=False):
                    if "card" in self.node_classes(grand):
                        cards.append(grand)
        return cards

    def render_two_col_panels(self, document: Any, node: Any, base_href: str):
        rendered = False
        for child in node.find_all(recursive=False):
            if "row" in self.node_classes(child):
                cards = self.direct_card_children(child)
                if cards:
                    self.render_card_row(document, cards, base_href)
                    rendered = True
        if not rendered:
            cards = self.direct_card_children(node)
            if cards:
                self.render_card_row(document, cards, base_href)

    def render_card_row(self, document: Any, cards: list[Any], base_href: str):
        if not cards:
            return
        table = document.add_table(rows=1, cols=len(cards))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        for index, card in enumerate(cards):
            cell = table.cell(0, index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            self.set_cell_fill(cell, "F8F9FA")
            self.set_cell_border(cell, color="D9DEE7", size="8")
            self.set_cell_width(cell, 3.05)
            self.clear_cell(cell)
            self.render_card_contents(cell, card, base_href, centered=True)
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)

    def render_card(self, document: Any, node: Any, base_href: str):
        table = document.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        is_graphic = "card-graphic" in self.node_classes(node)
        self.set_cell_border(cell, color="008C8C" if is_graphic else "D9DEE7", size="16" if is_graphic else "8")
        self.set_cell_fill(cell, "FFFFFF" if is_graphic else "F8F9FA")
        self.clear_cell(cell)
        centered = bool(node.find(style=lambda value: value and "text-align: center" in value))
        self.render_card_contents(cell, node, base_href, centered=centered)
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(8)

    def render_card_contents(self, container: Any, node: Any, base_href: str, centered: bool):
        content = node.find(class_=lambda value: value and "card-body" in str(value).split()) or node
        for child in content.children:
            if isinstance(child, english_builder.NavigableString):
                text = english_builder.clean_text(str(child))
                if text:
                    para = container.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else None
                    run = para.add_run(text)
                    run.font.size = Pt(12.5)
                continue
            if not isinstance(child, english_builder.Tag):
                continue
            name = (child.name or "").lower()
            if name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                para = container.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_after = Pt(4)
                run = para.add_run(english_builder.clean_text(child.get_text(" ")))
                run.bold = True
                run.font.size = Pt(13 if name == "h4" else 14)
                run.font.color.rgb = RGBColor(53, 62, 72)
            elif name in ["p", "div"]:
                if self.node_classes(child) & {"card-text", "card-body"}:
                    self.render_card_contents(container, child, base_href, centered=centered)
                else:
                    para = container.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else None
                    para.paragraph_format.space_after = Pt(4)
                    self.add_inline(para, child, base_href, font_size=self.inline_font_size(child) or 12.5)
            else:
                self.render_block(container, child, base_href)

    def clear_cell(self, cell: Any):
        for paragraph in list(cell.paragraphs):
            self.remove_paragraph(paragraph)

    def set_cell_fill(self, cell: Any, fill: str):
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = english_builder.OxmlElement("w:shd")
        shading.set(english_builder.qn("w:fill"), fill)
        tc_pr.append(shading)

    def set_cell_width(self, cell: Any, width_inches: float):
        tc_pr = cell._tc.get_or_add_tcPr()
        width = english_builder.OxmlElement("w:tcW")
        width.set(english_builder.qn("w:w"), str(int(width_inches * 1440)))
        width.set(english_builder.qn("w:type"), "dxa")
        tc_pr.append(width)

    def set_cell_border(self, cell: Any, color: str, size: str):
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.first_child_found_in("w:tcBorders")
        if borders is None:
            borders = english_builder.OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        for edge in ["top", "left", "bottom", "right"]:
            element = english_builder.OxmlElement(f"w:{edge}")
            element.set(english_builder.qn("w:val"), "single")
            element.set(english_builder.qn("w:sz"), size)
            element.set(english_builder.qn("w:space"), "0")
            element.set(english_builder.qn("w:color"), color)
            borders.append(element)

    def declared_image_width_inches(self, node: Any | None) -> float | None:
        if not isinstance(node, english_builder.Tag):
            return None
        raw_width = node.get("width") or self.style_map(node).get("width")
        if not raw_width:
            return None
        match = english_builder.re.search(r"([0-9.]+)", str(raw_width))
        if not match:
            return None
        return float(match.group(1)) / 96

    def image_alignment(self, node: Any | None):
        if not isinstance(node, english_builder.Tag):
            return WD_ALIGN_PARAGRAPH.CENTER
        style = self.style_map(node)
        if style.get("float") == "right":
            return WD_ALIGN_PARAGRAPH.RIGHT
        if style.get("float") == "left":
            return WD_ALIGN_PARAGRAPH.LEFT
        return WD_ALIGN_PARAGRAPH.CENTER

    def should_skip_noise_node(self, node: Any) -> bool:
        if isinstance(node, english_builder.Tag) and node.find(["img", "iframe", "video", "audio", "embed", "object"]):
            return False
        text = english_builder.normalize_key(node.get_text(" ") if hasattr(node, "get_text") else "")
        return self.is_noise_text(text)

    def is_noise_text(self, text: str) -> bool:
        normalized = english_builder.normalize_key(text)
        return normalized in {
            "",
            "image source",
            "image sources",
            "source",
            "template javascript",
        }

    def is_media_only_node(self, node: Any) -> bool:
        text = english_builder.normalize_key(node.get_text(" ") if hasattr(node, "get_text") else "")
        return text in {"", "embedded media", "ted", "youtube", "youtube video player"}

    def clean_non_video_href(self, href: str) -> str:
        if not href:
            return ""
        href = href.strip()
        if "d2lsessionval=" in href.casefold():
            href = href.split("?", 1)[0]
        return href

    def is_public_href(self, href: str) -> bool:
        return href.startswith("http://") or href.startswith("https://")

    def write_resource_indexes(self):
        rows = list(self.audit["copiedSupportFiles"])
        csv_path = EXPORT_DIR / "practice-supporting-files-index.csv"
        with csv_path.open("w", encoding="utf-8", newline="") as handle:
            writer = english_builder.csv.DictWriter(handle, fieldnames=["sourceHref", "outputPath", "bytes"])
            writer.writeheader()
            writer.writerows(rows)
        map_payload = {
            "requestedUnits": self.audit["requestedUnits"],
            "sourceTopLevelModules": [
                {
                    "title": english_builder.item_title(module),
                    "resourceItemCount": len(
                        [item for item in english_builder.walk_items(module) if item.attrib.get("identifierref")]
                    ),
                }
                for module in self.top_level_modules()
            ],
        }
        (META_DIR / "practice-conversion-map.json").write_text(
            json.dumps(map_payload, indent=2), encoding="utf-8"
        )

    def verify_outputs(self):
        verification = {
            "expectedDocxCount": 1,
            "actualDocxCount": 0,
            "invalidDocx": [],
            "docx": [],
            "mediaReferences": len(self.audit["mediaReferences"]),
            "embeddedImages": self.audit["embeddedImages"],
        }
        for path in sorted(DOCX_DIR.glob("*.docx")):
            verification["actualDocxCount"] += 1
            record = {"path": english_builder.rel_posix(path, PROJECT_ROOT), "bytes": path.stat().st_size}
            try:
                with zipfile.ZipFile(path) as docx_zip:
                    docx_zip.getinfo("word/document.xml")
                    record["wordMediaFiles"] = len(
                        [name for name in docx_zip.namelist() if name.startswith("word/media/")]
                    )
            except Exception as exc:
                record["error"] = str(exc)
                verification["invalidDocx"].append(record)
            verification["docx"].append(record)
        verification["passed"] = verification["actualDocxCount"] == 1 and not verification["invalidDocx"]
        self.audit["verification"] = verification
        (META_DIR / "practice-docx-export-verification.json").write_text(
            json.dumps(verification, indent=2), encoding="utf-8"
        )
        if not verification["passed"]:
            raise SystemExit("DOCX verification failed. See meta/practice-docx-export-verification.json")

    def write_audit_markdown(self):
        lines = [
            "# English Language Arts 10-2 Native Practice DOCX Export Audit",
            "",
            f"- Generated: {self.audit['generatedAt']}",
            f"- Source ZIP: `{self.audit['sourceZip']}`",
            f"- Target unit: `{PRACTICE_UNIT_TITLE}`",
            f"- DOCX files generated: {self.audit['verification']['actualDocxCount']}",
            f"- Embedded HTML images: {self.audit['embeddedImages']}",
            f"- Media references: {len(self.audit['mediaReferences'])}",
            f"- Raw media URL lines: {self.audit['mediaRawUrlLines']}",
            f"- Video links normalized: {len(self.audit['videoLinkReferences'])}",
            f"- Unresolved HTML assets: {len(self.audit['unresolvedHtmlAssets'])}",
            "",
            "## Generated DOCX Files",
            "",
        ]
        for docx_record in self.audit["verification"]["docx"]:
            lines.append(f"- `{docx_record['path']}` ({docx_record['bytes'] / 1024 / 1024:.2f} MB)")
        if self.audit["mediaReferences"]:
            lines.extend(["", "## Embedded Media Handoff Links", ""])
            for media in self.audit["mediaReferences"]:
                lines.append(f"- `{media['title']}` -> {media.get('handoffUrl') or media['src']}")
        if self.audit["videoLinkReferences"]:
            lines.extend(["", "## Video Links Normalized", ""])
            for link in self.audit["videoLinkReferences"]:
                lines.append(f"- `{link.get('text') or link.get('href')}` -> {link.get('handoffUrl')}")
        if self.audit["unresolvedHtmlAssets"]:
            lines.extend(["", "## Audit-Only Missing/Fallback Items", ""])
            for asset in self.audit["unresolvedHtmlAssets"][:40]:
                lines.append(f"- `{asset}`")
        lines.append("")
        (META_DIR / "practice-docx-export-audit.md").write_text("\n".join(lines), encoding="utf-8")
        (META_DIR / "practice-docx-export-audit.json").write_text(
            json.dumps(self.audit, indent=2), encoding="utf-8"
        )

    def youtube_video_id(self, src: str) -> str | None:
        return social_builder.UnitOneDocxExporter.youtube_video_id(self, src)

    def media_provider_label(self, src: str) -> str:
        return social_builder.UnitOneDocxExporter.media_provider_label(self, src)

    def media_thumbnail_url(self, src: str) -> str | None:
        return social_builder.UnitOneDocxExporter.media_thumbnail_url(self, src)

    def media_display_title(self, title: str, src: str) -> str:
        return social_builder.UnitOneDocxExporter.media_display_title(self, title, src)

    def media_handoff_url(self, src: str) -> str:
        return social_builder.UnitOneDocxExporter.media_handoff_url(self, src)

    def media_oembed_metadata(self, src: str) -> dict[str, Any] | None:
        return social_builder.UnitOneDocxExporter.media_oembed_metadata(self, src)

    def fetch_remote_json(self, url: str) -> dict[str, Any] | None:
        return social_builder.UnitOneDocxExporter.fetch_remote_json(self, url)

    def fetch_remote_image_bytes(self, url: str) -> bytes | None:
        return social_builder.UnitOneDocxExporter.fetch_remote_image_bytes(self, url)

    def fitted_video_thumbnail_bytes(self, image_bytes: bytes) -> bytes:
        return social_builder.UnitOneDocxExporter.fitted_video_thumbnail_bytes(self, image_bytes)

    def media_preview_image_bytes(self, title: str, src: str) -> bytes:
        return social_builder.UnitOneDocxExporter.media_preview_image_bytes(self, title, src)

    def add_hyperlinked_picture(self, paragraph: Any, image_bytes: bytes, width: float, url: str) -> None:
        return social_builder.UnitOneDocxExporter.add_hyperlinked_picture(self, paragraph, image_bytes, width, url)


def main() -> None:
    if not ZIP_PATH.exists():
        raise SystemExit(f"Source ZIP not found: {ZIP_PATH}")
    with zipfile.ZipFile(ZIP_PATH) as zip_file:
        manifest_root = ET.fromstring(zip_file.read("imsmanifest.xml"))
        exporter = NativePracticeDocxExporter(zip_file, manifest_root)
        exporter.build_all()
        for record in exporter.audit["verification"]["docx"]:
            print(f"Wrote {PROJECT_ROOT / record['path']}")


if __name__ == "__main__":
    main()
