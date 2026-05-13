from __future__ import annotations

import csv
import hashlib
import json
import os
import posixpath
import re
import shutil
import sys
import unicodedata
import zipfile
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any
from urllib.parse import parse_qs, quote, unquote, urlparse
from urllib.request import Request, urlopen
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import html as lxml_html
from lxml.html import HtmlElement
from PIL import Image, ImageDraw, ImageFont

PYMUPDF_TARGET = Path("/private/tmp/codex-pymupdf")
if PYMUPDF_TARGET.exists():
    sys.path.insert(0, str(PYMUPDF_TARGET))

try:
    import fitz  # type: ignore
except Exception:  # pragma: no cover - recorded in audit when unavailable
    fitz = None


REPO_ROOT = Path(__file__).resolve().parents[3]
PROJECT_ROOT = REPO_ROOT / "projects" / "social-studies-10-1-docx-export"
SOURCE_ZIP_NAME = "D2LCCExport_149634_25-26 _ S2 _ Social Studies 10-1 _ Per 1(A) _ Sec _202651213.ZIP"
STYLE_REFERENCE_ZIP_NAME = "U1P02overviewsurvey.html.zip"


def first_existing_path(env_var: str, candidates: list[Path]) -> Path:
    override = os.environ.get(env_var)
    paths = ([Path(override)] if override else []) + candidates
    for path in paths:
        if path.exists():
            return path
    return paths[0]


ZIP_PATH = first_existing_path(
    "SOCIAL10_SOURCE_ZIP",
    [
        Path.home() / "Downloads" / SOURCE_ZIP_NAME,
        Path("/Users/deanguedo/Downloads") / SOURCE_ZIP_NAME,
    ],
)
META_DIR = PROJECT_ROOT / "meta"
EXPORT_DIR = PROJECT_ROOT / "exports"
DOCX_DIR = EXPORT_DIR / "docx"
SUPPORT_DIR = EXPORT_DIR / "supporting-files"
QA_DIR = EXPORT_DIR / "qa"

UNIT_TITLE = "1. Globalization and Identity"
STYLE_REFERENCE_ZIP = first_existing_path(
    "SOCIAL10_STYLE_REFERENCE_ZIP",
    [
        Path.home() / "Downloads" / STYLE_REFERENCE_ZIP_NAME,
        Path("/Users/deanguedo/Downloads") / STYLE_REFERENCE_ZIP_NAME,
    ],
)
SITE_BLUE = "6096BF"
SITE_PURPLE = "4B4665"
SITE_TEXT = "333333"
SITE_MUTED = "666666"
SITE_OUTER_BORDER = "A7A7A7"
SITE_INNER_BORDER = "727EA3"
SITE_FEATURE_BORDER = "CCCCCC"
SITE_GREEN = "EBF5EB"
SITE_LAVENDER = "E2E1EE"
SITE_GOLD = "FFF1C9"
SITE_LIGHT_BLUE = "DFF1FF"
SITE_FRAME_WIDTH_DXA = 10800
SITE_TABLE_WIDTH_DXA = 10080
SITE_LEFT_IMAGE_DXA = 3900
SITE_TEXT_IMAGE_DXA = SITE_TABLE_WIDTH_DXA - SITE_LEFT_IMAGE_DXA
SITE_HEADER_WIDTH_IN = 7.15
SITE_FULL_IMAGE_WIDTH_IN = 7.0
TREBUCHET_FONT_PATH = Path("/System/Library/Fonts/Supplemental/Trebuchet MS.ttf")

BLOCK_TAGS = {
    "address",
    "article",
    "aside",
    "blockquote",
    "center",
    "dd",
    "details",
    "div",
    "dl",
    "dt",
    "fieldset",
    "figcaption",
    "figure",
    "footer",
    "form",
    "h1",
    "h2",
    "h3",
    "h4",
    "h5",
    "h6",
    "header",
    "hr",
    "iframe",
    "li",
    "main",
    "nav",
    "object",
    "ol",
    "p",
    "pre",
    "section",
    "table",
    "ul",
    "video",
    "audio",
    "embed",
}

CALLOUT_IDS = {
    "readingassignment": ("", SITE_GREEN),
    "internet": ("", SITE_GREEN),
    "email": ("", SITE_GREEN),
    "vocabulary": ("", SITE_GREEN),
    "media": ("", SITE_LAVENDER),
    "portfolio": ("", SITE_LAVENDER),
    "multipleperspectives": ("", SITE_LAVENDER),
    "trackyourprogress": ("", SITE_LAVENDER),
    "help": ("", SITE_LAVENDER),
    "tools": ("", "D6E1ED"),
    "skills": ("", "D6E1ED"),
    "tip": ("", SITE_LIGHT_BLUE),
    "discussions": ("", SITE_LIGHT_BLUE),
    "journal": ("", "EDE2EF"),
    "assignmentdrop": ("", SITE_GOLD),
    "quiz": ("", SITE_GOLD),
    "bonus": ("", "EBF5EB"),
}


@dataclass(frozen=True)
class RenderContext:
    base_href: str
    heading_base: int
    source_title: str


def local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1]


def item_title(item: ET.Element) -> str:
    for child in item:
        if local_name(child.tag) == "title":
            return "".join(child.itertext()).strip()
    return ""


def item_children(item: ET.Element) -> list[ET.Element]:
    return [child for child in item if local_name(child.tag) == "item"]


def walk_items(item: ET.Element) -> list[ET.Element]:
    result = [item]
    for child in item_children(item):
        result.extend(walk_items(child))
    return result


def normalize_key(value: str) -> str:
    return re.sub(r"\s+", " ", value.strip()).casefold()


def safe_name(value: str, max_len: int = 96) -> str:
    normalized = unicodedata.normalize("NFKD", value)
    ascii_value = normalized.encode("ascii", "ignore").decode("ascii")
    ascii_value = re.sub(r"[^\w\s().,&+-]+", "", ascii_value)
    ascii_value = re.sub(r"\s+", " ", ascii_value).strip()
    ascii_value = ascii_value.replace("&", "and")
    ascii_value = ascii_value.strip(" .")
    if not ascii_value:
        ascii_value = "untitled"
    return ascii_value[:max_len].strip(" .")


def rel_posix(path: Path, start: Path) -> str:
    return Path(os.path.relpath(path, start)).as_posix()


def clean_text(value: str) -> str:
    value = value.replace("\xa0", " ")
    value = value.replace("\u200b", "")
    return re.sub(r"\s+", " ", value).strip()


def decode_html(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "windows-1252", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def package_path_exists(zip_file: zipfile.ZipFile, href: str) -> bool:
    try:
        zip_file.getinfo(href)
        return True
    except KeyError:
        return False


def package_match_score(base_href: str, name: str, suffix: str) -> tuple[int, int, int, int]:
    base_segments = [segment.casefold() for segment in posixpath.dirname(base_href).split("/") if segment]
    name_segments = [segment.casefold() for segment in posixpath.dirname(name).split("/") if segment]
    common = 0
    for left, right in zip(base_segments, name_segments):
        if left != right:
            break
        common += 1
    suffix_depth = len([segment for segment in suffix.split("/") if segment])
    distance = abs(len(name_segments) - len(base_segments))
    return (common, suffix_depth, -distance, -len(name))


def resolve_package_href(base_href: str, relative_href: str, zip_file: zipfile.ZipFile) -> str | None:
    if not relative_href:
        return None
    if re.match(r"^[a-zA-Z][a-zA-Z0-9+.-]*:", relative_href):
        return None
    href = relative_href.split("#", 1)[0].split("?", 1)[0].strip()
    if not href:
        return None
    href = unquote(href.replace("\\", "/"))
    candidate = posixpath.normpath(posixpath.join(posixpath.dirname(base_href), href))
    for possible in (candidate, href):
        if package_path_exists(zip_file, possible):
            return possible
    suffixes = []
    stripped = re.sub(r"^(\.\./)+", "", href)
    if stripped and stripped != href:
        suffixes.append(stripped)
    basename = posixpath.basename(href)
    if basename and "." in basename:
        suffixes.append(basename)
    names = zip_file.namelist()
    for suffix in suffixes:
        suffix_key = suffix.casefold()
        matches = [
            name
            for name in names
            if not name.endswith("/") and unquote(name).casefold().endswith(suffix_key)
        ]
        if len(matches) == 1:
            return matches[0]
        if len(matches) > 1:
            ranked = sorted(
                ((package_match_score(base_href, name, suffix), name) for name in matches),
                reverse=True,
            )
            if len(ranked) == 1 or ranked[0][0] != ranked[1][0]:
                return ranked[0][1]
    return None


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell: Any, top: int = 100, bottom: int = 100, start: int = 140, end: int = 140) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell: Any, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def set_table_width(table: Any, width_dxa: int = SITE_TABLE_WIDTH_DXA, indent_dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))


def set_table_width_pct(table: Any, width_pct: int = 5000, indent_dxa: int = 0) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), str(width_pct))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))


def set_table_grid(table: Any, widths_dxa: list[int]) -> None:
    tbl = table._tbl
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)


def set_table_layout_fixed(table: Any) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_table_borders(table: Any, color: str = "D9D9D9", size: str = "8", inside: bool = True) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    names = ["top", "left", "bottom", "right"]
    if inside:
        names.extend(["insideH", "insideV"])
    for name in names:
        border = borders.find(qn(f"w:{name}"))
        if border is None:
            border = OxmlElement(f"w:{name}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def clear_table_borders(table: Any) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = borders.find(qn(f"w:{name}"))
        if border is None:
            border = OxmlElement(f"w:{name}")
            borders.append(border)
        border.set(qn("w:val"), "nil")


def set_paragraph_spacing(paragraph: Any, before: int = 0, after: int = 120, line: int | None = 300) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line is not None:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "auto")


def css_color(value: str | None, fallback: str = "D9D9D9") -> str:
    if not value:
        return fallback
    value = value.strip().lstrip("#")
    if re.fullmatch(r"[0-9a-fA-F]{6}", value):
        return value.upper()
    return fallback


def int_attr(value: str | None) -> int | None:
    if not value:
        return None
    match = re.search(r"\d+", value)
    return int(match.group(0)) if match else None


def add_hyperlink(paragraph: Any, text: str, url: str) -> None:
    text = text or url
    if not text:
        return
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(color)
    properties.append(underline)
    run.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_container_heading(container: Any, text: str, level: int) -> Any:
    if hasattr(container, "add_heading"):
        return container.add_heading(text, level=level)
    style_name = "Title" if level == 0 else f"Heading {min(max(level, 1), 4)}"
    return container.add_paragraph(text, style=style_name)


class UnitOneDocxExporter:
    def __init__(self, zip_file: zipfile.ZipFile, manifest_root: ET.Element):
        self.zip_file = zip_file
        self.manifest_root = manifest_root
        self.resources = self._read_resources()
        self._copied_support: dict[str, Path] = {}
        self._media_metadata_cache: dict[str, dict[str, Any] | None] = {}
        self.audit: dict[str, Any] = {
            "schemaVersion": 1,
            "generatedAt": datetime.now().isoformat(timespec="seconds"),
            "sourceZip": str(ZIP_PATH),
            "sourceZipBytes": ZIP_PATH.stat().st_size,
            "styleReferenceZip": str(STYLE_REFERENCE_ZIP),
            "unitTitle": UNIT_TITLE,
            "docxStylePreset": "brightspace_site_frame",
            "includedItems": [],
            "supportFiles": [],
            "embeddedImages": 0,
            "mediaReferences": [],
            "renderedPdfPages": 0,
            "linkedLocalResources": [],
            "missingResources": [],
            "unresolvedAssets": [],
            "verification": {},
            "notes": [
                "Brightspace hierarchy is read from imsmanifest.xml.",
                "Word pages are styled from the original Brightspace CSS: bordered shell, blue header, Trebuchet/Lucida type, and feature/callout boxes.",
                "Blue Brightspace headers are generated as Word-stable banner images from the CSS colors and Trebuchet font to avoid DOCX nested-table rendering drift.",
                "Embedded iframes/videos are represented as playable links because DOCX cannot preserve live iframe execution.",
            ],
        }

    def _read_resources(self) -> dict[str, dict[str, Any]]:
        resources: dict[str, dict[str, Any]] = {}
        for resource in self.manifest_root.iter():
            if local_name(resource.tag) != "resource":
                continue
            identifier = resource.attrib.get("identifier")
            if not identifier:
                continue
            files = [
                file_node.attrib.get("href")
                for file_node in resource
                if local_name(file_node.tag) == "file" and file_node.attrib.get("href")
            ]
            resources[identifier] = {
                "identifier": identifier,
                "type": resource.attrib.get("type"),
                "files": files,
            }
        return resources

    def top_level_modules(self) -> list[ET.Element]:
        organization = next(
            node for node in self.manifest_root.iter() if local_name(node.tag) == "organization"
        )
        root_items = item_children(organization)
        if len(root_items) == 1 and not item_title(root_items[0]):
            return item_children(root_items[0])
        return root_items

    def find_unit(self) -> ET.Element:
        wanted = normalize_key(UNIT_TITLE)
        for item in self.top_level_modules():
            if normalize_key(item_title(item)) == wanted:
                return item
        raise SystemExit(f"Could not find Unit 1 module in imsmanifest.xml: {UNIT_TITLE}")

    def build(self) -> Path:
        self.prepare_outputs()
        unit = self.find_unit()
        document = self.new_document()

        children = item_children(unit)
        for index, child in enumerate(children):
            self.render_manifest_item(document, child, root=unit, depth=1)
            if index < len(children) - 1:
                document.add_page_break()

        output_path = self.available_output_path(DOCX_DIR / "01 - 1. Globalization and Identity.docx")
        document.save(output_path)
        self.audit["outputPath"] = rel_posix(output_path, PROJECT_ROOT)
        self.audit["outputBytes"] = output_path.stat().st_size
        self.write_indexes()
        self.verify_output(output_path)
        self.write_audit_markdown()
        return output_path

    def prepare_outputs(self) -> None:
        for target in (DOCX_DIR, SUPPORT_DIR, QA_DIR):
            if target.exists():
                try:
                    shutil.rmtree(target)
                except PermissionError:
                    if target != DOCX_DIR:
                        raise
            target.mkdir(parents=True, exist_ok=True)
        META_DIR.mkdir(parents=True, exist_ok=True)

    def output_path_is_locked(self, path: Path) -> bool:
        if not path.exists():
            return False
        try:
            with path.open("ab"):
                return False
        except PermissionError:
            return True

    def available_output_path(self, canonical_path: Path) -> Path:
        if not self.output_path_is_locked(canonical_path):
            return canonical_path
        fallback = canonical_path.with_name(f"{canonical_path.stem} - refreshed{canonical_path.suffix}")
        counter = 2
        while fallback.exists() and self.output_path_is_locked(fallback):
            fallback = canonical_path.with_name(f"{canonical_path.stem} - refreshed {counter}{canonical_path.suffix}")
            counter += 1
        self.audit.setdefault("outputWarnings", []).append(
            {
                "kind": "locked-canonical-output",
                "canonicalPath": rel_posix(canonical_path, PROJECT_ROOT),
                "fallbackPath": rel_posix(fallback, PROJECT_ROOT),
            }
        )
        return fallback

    def new_document(self) -> Document:
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)

        styles = document.styles
        styles["Normal"].font.name = "Trebuchet MS"
        styles["Normal"].font.size = Pt(11.5)
        styles["Normal"].font.color.rgb = RGBColor.from_string(SITE_TEXT)
        styles["Normal"].paragraph_format.space_after = Pt(6)
        styles["Normal"].paragraph_format.line_spacing = 1.3

        styles["Title"].font.name = "Trebuchet MS"
        styles["Title"].font.size = Pt(27)
        styles["Title"].font.bold = False
        styles["Title"].font.color.rgb = RGBColor.from_string("FFFFFF")

        styles["Subtitle"].font.name = "Trebuchet MS"
        styles["Subtitle"].font.size = Pt(12)
        styles["Subtitle"].font.color.rgb = RGBColor.from_string(SITE_MUTED)

        heading_specs = {
            "Heading 1": (20, SITE_PURPLE, 12, 8),
            "Heading 2": (16, SITE_PURPLE, 8, 7),
            "Heading 3": (12, SITE_MUTED, 6, 4),
            "Heading 4": (11, SITE_MUTED, 6, 4),
        }
        for style_name, (size, color, before, after) in heading_specs.items():
            style = styles[style_name]
            style.font.name = "Lucida Grande"
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = RGBColor.from_string(color)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True

        if "Course Note" not in styles:
            note_style = styles.add_style("Course Note", 1)
            note_style.base_style = styles["Normal"]
            note_style.font.name = "Trebuchet MS"
            note_style.font.size = Pt(10)
            note_style.font.color.rgb = RGBColor(31, 58, 95)
            note_style.paragraph_format.space_after = Pt(6)

        if "Small Caption" not in styles:
            caption = styles.add_style("Small Caption", 1)
            caption.base_style = styles["Normal"]
            caption.font.name = "Trebuchet MS"
            caption.font.size = Pt(9)
            caption.font.italic = True
            caption.font.color.rgb = RGBColor(85, 85, 85)
            caption.paragraph_format.space_after = Pt(6)

        if "Site Body" not in styles:
            site_body = styles.add_style("Site Body", 1)
            site_body.base_style = styles["Normal"]
            site_body.font.name = "Trebuchet MS"
            site_body.font.size = Pt(11.5)
            site_body.font.color.rgb = RGBColor.from_string(SITE_TEXT)
            site_body.paragraph_format.space_after = Pt(8)
            site_body.paragraph_format.line_spacing = 1.3

        if "Site H2" not in styles:
            site_h2 = styles.add_style("Site H2", 1)
            site_h2.base_style = styles["Normal"]
            site_h2.font.name = "Lucida Grande"
            site_h2.font.size = Pt(18)
            site_h2.font.bold = True
            site_h2.font.color.rgb = RGBColor.from_string(SITE_PURPLE)
            site_h2.paragraph_format.space_before = Pt(4)
            site_h2.paragraph_format.space_after = Pt(12)
            site_h2.paragraph_format.keep_with_next = True

        if "Site H3" not in styles:
            site_h3 = styles.add_style("Site H3", 1)
            site_h3.base_style = styles["Normal"]
            site_h3.font.name = "Lucida Grande"
            site_h3.font.size = Pt(12)
            site_h3.font.bold = True
            site_h3.font.color.rgb = RGBColor.from_string(SITE_MUTED)
            site_h3.paragraph_format.space_before = Pt(4)
            site_h3.paragraph_format.space_after = Pt(4)
            site_h3.paragraph_format.keep_with_next = True

        return document

    def add_manifest_outline(self, document: Document, unit: ET.Element) -> None:
        document.add_heading("Brightspace Unit Structure", level=1)
        for child in item_children(unit):
            self.add_outline_line(document, child, depth=0)

    def add_outline_line(self, document: Document, item: ET.Element, depth: int) -> None:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.25 * depth)
        paragraph.add_run(item_title(item) or "Untitled item")
        for child in item_children(item):
            self.add_outline_line(document, child, depth + 1)

    def item_depth(self, root: ET.Element, target: ET.Element, depth: int = 0) -> int:
        if root is target:
            return depth
        for child in item_children(root):
            found = self.item_depth(child, target, depth + 1)
            if found >= 0:
                return found
        return -1

    def render_manifest_item(self, document: Document, item: ET.Element, root: ET.Element, depth: int) -> None:
        title = item_title(item) or "Untitled item"
        heading_level = min(max(depth + 1, 2), 4)

        ref = item.attrib.get("identifierref")
        record: dict[str, Any] = {
            "title": title,
            "identifier": item.attrib.get("identifier"),
            "identifierref": ref,
            "files": [],
            "children": [item_title(child) for child in item_children(item)],
        }
        self.audit["includedItems"].append(record)

        if ref:
            resource = self.resources.get(ref)
            if not resource:
                self.audit["missingResources"].append({"title": title, "identifierref": ref})
                document.add_heading(title, level=heading_level)
                document.add_paragraph(f"Missing manifest resource: {ref}", style="Course Note")
            else:
                files = resource.get("files") or []
                if not files:
                    document.add_heading(title, level=heading_level)
                    document.add_paragraph("This manifest resource does not list a source file.", style="Course Note")
                for href in files:
                    record["files"].append(href)
                    self.render_resource(document, href, title, heading_level)
        elif item_children(item):
            self.add_site_section_page(document, title)
        else:
            document.add_heading(title, level=heading_level)
            document.add_paragraph("No direct resource is attached to this Brightspace item.", style="Course Note")

        children = item_children(item)
        for child_index, child in enumerate(children):
            if child_index > 0:
                document.add_page_break()
            self.render_manifest_item(document, child, root=root, depth=depth + 1)

    def render_resource(self, document: Document, href: str, item_name: str, heading_level: int) -> None:
        ext = Path(href).suffix.lower()
        if not package_path_exists(self.zip_file, href):
            self.audit["missingResources"].append({"title": item_name, "href": href})
            document.add_paragraph(f"Missing source file: {href}", style="Course Note")
            return

        if ext in (".html", ".htm"):
            self.render_html(document, href, item_name, heading_level)
        elif ext == ".pdf":
            self.add_site_section_page(document, item_name)
            support_path = self.copy_support_file(href, UNIT_TITLE, item_name)
            self.add_support_reference(document, href, support_path, "Original PDF resource")
            self.render_pdf(document, href)
        elif ext in (".png", ".jpg", ".jpeg", ".gif"):
            self.add_site_section_page(document, item_name)
            self.add_image_from_href(document, href, max_width=SITE_FULL_IMAGE_WIDTH_IN)
        else:
            self.add_site_section_page(document, item_name)
            support_path = self.copy_support_file(href, UNIT_TITLE, item_name)
            self.add_support_reference(document, href, support_path, f"Original {ext.lstrip('.').upper() or 'resource'} file")

    def render_html(self, container: Any, href: str, item_name: str, heading_level: int) -> None:
        html_text = decode_html(self.zip_file.read(href))
        try:
            root = lxml_html.fromstring(html_text)
        except Exception as exc:
            self.audit["unresolvedAssets"].append({"href": href, "error": f"HTML parse failed: {exc}"})
            container.add_paragraph(f"HTML could not be parsed: {exc}", style="Course Note")
            return

        for node in root.xpath("//script|//style|//meta|//link|//title"):
            parent = node.getparent()
            if parent is not None:
                parent.remove(node)
        body = root.find("body") if root.tag.lower() != "body" else root
        if body is None:
            body = root
        context = RenderContext(base_href=href, heading_base=min(heading_level + 1, 4), source_title=item_name)
        header_nodes = root.xpath("//*[@id='header']//h1")
        content_nodes = root.xpath("//*[@id='content']")
        footer_nodes = root.xpath("//*[@id='footer']")
        if header_nodes and content_nodes:
            header_title = clean_text(header_nodes[0].text_content()) or item_name
            footer_text = clean_text(footer_nodes[0].text_content()) if footer_nodes else ""
            self.render_site_page(container, header_title, content_nodes[0], context, footer_text=footer_text)
            return
        self.render_site_page(container, item_name, body, context, footer_text="")

    def add_site_section_page(self, document: Document, title: str) -> None:
        frame_table, frame_cell = self.add_site_frame(document)
        self.add_site_header(frame_cell, title)

    def add_site_frame(self, container: Any) -> tuple[Any, Any]:
        table = container.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_width(table, width_dxa=SITE_FRAME_WIDTH_DXA, indent_dxa=0)
        set_table_grid(table, [SITE_FRAME_WIDTH_DXA])
        set_table_layout_fixed(table)
        set_table_borders(table, color=SITE_OUTER_BORDER, size="16", inside=False)
        cell = table.cell(0, 0)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_width(cell, SITE_FRAME_WIDTH_DXA)
        set_cell_margins(cell, top=180, bottom=180, start=180, end=180)
        self.clear_cell(cell)
        return table, cell

    def add_site_header(self, container: Any, title: str) -> None:
        table = container.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_width(table, width_dxa=SITE_TABLE_WIDTH_DXA, indent_dxa=0)
        set_table_grid(table, [SITE_TABLE_WIDTH_DXA])
        set_table_layout_fixed(table)
        clear_table_borders(table)
        cell = table.cell(0, 0)
        set_cell_width(cell, SITE_TABLE_WIDTH_DXA)
        set_cell_shading(cell, SITE_BLUE)
        set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
        self.clear_cell(cell)
        paragraph = cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(paragraph, before=0, after=0, line=None)
        run = paragraph.add_run(title)
        run.font.name = "Trebuchet MS"
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor.from_string("FFFFFF")
        container.add_paragraph()

    def site_header_image_bytes(self, title: str) -> bytes:
        width, height = 1600, 200
        image = Image.new("RGB", (width, height), f"#{SITE_BLUE}")
        draw = ImageDraw.Draw(image)
        font_size = 84
        while font_size >= 48:
            if TREBUCHET_FONT_PATH.exists():
                font = ImageFont.truetype(str(TREBUCHET_FONT_PATH), font_size)
            else:
                font = ImageFont.load_default()
            bbox = draw.textbbox((0, 0), title, font=font)
            text_width = bbox[2] - bbox[0]
            text_height = bbox[3] - bbox[1]
            if text_width <= width - 90 and text_height <= height - 42:
                break
            font_size -= 2
        x = 44
        y = (height - text_height) / 2 - bbox[1]
        draw.text((x, y), title, fill="#FFFFFF", font=font)
        output = BytesIO()
        image.save(output, format="PNG", dpi=(144, 144))
        return output.getvalue()

    def render_site_page(
        self,
        container: Any,
        header_title: str,
        content_node: HtmlElement,
        context: RenderContext,
        footer_text: str = "",
    ) -> None:
        _frame_table, frame_cell = self.add_site_frame(container)
        self.add_site_header(frame_cell, header_title)
        children = [child for child in content_node if isinstance(child.tag, str)]
        index = 0
        while index < len(children):
            child = children[index]
            if self.is_left_image_paragraph(child):
                index = self.render_left_image_cluster(frame_cell, children, index, context)
                continue
            self.render_block(frame_cell, child, context)
            index += 1
        if footer_text:
            footer_table = frame_cell.add_table(rows=1, cols=1)
            footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            footer_table.autofit = False
            set_table_width_pct(footer_table)
            set_table_grid(footer_table, [SITE_TABLE_WIDTH_DXA])
            set_table_layout_fixed(footer_table)
            clear_table_borders(footer_table)
            cell = footer_table.cell(0, 0)
            set_cell_width(cell, SITE_TABLE_WIDTH_DXA)
            set_cell_shading(cell, SITE_BLUE)
            set_cell_margins(cell, top=100, bottom=100, start=140, end=140)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(footer_text)
            run.font.name = "Trebuchet MS"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string("FFFFFF")

    def render_block(self, container: Any, node: HtmlElement, context: RenderContext) -> None:
        if not isinstance(node.tag, str):
            return
        name = node.tag.lower()
        if name in {"script", "style", "meta", "link", "title"}:
            return
        if self.is_back_to_top(node):
            return
        node_id = (node.get("id") or "").strip().lower()
        if node_id in CALLOUT_IDS:
            self.render_callout(container, node, context, *CALLOUT_IDS[node_id])
            return
        if node_id == "feature":
            self.render_feature_box(container, node, context)
            return
        if name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            text = clean_text(node.text_content())
            if text:
                if name == "h2":
                    paragraph = container.add_paragraph(style="Site H2")
                elif name in {"h3", "h4", "h5", "h6"}:
                    paragraph = container.add_paragraph(style="Site H3")
                else:
                    paragraph = add_container_heading(container, text, level=2)
                    return
                paragraph.add_run(text)
            return
        if name == "hr":
            paragraph = container.add_paragraph()
            paragraph.add_run(" " * 40).underline = True
            return
        if name in {"iframe", "video", "audio", "embed", "object"}:
            self.add_media_card(container, node, context)
            return
        if name == "img":
            paragraph = container.add_paragraph()
            self.apply_alignment(paragraph, node)
            embedded = self.add_image_to_paragraph(paragraph, node, context, max_width=SITE_FULL_IMAGE_WIDTH_IN)
            if not embedded and not clean_text(paragraph.text) and not paragraph.runs:
                self.remove_paragraph(paragraph)
            return
        if name in {"ul", "ol"}:
            self.render_list(container, node, context, ordered=(name == "ol"))
            return
        if name == "table":
            self.render_table(container, node, context)
            return
        if name == "pre":
            text = clean_text(node.text_content())
            if text:
                paragraph = container.add_paragraph()
                run = paragraph.add_run(text)
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            return
        nested_media = node.xpath(".//iframe|.//video|.//audio|.//embed|.//object")
        if nested_media and name in {"p", "center", "blockquote"}:
            leading_text = clean_text(node.text or "")
            if leading_text:
                paragraph = container.add_paragraph()
                self.apply_alignment(paragraph, node)
                self.add_text_run(paragraph, leading_text)
            for media_node in nested_media:
                self.add_media_card(container, media_node, context)
            return

        block_children = [child for child in node if isinstance(child.tag, str) and child.tag.lower() in BLOCK_TAGS]
        direct_text = clean_text(node.text or "")
        if name in {"p", "blockquote", "figcaption", "center"} or (direct_text and not block_children):
            paragraph = container.add_paragraph(style="Site Body")
            if name == "blockquote":
                paragraph.style = "Intense Quote"
            if name == "figcaption":
                paragraph.style = "Small Caption"
            self.apply_alignment(paragraph, node)
            self.add_inline(paragraph, node, context)
            if not clean_text(paragraph.text) and not paragraph.runs:
                self.remove_paragraph(paragraph)
            return

        if direct_text:
            paragraph = container.add_paragraph()
            self.add_text_run(paragraph, direct_text)
        for child in node:
            self.render_block(container, child, context)
            if child.tail and clean_text(child.tail):
                paragraph = container.add_paragraph()
                self.add_text_run(paragraph, clean_text(child.tail))

    def is_left_image_paragraph(self, node: HtmlElement) -> bool:
        if not isinstance(node.tag, str) or node.tag.lower() != "p":
            return False
        direct_images = [child for child in node if isinstance(child.tag, str) and child.tag.lower() == "img"]
        if not direct_images:
            return False
        first = direct_images[0]
        align = (first.get("align") or node.get("align") or "").strip().lower()
        class_name = (first.get("class") or "").strip().lower()
        style = (first.get("style") or "").strip().lower()
        return align == "left" or "img_left" in class_name or "float: left" in style

    def render_left_image_cluster(
        self,
        container: Any,
        siblings: list[HtmlElement],
        start_index: int,
        context: RenderContext,
    ) -> int:
        first_paragraph = siblings[start_index]
        image = next(
            child
            for child in first_paragraph
            if isinstance(child.tag, str) and child.tag.lower() == "img"
        )
        table = container.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_width(table, width_dxa=SITE_TABLE_WIDTH_DXA, indent_dxa=0)
        set_table_grid(table, [SITE_LEFT_IMAGE_DXA, SITE_TEXT_IMAGE_DXA])
        clear_table_borders(table)
        image_cell = table.cell(0, 0)
        text_cell = table.cell(0, 1)
        for cell, width in ((image_cell, SITE_LEFT_IMAGE_DXA), (text_cell, SITE_TEXT_IMAGE_DXA)):
            self.clear_cell(cell)
            set_cell_width(cell, width)
            set_cell_margins(cell, top=0, bottom=0, start=0, end=120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        image_paragraph = image_cell.add_paragraph()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self.add_image_to_paragraph(image_paragraph, image, context, max_width=2.7, display_width=2.7)

        consumed = 0
        max_cluster_paragraphs = 3
        index = start_index
        while index < len(siblings) and consumed < max_cluster_paragraphs:
            paragraph_node = siblings[index]
            if not isinstance(paragraph_node.tag, str) or paragraph_node.tag.lower() != "p":
                break
            if consumed > 0 and paragraph_node.xpath("./img"):
                break
            paragraph = text_cell.add_paragraph(style="Site Body")
            self.apply_alignment(paragraph, paragraph_node)
            self.add_inline_without_images(paragraph, paragraph_node, context)
            if not clean_text(paragraph.text) and not paragraph.runs:
                self.remove_paragraph(paragraph)
            consumed += 1
            index += 1

        container.add_paragraph()
        return index

    def add_inline_without_images(self, paragraph: Any, node: HtmlElement, context: RenderContext) -> None:
        if getattr(node, "text", None):
            self.add_text_run(paragraph, node.text)
        for child in node:
            if not isinstance(child.tag, str):
                continue
            name = child.tag.lower()
            if name == "img":
                if child.tail:
                    self.add_text_run(paragraph, child.tail)
                continue
            if name == "a":
                self.add_link_or_local_resource(paragraph, child, context)
            else:
                self.add_inline(paragraph, child, context)
            if child.tail:
                self.add_text_run(paragraph, child.tail)

    def render_feature_box(self, container: Any, node: HtmlElement, context: RenderContext) -> None:
        table = container.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_width(table, width_dxa=SITE_TABLE_WIDTH_DXA, indent_dxa=0)
        set_table_borders(table, color=SITE_FEATURE_BORDER, size="18", inside=False)
        cell = table.cell(0, 0)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_shading(cell, "FFFFFF")
        set_cell_margins(cell, top=130, bottom=130, start=150, end=150)
        self.clear_cell(cell)
        for child in node:
            self.render_block(cell, child, context)
        container.add_paragraph()

    def render_callout(self, container: Any, node: HtmlElement, context: RenderContext, label: str, fill: str) -> None:
        table = container.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_width(table, width_dxa=SITE_TABLE_WIDTH_DXA, indent_dxa=0)
        clear_table_borders(table)
        cell = table.cell(0, 0)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_shading(cell, fill)
        set_cell_margins(cell, top=150, bottom=140, start=170, end=170)
        self.clear_cell(cell)
        if label:
            heading = cell.add_paragraph(style="Course Note")
            run = heading.add_run(label)
            run.bold = True
            run.font.color.rgb = RGBColor(31, 58, 95)
        for child in node:
            self.render_block(cell, child, context)
        container.add_paragraph()

    def render_list(self, container: Any, node: HtmlElement, context: RenderContext, ordered: bool) -> None:
        style = "List Number" if ordered else "List Bullet"
        for li in [child for child in node if isinstance(child.tag, str) and child.tag.lower() == "li"]:
            paragraph = container.add_paragraph(style=style)
            self.add_inline(paragraph, li, context, skip_nested_blocks=True)
            for child in li:
                if isinstance(child.tag, str) and child.tag.lower() in {"ul", "ol"}:
                    self.render_list(container, child, context, ordered=(child.tag.lower() == "ol"))

    def render_table(self, container: Any, table_node: HtmlElement, context: RenderContext) -> None:
        rows = table_node.xpath(".//tr")
        if not rows:
            text = clean_text(table_node.text_content())
            if text:
                container.add_paragraph(text)
            return
        row_cells = [row.xpath("./th|./td") for row in rows]
        max_cols = max((len(cells) for cells in row_cells), default=0)
        if max_cols == 0:
            return
        width_px = int_attr(table_node.get("width"))
        width_dxa = min(SITE_TABLE_WIDTH_DXA, max(3600, (width_px * 15 if width_px else SITE_TABLE_WIDTH_DXA)))
        table = container.add_table(rows=len(rows), cols=max_cols)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        set_table_width(table, width_dxa=width_dxa, indent_dxa=120)
        border_color = css_color(table_node.get("bordercolor"), "D9D9D9")
        set_table_borders(table, color=border_color, size="8", inside=True)
        widths = []
        for col_index in range(max_cols):
            width = None
            for cells in row_cells:
                if col_index < len(cells):
                    width = int_attr(cells[col_index].get("width"))
                    if width:
                        break
            widths.append(width * 15 if width else width_dxa // max_cols)
        width_total = sum(widths)
        if width_total and width_total != width_dxa:
            widths = [max(600, round(width * width_dxa / width_total)) for width in widths]
            drift = width_dxa - sum(widths)
            widths[-1] += drift
        set_table_grid(table, widths)
        for row_index, cells in enumerate(row_cells):
            for col_index, cell_node in enumerate(cells):
                cell = table.cell(row_index, col_index)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                set_cell_width(cell, widths[col_index])
                set_cell_margins(cell, top=90, bottom=90, start=110, end=110)
                if cell_node.tag.lower() == "th":
                    set_cell_shading(cell, "E8EEF5")
                self.clear_cell(cell)
                if clean_text(cell_node.text or ""):
                    paragraph = cell.add_paragraph(style="Site Body")
                    self.add_text_run(paragraph, clean_text(cell_node.text or ""), bold=(cell_node.tag.lower() == "th"))
                for child in cell_node:
                    if isinstance(child.tag, str) and child.tag.lower() in BLOCK_TAGS:
                        self.render_block(cell, child, context)
                    else:
                        paragraph = cell.add_paragraph(style="Site Body")
                        self.add_inline(paragraph, child, context)
                    if child.tail and clean_text(child.tail):
                        paragraph = cell.add_paragraph(style="Site Body")
                        self.add_text_run(paragraph, clean_text(child.tail))
                if not cell.paragraphs:
                    cell.add_paragraph(clean_text(cell_node.text_content()), style="Site Body")
        container.add_paragraph()

    def add_inline(
        self,
        paragraph: Any,
        node: Any,
        context: RenderContext,
        bold: bool = False,
        italic: bool = False,
        underline: bool = False,
        skip_nested_blocks: bool = False,
    ) -> None:
        if getattr(node, "text", None):
            self.add_text_run(paragraph, node.text, bold=bold, italic=italic, underline=underline)

        for child in node:
            if not isinstance(child.tag, str):
                continue
            name = child.tag.lower()
            if self.is_back_to_top(child):
                continue
            child_bold = bold or name in {"strong", "b"}
            child_italic = italic or name in {"em", "i", "cite"}
            child_underline = underline or name == "u"

            if skip_nested_blocks and name in {"ul", "ol", "table"}:
                continue
            if name == "br":
                paragraph.add_run().add_break()
            elif name == "img":
                self.add_image_to_paragraph(paragraph, child, context, max_width=4.0)
            elif name == "a":
                self.add_link_or_local_resource(paragraph, child, context)
            elif name in {"iframe", "video", "audio", "embed", "object"}:
                src = child.get("src") or child.get("data") or ""
                title = child.get("title") or clean_text(child.text_content()) or "Embedded media"
                if src:
                    add_hyperlink(paragraph, title, src)
                else:
                    self.add_text_run(paragraph, title)
            elif name in {"sup", "sub"}:
                run = paragraph.add_run(clean_text(child.text_content()))
                run.font.superscript = name == "sup"
                run.font.subscript = name == "sub"
            elif name in BLOCK_TAGS and name not in {"span", "strong", "b", "em", "i", "u", "small"}:
                text = clean_text(child.text_content())
                if text:
                    self.add_text_run(paragraph, text, bold=child_bold, italic=child_italic, underline=child_underline)
            else:
                self.add_inline(
                    paragraph,
                    child,
                    context,
                    bold=child_bold,
                    italic=child_italic,
                    underline=child_underline,
                    skip_nested_blocks=skip_nested_blocks,
                )
            if child.tail:
                self.add_text_run(paragraph, child.tail, bold=bold, italic=italic, underline=underline)

    def add_text_run(self, paragraph: Any, text: str, bold: bool = False, italic: bool = False, underline: bool = False) -> None:
        if not text:
            return
        normalized = re.sub(r"[ \t\r\n]+", " ", text.replace("\xa0", " "))
        if not normalized.strip():
            if normalized:
                paragraph.add_run(" ")
            return
        run = paragraph.add_run(normalized)
        run.bold = bold
        run.italic = italic
        run.underline = underline
        run.font.name = "Trebuchet MS"
        run.font.color.rgb = RGBColor.from_string(SITE_TEXT)

    def record_image_source_link(self, href: str, context: RenderContext, image_count: int) -> None:
        self.audit.setdefault("imageSourceLinks", []).append(
            {
                "sourceHtml": context.base_href,
                "href": href,
                "imageCount": image_count,
            }
        )

    def add_link_or_local_resource(self, paragraph: Any, node: HtmlElement, context: RenderContext) -> None:
        href = node.get("href") or ""
        text = clean_text(node.text_content()) or href
        image_children = node.xpath(".//img")
        embedded_image_count = 0
        for image_child in image_children:
            if self.add_image_to_paragraph(paragraph, image_child, context, max_width=4.8):
                embedded_image_count += 1
        if not href:
            if not image_children:
                self.add_inline(paragraph, node, context)
            return
        if href.startswith("#"):
            if text and normalize_key(text) != "back to top":
                paragraph.add_run(text)
            return
        if re.match(r"^[a-zA-Z][a-zA-Z0-9+.-]*:", href):
            if image_children:
                self.record_image_source_link(href, context, embedded_image_count)
                return
            add_hyperlink(paragraph, text, href)
            return
        package_href = resolve_package_href(context.base_href, href, self.zip_file)
        if package_href:
            support_path = self.copy_support_file(package_href, UNIT_TITLE, text or Path(package_href).stem)
            relative = rel_posix(support_path, DOCX_DIR)
            if image_children:
                self.record_image_source_link(relative, context, embedded_image_count)
                return
            add_hyperlink(paragraph, text, relative)
            self.audit["linkedLocalResources"].append(
                {
                    "sourceHtml": context.base_href,
                    "href": href,
                    "resolvedPackageHref": package_href,
                    "supportPath": rel_posix(support_path, PROJECT_ROOT),
                }
            )
            return
        add_hyperlink(paragraph, text, href)
        self.audit["unresolvedAssets"].append({"sourceHtml": context.base_href, "href": href, "kind": "link"})

    def add_media_card(self, container: Any, node: HtmlElement, context: RenderContext) -> None:
        src = node.get("src") or node.get("data") or ""
        raw_title = node.get("title") or clean_text(node.text_content()) or "Embedded media"
        title = self.media_display_title(raw_title, src)
        handoff_url = self.media_handoff_url(src)
        link_url = handoff_url or src
        media_type = node.tag.lower()
        thumbnail_url = self.media_thumbnail_url(src)
        used_remote_thumbnail = False
        table = container.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        set_table_width(table, width_dxa=SITE_TABLE_WIDTH_DXA, indent_dxa=0)
        cell = table.cell(0, 0)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_shading(cell, "EAF4EA")
        set_cell_margins(cell, top=170, bottom=170, start=180, end=180)
        self.clear_cell(cell)
        preview = cell.add_paragraph()
        preview.alignment = WD_ALIGN_PARAGRAPH.CENTER
        preview_bytes = self.media_preview_image_bytes(title, src)
        used_remote_thumbnail = bool(thumbnail_url and self._last_media_preview_used_remote)
        self.add_hyperlinked_picture(preview, preview_bytes, width=6.65, url=link_url)
        if handoff_url:
            handoff = cell.add_paragraph(style="Course Note")
            add_hyperlink(handoff, handoff_url, handoff_url)
            self.audit.setdefault("mediaRawUrlLines", 0)
            self.audit["mediaRawUrlLines"] += 1
        container.add_paragraph()
        self.audit.setdefault("mediaPreviewCards", 0)
        self.audit["mediaPreviewCards"] += 1
        self.audit["mediaReferences"].append(
            {
                "sourceHtml": context.base_href,
                "sourceTitle": context.source_title,
                "type": media_type,
                "title": title,
                "src": src,
                "linkUrl": link_url,
                "handoffUrl": handoff_url,
                "thumbnailUrl": thumbnail_url,
                "usedRemoteThumbnail": used_remote_thumbnail,
            }
        )

    def add_hyperlinked_picture(self, paragraph: Any, image_bytes: bytes, width: float, url: str) -> None:
        run = paragraph.add_run()
        run.add_picture(BytesIO(image_bytes), width=Inches(width))
        if not url:
            return
        relationship_id = paragraph.part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), relationship_id)
        run_element = run._r
        parent = run_element.getparent()
        parent.remove(run_element)
        hyperlink.append(run_element)
        parent.append(hyperlink)

    def media_provider_label(self, src: str) -> str:
        host = (urlparse(src).netloc or "").casefold()
        if "youtube" in host or "youtu.be" in host:
            return "YouTube"
        if "ted.com" in host:
            return "TED"
        if host:
            return host.removeprefix("www.")
        return "Embedded media"

    def youtube_video_id(self, src: str) -> str | None:
        parsed = urlparse(src)
        host = parsed.netloc.casefold()
        path_parts = [part for part in parsed.path.split("/") if part]
        if "youtu.be" in host and path_parts:
            return path_parts[0]
        if "youtube" in host or "youtube-nocookie" in host:
            if len(path_parts) >= 2 and path_parts[0] in {"embed", "shorts", "v"}:
                return path_parts[1]
            query_id = parse_qs(parsed.query).get("v", [None])[0]
            if query_id:
                return query_id
        return None

    def media_thumbnail_url(self, src: str) -> str | None:
        video_id = self.youtube_video_id(src)
        if video_id:
            return f"https://img.youtube.com/vi/{video_id}/hqdefault.jpg"
        metadata = self.media_oembed_metadata(src)
        thumbnail = metadata.get("thumbnail_url") if metadata else None
        return clean_text(thumbnail) if isinstance(thumbnail, str) and thumbnail else None

    def media_display_title(self, title: str, src: str) -> str:
        generic_titles = {"embedded media", "youtube video player", "video player"}
        if title and normalize_key(title) not in generic_titles:
            return title
        metadata = self.media_oembed_metadata(src)
        metadata_title = metadata.get("title") if metadata else None
        if isinstance(metadata_title, str) and clean_text(metadata_title):
            return clean_text(metadata_title)
        provider = self.media_provider_label(src)
        if provider and provider != "Embedded media":
            return f"{provider} video"
        return "Embedded media"

    def media_handoff_url(self, src: str) -> str:
        if not src:
            return ""
        video_id = self.youtube_video_id(src)
        if video_id:
            return f"https://www.youtube.com/watch?v={video_id}"
        parsed = urlparse(src)
        if parsed.netloc.casefold() == "embed.ted.com":
            return src.replace("//embed.ted.com/", "//www.ted.com/", 1)
        return src

    def media_oembed_metadata(self, src: str) -> dict[str, Any] | None:
        parsed = urlparse(src)
        host = parsed.netloc.casefold()
        video_id = self.youtube_video_id(src)
        endpoint = ""
        if video_id:
            watch_url = f"https://www.youtube.com/watch?v={video_id}"
            endpoint = f"https://www.youtube.com/oembed?format=json&url={quote(watch_url, safe='')}"
        elif "ted.com" in host:
            endpoint = f"https://www.ted.com/services/v1/oembed.json?url={quote(src, safe='')}"
        else:
            return None
        if not hasattr(self, "_media_metadata_cache"):
            self._media_metadata_cache = {}
        if src in self._media_metadata_cache:
            return self._media_metadata_cache[src]
        metadata = self.fetch_remote_json(endpoint)
        self._media_metadata_cache[src] = metadata
        return metadata

    def fetch_remote_json(self, url: str) -> dict[str, Any] | None:
        try:
            request = Request(url, headers={"User-Agent": "Mozilla/5.0"})
            with urlopen(request, timeout=8) as response:
                data = response.read(512 * 1024)
            parsed = json.loads(data.decode("utf-8"))
            return parsed if isinstance(parsed, dict) else None
        except Exception:
            return None

    def fetch_remote_image_bytes(self, url: str) -> bytes | None:
        try:
            request = Request(url, headers={"User-Agent": "Mozilla/5.0"})
            with urlopen(request, timeout=8) as response:
                content_type = (response.headers.get("content-type") or "").casefold()
                data = response.read(8 * 1024 * 1024)
            if "image" not in content_type:
                return None
            return data
        except Exception:
            return None

    def fitted_video_thumbnail_bytes(self, image_bytes: bytes) -> bytes:
        image = Image.open(BytesIO(image_bytes))
        image.load()
        image = image.convert("RGB")
        target_w, target_h = 1280, 720
        scale = max(target_w / image.width, target_h / image.height)
        new_size = (max(1, int(image.width * scale)), max(1, int(image.height * scale)))
        resample = getattr(getattr(Image, "Resampling", Image), "LANCZOS")
        image = image.resize(new_size, resample=resample)
        left = max(0, (image.width - target_w) // 2)
        top = max(0, (image.height - target_h) // 2)
        image = image.crop((left, top, left + target_w, top + target_h))
        draw = ImageDraw.Draw(image, "RGBA")
        center_x, center_y = target_w // 2, target_h // 2
        draw.ellipse((center_x - 82, center_y - 82, center_x + 82, center_y + 82), fill=(255, 255, 255, 235))
        draw.polygon(
            [
                (center_x - 24, center_y - 42),
                (center_x - 24, center_y + 42),
                (center_x + 48, center_y),
            ],
            fill=(31, 52, 64, 255),
        )
        output = BytesIO()
        image.save(output, format="PNG", dpi=(144, 144))
        return output.getvalue()

    def media_preview_image_bytes(self, title: str, src: str) -> bytes:
        self._last_media_preview_used_remote = False
        thumbnail_url = self.media_thumbnail_url(src)
        if thumbnail_url:
            thumbnail_bytes = self.fetch_remote_image_bytes(thumbnail_url)
            if thumbnail_bytes:
                try:
                    self._last_media_preview_used_remote = True
                    return self.fitted_video_thumbnail_bytes(thumbnail_bytes)
                except Exception:
                    self._last_media_preview_used_remote = False
        width, height = 1280, 720
        image = Image.new("RGB", (width, height), "#1F3440")
        draw = ImageDraw.Draw(image)
        for y in range(height):
            blend = y / max(height - 1, 1)
            r = int(31 + blend * 28)
            g = int(52 + blend * 56)
            b = int(64 + blend * 62)
            draw.line([(0, y), (width, y)], fill=(r, g, b))
        try:
            title_font = ImageFont.truetype(str(TREBUCHET_FONT_PATH), 58) if TREBUCHET_FONT_PATH.exists() else ImageFont.load_default()
            label_font = ImageFont.truetype(str(TREBUCHET_FONT_PATH), 32) if TREBUCHET_FONT_PATH.exists() else ImageFont.load_default()
        except Exception:
            title_font = ImageFont.load_default()
            label_font = ImageFont.load_default()
        label = self.media_provider_label(src)
        draw.text((68, 58), label.upper(), fill="#D7EEF0", font=label_font)
        words = clean_text(title).split()
        lines: list[str] = []
        current = ""
        for word in words:
            candidate = f"{current} {word}".strip()
            bbox = draw.textbbox((0, 0), candidate, font=title_font)
            if bbox[2] - bbox[0] <= width - 470:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = word
            if len(lines) == 3:
                break
        if current and len(lines) < 3:
            lines.append(current)
        for index, line in enumerate(lines[:3]):
            draw.text((68, 160 + index * 70), line, fill="#FFFFFF", font=title_font)
        center_x, center_y = width - 210, height // 2
        draw.ellipse((center_x - 82, center_y - 82, center_x + 82, center_y + 82), fill="#FFFFFF")
        draw.polygon(
            [
                (center_x - 24, center_y - 42),
                (center_x - 24, center_y + 42),
                (center_x + 48, center_y),
            ],
            fill="#1F3440",
        )
        output = BytesIO()
        image.save(output, format="PNG", dpi=(144, 144))
        return output.getvalue()

    def add_image_to_paragraph(
        self,
        paragraph: Any,
        img: HtmlElement,
        context: RenderContext,
        max_width: float,
        display_width: float | None = None,
    ) -> bool:
        src = img.get("src") or ""
        if not src:
            return False
        package_href = resolve_package_href(context.base_href, src, self.zip_file)
        if not package_href:
            self.audit["unresolvedAssets"].append({"sourceHtml": context.base_href, "src": src, "kind": "image"})
            return False
        return self.add_image_run(
            paragraph,
            package_href,
            max_width=max_width,
            alt=img.get("alt") or "",
            display_width=display_width,
        )

    def add_image_from_href(self, container: Any, href: str, max_width: float) -> None:
        paragraph = container.add_paragraph()
        self.add_image_run(paragraph, href, max_width=max_width, alt=Path(href).name)

    def add_image_run(
        self,
        paragraph: Any,
        href: str,
        max_width: float,
        alt: str = "",
        display_width: float | None = None,
    ) -> bool:
        try:
            image_bytes = self.zip_file.read(href)
            normalized_bytes, width_px, height_px = self.normalized_image_bytes(image_bytes)
            if width_px <= 0 or height_px <= 0:
                raise ValueError("invalid image size")
            width_inches = min(max_width, display_width if display_width else width_px / 96)
            paragraph.add_run().add_picture(BytesIO(normalized_bytes), width=Inches(width_inches))
            self.audit["embeddedImages"] += 1
            return True
        except Exception as exc:
            if package_path_exists(self.zip_file, href):
                support_path = self.copy_support_file(href, UNIT_TITLE, "Unembedded image")
                self.audit["unresolvedAssets"].append(
                    {
                        "href": href,
                        "kind": "image",
                        "error": str(exc),
                        "preservedOriginal": rel_posix(support_path, PROJECT_ROOT),
                    }
                )
            else:
                self.audit["unresolvedAssets"].append({"href": href, "kind": "image", "error": str(exc)})
            return False

    def normalized_image_bytes(self, image_bytes: bytes) -> tuple[bytes, int, int]:
        image = Image.open(BytesIO(image_bytes))
        image.load()
        if image.mode not in ("RGB", "RGBA"):
            image = image.convert("RGBA" if "A" in image.getbands() else "RGB")
        output = BytesIO()
        image.save(output, format="PNG", dpi=(144, 144))
        return output.getvalue(), image.size[0], image.size[1]

    def render_pdf(self, container: Any, href: str) -> None:
        if fitz is None:
            container.add_paragraph("PDF rendering is unavailable in this runtime; the original PDF is linked above.", style="Course Note")
            self.audit["unresolvedAssets"].append({"href": href, "kind": "pdf-render", "error": "PyMuPDF unavailable"})
            return
        try:
            pdf = fitz.open(stream=self.zip_file.read(href), filetype="pdf")
        except Exception as exc:
            container.add_paragraph(f"PDF could not be opened for rendering: {exc}", style="Course Note")
            self.audit["unresolvedAssets"].append({"href": href, "kind": "pdf-render", "error": str(exc)})
            return
        for page_index in range(pdf.page_count):
            page = pdf.load_page(page_index)
            pixmap = page.get_pixmap(matrix=fitz.Matrix(1.6, 1.6), alpha=False)
            png_bytes = pixmap.tobytes("png")
            caption = container.add_paragraph(style="Small Caption")
            caption.add_run(f"Rendered PDF page {page_index + 1} of {pdf.page_count}")
            container.add_picture(BytesIO(png_bytes), width=Inches(6.3))
            self.audit["renderedPdfPages"] += 1
        pdf.close()

    def copy_support_file(self, href: str, unit_title: str, item_name: str) -> Path:
        if href in self._copied_support:
            return self._copied_support[href]
        ext = Path(href).suffix or ".resource"
        unit_dir = SUPPORT_DIR / safe_name(unit_title, 48)
        unit_dir.mkdir(parents=True, exist_ok=True)
        digest = hashlib.sha1(href.encode("utf-8")).hexdigest()[:8]
        basename = f"{safe_name(item_name, 58)}-{digest}{ext}"
        destination = unit_dir / basename
        counter = 2
        while destination.exists():
            destination = unit_dir / f"{Path(basename).stem}-{counter}{ext}"
            counter += 1
        with self.zip_file.open(href) as src, destination.open("wb") as dest:
            shutil.copyfileobj(src, dest, length=1024 * 1024)
        self._copied_support[href] = destination
        self.audit["supportFiles"].append(
            {
                "sourceHref": href,
                "outputPath": rel_posix(destination, PROJECT_ROOT),
                "bytes": destination.stat().st_size,
            }
        )
        return destination

    def add_support_reference(self, container: Any, href: str, support_path: Path, label: str) -> None:
        paragraph = container.add_paragraph(style="Course Note")
        paragraph.add_run(f"{label}: ").bold = True
        relative = rel_posix(support_path, DOCX_DIR)
        add_hyperlink(paragraph, relative, relative)
        details = container.add_paragraph(style="Small Caption")
        details.add_run("Original package path: ").bold = True
        details.add_run(href)

    def apply_alignment(self, paragraph: Any, node: HtmlElement) -> None:
        value = (node.get("align") or "").lower()
        style = (node.get("style") or "").lower()
        if "text-align: center" in style or value == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "text-align: right" in style or value == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def is_back_to_top(self, node: HtmlElement) -> bool:
        text = normalize_key(node.text_content() if hasattr(node, "text_content") else "")
        href = (node.get("href") or "").strip().lower()
        return text == "back to top" or (href == "#top" and "back" in text)

    def clear_cell(self, cell: Any) -> None:
        for paragraph in list(cell.paragraphs):
            self.remove_paragraph(paragraph)

    def remove_paragraph(self, paragraph: Any) -> None:
        element = paragraph._element
        element.getparent().remove(element)
        paragraph._p = paragraph._element = None

    def write_indexes(self) -> None:
        rows = list(self.audit["supportFiles"])
        csv_path = EXPORT_DIR / "supporting-files-index.csv"
        with csv_path.open("w", encoding="utf-8", newline="") as handle:
            writer = csv.DictWriter(handle, fieldnames=["sourceHref", "outputPath", "bytes"])
            writer.writeheader()
            writer.writerows(rows)

        structure = self.structure_record(self.find_unit())
        (META_DIR / "unit-1-conversion-map.json").write_text(
            json.dumps(
                {
                    "unitTitle": UNIT_TITLE,
                    "sourceZip": str(ZIP_PATH),
                    "structure": structure,
                },
                indent=2,
            ),
            encoding="utf-8",
        )

    def structure_record(self, item: ET.Element) -> dict[str, Any]:
        ref = item.attrib.get("identifierref")
        return {
            "title": item_title(item),
            "identifier": item.attrib.get("identifier"),
            "identifierref": ref,
            "files": (self.resources.get(ref, {}) or {}).get("files", []) if ref else [],
            "children": [self.structure_record(child) for child in item_children(item)],
        }

    def verify_output(self, output_path: Path) -> None:
        verification: dict[str, Any] = {
            "docxPath": rel_posix(output_path, PROJECT_ROOT),
            "docxBytes": output_path.stat().st_size,
            "validDocxZip": False,
            "documentXmlPresent": False,
            "embeddedImages": self.audit["embeddedImages"],
            "mediaReferences": len(self.audit["mediaReferences"]),
            "renderedPdfPages": self.audit["renderedPdfPages"],
            "supportFileCount": len(self.audit["supportFiles"]),
            "passed": False,
        }
        try:
            with zipfile.ZipFile(output_path) as docx_zip:
                verification["validDocxZip"] = True
                doc_xml = docx_zip.getinfo("word/document.xml")
                verification["documentXmlPresent"] = doc_xml.file_size > 0
                verification["wordMediaFiles"] = len(
                    [name for name in docx_zip.namelist() if name.startswith("word/media/")]
                )
        except Exception as exc:
            verification["error"] = str(exc)
        verification["passed"] = bool(
            verification["validDocxZip"]
            and verification["documentXmlPresent"]
            and verification["embeddedImages"] >= 1
            and verification["mediaReferences"] >= 1
        )
        self.audit["verification"] = verification
        (META_DIR / "unit-1-docx-export-verification.json").write_text(
            json.dumps(verification, indent=2), encoding="utf-8"
        )
        (META_DIR / "unit-1-docx-export-audit.json").write_text(
            json.dumps(self.audit, indent=2), encoding="utf-8"
        )
        if not verification["passed"]:
            raise SystemExit("DOCX verification failed. See meta/unit-1-docx-export-verification.json")

    def write_audit_markdown(self) -> None:
        support_bytes = sum(entry["bytes"] for entry in self.audit["supportFiles"])
        lines = [
            "# Social Studies 10-1 Unit 1 DOCX Export Audit",
            "",
            f"- Generated: {self.audit['generatedAt']}",
            f"- Source ZIP: `{self.audit['sourceZip']}`",
            f"- Unit: `{UNIT_TITLE}`",
            f"- Output: `{self.audit.get('outputPath')}`",
            f"- Output size: {self.audit.get('outputBytes', 0) / 1024 / 1024:.2f} MB",
            f"- Included Brightspace items: {len(self.audit['includedItems'])}",
            f"- Embedded package images: {self.audit['embeddedImages']}",
            f"- Preserved embedded media links: {len(self.audit['mediaReferences'])}",
            f"- Rendered direct PDF pages: {self.audit['renderedPdfPages']}",
            f"- Supporting files copied: {len(self.audit['supportFiles'])}",
            f"- Supporting file size: {support_bytes / 1024 / 1024:.2f} MB",
            "",
            "## Brightspace Item Order",
            "",
        ]
        for item in self.audit["includedItems"]:
            files = ", ".join(item.get("files") or [])
            lines.append(f"- `{item['title']}`{f' -> `{files}`' if files else ''}")
        if self.audit["mediaReferences"]:
            lines.extend(["", "## Embedded Media Preserved As Links", ""])
            for media in self.audit["mediaReferences"]:
                lines.append(f"- `{media['sourceTitle']}`: {media['title']} -> {media['src']}")
        if self.audit["unresolvedAssets"]:
            lines.extend(["", "## Unresolved / Fallback Items", ""])
            for asset in self.audit["unresolvedAssets"]:
                lines.append(f"- `{asset}`")
        lines.append("")
        (META_DIR / "unit-1-docx-export-audit.md").write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    if not ZIP_PATH.exists():
        raise SystemExit(f"Source ZIP not found: {ZIP_PATH}")
    with zipfile.ZipFile(ZIP_PATH) as zip_file:
        manifest_root = ET.fromstring(zip_file.read("imsmanifest.xml"))
        exporter = UnitOneDocxExporter(zip_file, manifest_root)
        output = exporter.build()
        print(f"Wrote {output}")


if __name__ == "__main__":
    main()
