from __future__ import annotations

import html
import json
import re
import shutil
import textwrap
import zipfile
from datetime import datetime
from io import BytesIO
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from lxml import html as lxml_html


REPO_ROOT = Path(__file__).resolve().parents[3]
PROJECT_DIR = Path(__file__).resolve().parents[1]
META_DIR = PROJECT_DIR / "meta"
WORKSPACE_DIR = PROJECT_DIR / "workspace"
RAW_DIR = PROJECT_DIR / "raw"

SOURCE_ZIP = Path(
    r"C:\Users\dean.guedo\Downloads\D2LExport_149443_24-25 _ Aboriginal Studies 30 _ Per 1(A-B) _ Sec S_202652511.zip"
)
COURSE_MATERIALS_DIR = Path(r"C:\Users\dean.guedo\Documents\ONLINE COURSES\Aboriginal Studies 30\Course Materials")
THEME_1_BOOKLET_PDF = Path(r"C:\Users\dean.guedo\Documents\AB STUDIES\ab30theme1BOOKLET.pdf")
THEME_1_ORAL_TRADITION_PDF = Path(r"C:\Users\dean.guedo\Downloads\Indigenous-Worldviews.pdf")
THEME_1_ORAL_TRADITION_URL = "./assets/theme-1/readings/indigenous-worldviews.pdf"

IMS_NS = {"ims": "http://www.imsglobal.org/xsd/imscp_v1p1"}
D2L_NS = "http://desire2learn.com/xsd/d2lcp_v2p0"


def rel(path: Path) -> str:
    return str(path.resolve())


def rel_workspace(path: Path) -> str:
    return "./" + path.relative_to(WORKSPACE_DIR).as_posix()


def item_title(item: ET.Element) -> str:
    title = item.find("ims:title", IMS_NS)
    return clean_text(title.text or "") if title is not None else ""


def attr_local(item: ET.Element, name: str) -> str:
    return item.get(name) or item.get(f"{{{D2L_NS}}}{name}") or ""


def safe_slug(value: str, fallback: str = "item") -> str:
    lowered = value.strip().lower()
    lowered = lowered.replace("&", "and")
    lowered = re.sub(r"[^a-z0-9]+", "-", lowered)
    lowered = re.sub(r"-+", "-", lowered).strip("-")
    return lowered or fallback


def clean_text(value: str) -> str:
    value = html.unescape(value or "")
    value = value.replace("\xa0", " ")
    value = re.sub(r"\s+", " ", value).strip()
    value = value.replace("MÃ©tis", "Métis")
    value = value.replace("MÃ¢", "Mâ")
    value = value.replace("Ã¢", "â")
    value = value.replace("Ã´", "ô")
    return value


def clean_student_html(raw_html: str) -> str:
    raw_html = html.unescape(raw_html or "")
    raw_html = raw_html.replace("\xa0", " ")
    raw_html = raw_html.replace("BrightSpace page", "course resources")
    raw_html = raw_html.replace("Brightspace page", "course resources")
    raw_html = raw_html.replace("BrightSpace", "course")
    raw_html = raw_html.replace("Brightspace", "course")
    raw_html = raw_html.replace("Aboriginal Studies 10 Theme 1 Assignment Booklet", "Aboriginal Studies 30 assignment")
    raw_html = raw_html.replace("Aboriginal Studies 10 Theme 1 Assignment", "Aboriginal Studies 30 assignment")
    raw_html = raw_html.replace("Please return here to upload (hand in) the assignment when it is complete.", "Submit the completed assignment using your course submission process.")
    raw_html = raw_html.replace("Please make a copy of this assignment and replace \"Copy of\" in the file name to your name.", "Make a copy of the assignment document and rename it with your name.")

    if not clean_text(re.sub(r"<[^>]+>", " ", raw_html)):
        raw_html = "<p>Open the linked assignment document and follow the instructions.</p>"

    wrapper = lxml_html.fromstring(f"<div>{raw_html}</div>")
    for bad in wrapper.xpath(".//script|.//style"):
        bad.drop_tree()
    for element in wrapper.iter():
        for attribute in list(element.attrib):
            if attribute in {"style", "class", "id", "lang"}:
                del element.attrib[attribute]

    output = lxml_html.tostring(wrapper, encoding="unicode", method="html")
    output = re.sub(r"^<div>|</div>$", "", output.strip())
    output = re.sub(r"<p>\s*</p>", "", output)
    output = re.sub(r"\s{2,}", " ", output)
    return output.strip() or "<p>Open the linked assignment document and follow the instructions.</p>"


def html_to_plain_blocks(cleaned_html: str) -> list[tuple[str, str]]:
    wrapper = lxml_html.fromstring(f"<div>{cleaned_html}</div>")
    blocks: list[tuple[str, str]] = []
    for element in wrapper.iter():
        tag = element.tag.lower() if isinstance(element.tag, str) else ""
        if tag in {"p", "li"}:
            text = clean_text(element.text_content())
            if text:
                blocks.append(("bullet" if tag == "li" else "paragraph", text))
    if not blocks:
        text = clean_text(wrapper.text_content())
        if text:
            blocks.append(("paragraph", text))
    return blocks


def summary_from_html(cleaned_html: str) -> str:
    wrapper = lxml_html.fromstring(f"<div>{cleaned_html}</div>")
    text = clean_text(wrapper.text_content())
    if len(text) > 180:
        return text[:177].rstrip() + "..."
    return text


def copy_file(source: Path, target: Path) -> None:
    target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, target)


def copy_zip_entry(zip_file: zipfile.ZipFile, source_name: str, target: Path) -> None:
    target.parent.mkdir(parents=True, exist_ok=True)
    with zip_file.open(source_name) as source, target.open("wb") as destination:
        shutil.copyfileobj(source, destination)


def read_manifest(zip_file: zipfile.ZipFile) -> ET.Element:
    return ET.fromstring(zip_file.read("imsmanifest.xml"))


def read_resources(manifest_root: ET.Element) -> dict[str, str]:
    resources: dict[str, str] = {}
    for resource in manifest_root.findall(".//ims:resources/ims:resource", IMS_NS):
        identifier = resource.get("identifier") or ""
        href = resource.get("href") or ""
        if identifier:
            resources[identifier] = href
    return resources


def top_items(manifest_root: ET.Element) -> list[ET.Element]:
    organization = manifest_root.find(".//ims:organizations/ims:organization", IMS_NS)
    return list(organization.findall("ims:item", IMS_NS)) if organization is not None else []


def child_items(item: ET.Element) -> list[ET.Element]:
    return list(item.findall("ims:item", IMS_NS))


def is_dropbox_item(item: ET.Element) -> bool:
    return "Dropbox" in (item.get("resource_type_key") or "")


def infer_kind(title: str, href: str) -> str:
    text = f"{title} {href}".lower()
    if "youtube.com" in text or "youtu.be" in text:
        return "video"
    if "archive.org" in text or "cbc.ca/player" in text or "documentary" in text or "film" in text:
        return "film"
    if href.lower().endswith(".pdf"):
        return "reading"
    if href.lower().endswith((".html", ".htm")):
        return "article"
    return "resource"


def is_film_room_url(title: str, href: str) -> bool:
    text = f"{title} {href}".lower()
    return any(
        marker in text
        for marker in (
            "youtube.com",
            "youtu.be",
            "archive.org",
            "cbc.ca/player",
            "documentary",
            "rabbit-proof",
            "film",
        )
    )


def localize_resource(zip_file: zipfile.ZipFile, title: str, href: str, audit: dict) -> str:
    if not href or re.match(r"^https?://", href, re.I):
        return href
    if href.startswith("/d2l/"):
        return href
    if href.lower().endswith(".html"):
        target = WORKSPACE_DIR / "assets" / "source-pages" / f"{safe_slug(title)}.html"
        copy_zip_entry(zip_file, href, target)
        audit["copiedSourcePages"].append({"title": title, "sourceHref": href, "path": rel_workspace(target)})
        return rel_workspace(target)
    if href.lower().endswith(".pdf"):
        target = WORKSPACE_DIR / "assets" / "library" / f"{safe_slug(title)}.pdf"
        copy_zip_entry(zip_file, href, target)
        audit["copiedZipPdfs"].append({"title": title, "sourceHref": href, "path": rel_workspace(target)})
        return rel_workspace(target)
    return href


def build_library(zip_file: zipfile.ZipFile, audit: dict) -> list[dict]:
    library_dir = WORKSPACE_DIR / "assets" / "library"
    chapter_specs = [
        ("chapter-1", "Chapter 1", "CHAPTER 1.pdf", "chapter"),
        ("chapter-2", "Chapter 2", "CHAPTER 2.pdf", "chapter"),
        ("chapter-3", "Chapter 3", "CHAPTER 3.pdf", "chapter"),
        ("chapter-4", "Chapter 4", "CHAPTER 4.pdf", "chapter"),
        ("chapter-5", "Chapter 5", "CHAPTER 5.pdf", "chapter"),
        ("chapter-6", "Chapter 6", "CHAPTER 6.pdf", "chapter"),
        ("chapter-7", "Chapter 7", "CHAPTER 7.pdf", "chapter"),
        ("textbook", "Textbook", "ab30textbook.pdf", "textbook"),
        ("glossary", "Glossary", "GLOSSARY.pdf", "glossary"),
    ]
    library_items: list[dict] = []
    for code, title, source_name, kind in chapter_specs:
        source = COURSE_MATERIALS_DIR / source_name
        target = library_dir / f"{code}.pdf"
        copy_file(source, target)
        library_items.append(
            {
                "id": code,
                "code": "CH" + code.rsplit("-", 1)[-1] if kind == "chapter" else code.upper(),
                "title": title,
                "kind": kind,
                "description": "Open this PDF in the course viewer or download it for offline reading.",
                "file": rel_workspace(target),
            }
        )
        audit["copiedCourseMaterialPdfs"].append({"title": title, "sourcePath": rel(source), "path": rel_workspace(target)})

    zip_pdf_specs = [
        ("critical-response-criteria", "Critical Response Criteria", "Critical Response Criteria Aboriginal Studies (1).pdf", "support"),
        ("critical-response-rubric", "Critical Response Rubric", "Critical Response Rubric.pdf", "support"),
        ("halfbreed-maria-campbell", "Halfbreed - Maria Campbell", "Halfbreed - Maria Campbell (1).pdf", "novel"),
    ]
    for item_id, title, source_name, kind in zip_pdf_specs:
        if source_name in zip_file.namelist():
            target = library_dir / f"{item_id}.pdf"
            copy_zip_entry(zip_file, source_name, target)
            library_items.append(
                {
                    "id": item_id,
                    "code": "NOVEL" if kind == "novel" else "SUPPORT",
                    "title": title,
                    "kind": kind,
                    "description": "Student-facing support material from the source course.",
                    "file": rel_workspace(target),
                }
            )
            audit["copiedZipPdfs"].append({"title": title, "sourceHref": source_name, "path": rel_workspace(target)})

    return library_items


def build_units_and_film_room(zip_file: zipfile.ZipFile, manifest_root: ET.Element, resources: dict[str, str], audit: dict) -> tuple[list[dict], list[dict], dict[str, str]]:
    units: list[dict] = []
    film_room: list[dict] = []
    dropbox_unit_by_resource_code: dict[str, str] = {}
    seen_film_urls: set[str] = set()

    for top_index, top in enumerate(top_items(manifest_root), start=1):
        title = item_title(top)
        if not title.startswith("Theme "):
            continue

        unit_id = f"theme-{len(units) + 1}"
        unit_items: list[dict] = []
        for child in child_items(top):
            child_title = item_title(child)
            resource_code = attr_local(child, "resource_code")
            if resource_code and is_dropbox_item(child):
                dropbox_unit_by_resource_code[resource_code] = unit_id
                continue
            if is_dropbox_item(child):
                continue

            href = resources.get(child.get("identifierref") or "", "")
            if not href and not child_title:
                continue
            localized_href = localize_resource(zip_file, child_title, href, audit)
            kind = infer_kind(child_title, href)
            item_record = {
                "id": f"{unit_id}-{safe_slug(child_title, 'resource')}",
                "title": child_title,
                "kind": kind,
                "url": localized_href,
                "external": bool(re.match(r"^https?://", localized_href, re.I)),
            }
            unit_items.append(item_record)
            audit["unitItems"].append({"unitTitle": title, "title": child_title, "href": href, "kind": kind})

            if href and is_film_room_url(child_title, href) and href not in seen_film_urls:
                seen_film_urls.add(href)
                film_room.append(
                    {
                        "id": f"film-{safe_slug(child_title, str(len(film_room) + 1))}",
                        "title": child_title,
                        "unitId": unit_id,
                        "moduleCode": f"T{len(units) + 1}",
                        "moduleLabel": title,
                        "kind": kind if kind in {"video", "film"} else "media",
                        "url": href,
                        "description": f"Media resource from {title}.",
                    }
                )

        units.append(
            {
                "id": unit_id,
                "code": f"T{len(units) + 1}",
                "title": title,
                "description": f"Resources and readings from {title}.",
                "items": unit_items,
            }
        )
        audit["includedUnits"].append({"title": title, "itemCount": len(unit_items), "sourceIndex": top_index})

    return units, film_room, dropbox_unit_by_resource_code


def map_link(name: str, url: str) -> dict | None:
    clean_name = clean_text(name)
    clean_url = html.unescape(url or "")
    if not clean_url:
        return None
    if clean_url.startswith("/d2l/"):
        lowered = clean_name.lower()
        if "criteria" in lowered:
            clean_url = "./assets/library/critical-response-criteria.pdf"
        elif "rubric" in lowered:
            clean_url = "./assets/library/critical-response-rubric.pdf"
        else:
            return None
    return {"name": clean_name, "url": clean_url}


def infer_assignment_unit(name: str, fallback_units: list[dict]) -> str:
    match = re.search(r"(?:Theme|Assignment)\s*(\d)", name, re.I)
    if match:
        return f"theme-{match.group(1)}"
    match = re.match(r"(\d)\.", name.strip())
    if match:
        return f"theme-{match.group(1)}"
    return fallback_units[0]["id"] if fallback_units else "theme-1"


def write_assignment_docx(assignment: dict, target: Path) -> None:
    target.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading(assignment["title"], level=1)
    document.add_paragraph(assignment["unitTitle"])
    for block_type, text in html_to_plain_blocks(assignment["instructionsHtml"]):
        if block_type == "bullet":
            document.add_paragraph(text, style="List Bullet")
        else:
            document.add_paragraph(text)
    links = assignment.get("links") or []
    if links:
        document.add_heading("Resources", level=2)
        for link in links:
            document.add_paragraph(f"{link['name']}: {link['url']}")
    document.save(target)


def build_assignments(dropbox_unit_by_resource_code: dict[str, str], units: list[dict], audit: dict, zip_file: zipfile.ZipFile) -> list[dict]:
    dropbox_root = ET.fromstring(zip_file.read("dropbox_d2l.xml"))
    assignments: list[dict] = []
    unit_title_by_id = {unit["id"]: unit["title"] for unit in units}
    html_dir = WORKSPACE_DIR / "assets" / "assignments" / "html"
    docx_dir = WORKSPACE_DIR / "assets" / "assignments" / "docx"

    for folder in dropbox_root.findall("folder"):
        title = clean_text(folder.get("name") or "")
        if not title:
            continue
        instructions_node = folder.find("instructions/text")
        instructions_html = clean_student_html(instructions_node.text if instructions_node is not None else "")
        resource_code = folder.get("resource_code") or ""
        unit_id = dropbox_unit_by_resource_code.get(resource_code) or infer_assignment_unit(title, units)
        assignment_id = safe_slug(title, f"assignment-{len(assignments) + 1}")
        links: list[dict] = []
        for link in folder.findall("attachment_set/links/link"):
            mapped = map_link(link.get("name") or "", link.get("url") or "")
            if mapped and mapped not in links:
                links.append(mapped)
        replace_oral_tradition_links(links)

        html_path = html_dir / f"{assignment_id}.html"
        docx_path = docx_dir / f"{assignment_id}.docx"
        assignment = {
            "id": assignment_id,
            "title": title,
            "unitId": unit_id,
            "unitTitle": unit_title_by_id.get(unit_id, "Course Assignment"),
            "summary": summary_from_html(instructions_html),
            "instructionsHtml": instructions_html,
            "htmlPath": rel_workspace(html_path),
            "docxPath": rel_workspace(docx_path),
            "links": links,
        }
        html_path.parent.mkdir(parents=True, exist_ok=True)
        html_path.write_text(assignment_page_html(assignment), encoding="utf-8")
        write_assignment_docx(assignment, docx_path)
        assignments.append(assignment)
        audit["assignments"].append(
            {
                "title": title,
                "unitId": unit_id,
                "isHiddenInSource": folder.get("is_hidden") == "true",
                "docxPath": rel_workspace(docx_path),
                "linkCount": len(links),
            }
        )

    return assignments


def assignment_page_html(assignment: dict) -> str:
    links = "".join(
        f'<li><a href="{html.escape(link["url"])}" target="_blank" rel="noopener noreferrer">{html.escape(link["name"])}</a></li>'
        for link in assignment.get("links", [])
    )
    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(assignment["title"])}</title>
  <style>
    body {{ margin: 0; padding: 32px; font-family: Arial, sans-serif; line-height: 1.6; color: #1f2933; background: #ffffff; }}
    main {{ max-width: 820px; margin: 0 auto; }}
    h1 {{ font-size: 28px; margin: 0 0 8px; }}
    .unit {{ color: #53606d; margin: 0 0 24px; }}
    a {{ color: #0f766e; font-weight: 700; }}
  </style>
</head>
<body>
  <main>
    <h1>{html.escape(assignment["title"])}</h1>
    <p class="unit">{html.escape(assignment["unitTitle"])}</p>
    {assignment["instructionsHtml"]}
    {f"<h2>Resources</h2><ul>{links}</ul>" if links else ""}
  </main>
</body>
</html>
"""


def theme_resource(title: str, kind: str, url: str, action_label: str) -> dict:
    return {
        "title": title,
        "kind": kind,
        "url": url,
        "actionLabel": action_label,
    }


def copy_theme_one_oral_tradition_reading(audit: dict) -> str:
    target = WORKSPACE_DIR / "assets" / "theme-1" / "readings" / "indigenous-worldviews.pdf"
    if THEME_1_ORAL_TRADITION_PDF.exists():
        copy_file(THEME_1_ORAL_TRADITION_PDF, target)
        audit.setdefault("themeOneReadings", []).append(
            {
                "title": "Walking Together: The Oral Tradition",
                "sourcePath": rel(THEME_1_ORAL_TRADITION_PDF),
                "path": rel_workspace(target),
                "status": "copied-local-pdf",
            }
        )
        return rel_workspace(target)

    audit.setdefault("themeOneReadings", []).append(
        {
            "title": "Walking Together: The Oral Tradition",
            "sourcePath": rel(THEME_1_ORAL_TRADITION_PDF),
            "status": "missing-local-pdf",
        }
    )
    return "https://www.learnalberta.ca/content/aswt/oral_tradition/documents/oral_tradition.pdf"


def replace_oral_tradition_links(links: list[dict]) -> None:
    for link in links:
        title = str(link.get("name", ""))
        url = str(link.get("url", ""))
        if "Walking Together: The Oral Tradition" in title or "learnalberta.ca/content/aswt/oral_tradition" in url:
            link["url"] = THEME_1_ORAL_TRADITION_URL


def replace_oral_tradition_assignment_links(assignments: list[dict]) -> None:
    for assignment in assignments:
        replace_oral_tradition_links(assignment.get("links", []))


def activity_prompt(
    prompt_id: str,
    label: str,
    rows: int = 3,
    number: int | None = None,
    kind: str = "shortAnswer",
    **extra: object,
) -> dict:
    prompt = {
        "id": prompt_id,
        "label": label,
        "rows": rows,
        "kind": kind,
    }
    if number is not None:
        prompt["number"] = str(number)
    prompt.update(extra)
    return prompt


def q(number: int, label: str, rows: int = 3, **extra: object) -> dict:
    return activity_prompt(f"q{number}", label, rows, number, **extra)


def fill_blank_q(number: int, label: str, parts: list[str], blanks: list[str] | None = None) -> dict:
    blank_labels = blanks or [f"Blank {index + 1}" for index in range(max(1, len(parts) - 1))]
    return activity_prompt(
        f"q{number}",
        label,
        1,
        number,
        "fillBlank",
        textParts=parts,
        blanks=[{"id": f"blank-{index + 1}", "label": blank_label} for index, blank_label in enumerate(blank_labels)],
    )


def multiple_choice_q(number: int, label: str, choices: list[str]) -> dict:
    return activity_prompt(f"q{number}", label, 1, number, "multipleChoice", choices=choices)


def table_q(number: int, label: str, rows: list[str], columns: list[str]) -> dict:
    prompt = activity_prompt(f"q{number}", label, 1, number, "table", columns=columns)
    prompt["rows"] = rows
    return prompt


def extract_theme_one_images(audit: dict) -> dict[int, list[dict]]:
    page_images: dict[int, list[dict]] = {}
    if not THEME_1_BOOKLET_PDF.exists():
        audit["themeOneImageExtraction"] = {"status": "missing-pdf", "sourcePath": rel(THEME_1_BOOKLET_PDF)}
        return page_images

    try:
        import fitz
        from PIL import Image
    except Exception as error:
        audit["themeOneImageExtraction"] = {"status": "unavailable", "error": str(error)}
        return page_images

    image_dir = WORKSPACE_DIR / "assets" / "theme-1" / "images"
    image_dir.mkdir(parents=True, exist_ok=True)
    extracted = 0
    try:
        document = fitz.open(THEME_1_BOOKLET_PDF)
        for page_number, page in enumerate(document, start=1):
            for image_index, image_ref in enumerate(page.get_images(full=True), start=1):
                xref = image_ref[0]
                image_info = document.extract_image(xref)
                image_bytes = image_info.get("image")
                if not image_bytes:
                    continue
                with Image.open(BytesIO(image_bytes)) as image:
                    image = image.convert("RGB")
                    image.thumbnail((900, 900))
                    target = image_dir / f"theme-1-page-{page_number:02d}-image-{image_index:02d}.jpg"
                    image.save(target, "JPEG", quality=84, optimize=True)
                page_images.setdefault(page_number, []).append(
                    {
                        "src": rel_workspace(target),
                        "alt": "Theme 1 reference image",
                        "page": page_number,
                    }
                )
                extracted += 1
    except Exception as error:
        audit["themeOneImageExtraction"] = {"status": "error", "error": str(error)}
        return page_images

    audit["themeOneImageExtraction"] = {
        "status": "ok",
        "sourcePath": rel(THEME_1_BOOKLET_PDF),
        "imageCount": extracted,
        "pages": sorted(page_images.keys()),
    }
    return page_images


def images_for_pages(page_images: dict[int, list[dict]], pages: list[int]) -> list[dict]:
    images: list[dict] = []
    for page in pages:
        images.extend(page_images.get(page, []))
    return images


def build_theme_one_sections(page_images: dict[int, list[dict]], oral_tradition_url: str, road_allowance_url: str, metis_governance_url: str) -> list[dict]:
    return [
        {
            "id": "assignment-1-1",
            "title": "Assignment 1.1: Oral Tradition",
            "instructions": "Use the oral tradition reading before responding.",
            "sourceRef": "Walking Together reading.",
            "images": [],
            "prompts": [
                activity_prompt(
                    "assignment-1-1",
                    "Storytelling is part of Aboriginal oral tradition. Based on what you have read, define oral tradition, explain its importance, and explain its purpose. Respond in paragraph form using 2-3 paragraphs.",
                    8,
                    resources=[theme_resource("Walking Together: The Oral Tradition", "reading", oral_tradition_url, "Open Reading")],
                ),
            ],
        },
        {
            "id": "inherent-rights",
            "title": "Inherent Rights, Nations, Peoples, and Treaties",
            "instructions": "Complete questions 1-34 using the Theme 1 readings and Chapter 1.",
            "sourceRef": "Textbook pages 2-35.",
            "images": [],
            "prompts": [
                fill_blank_q(
                    1,
                    "Colonization, often with force, has led to ancient civilizations and ways of life being ________________.",
                    ["Colonization, often with force, has led to ancient civilizations and ways of life being ", "."],
                ),
                fill_blank_q(
                    2,
                    "Unlike the rest of Canada, the Dene and the Inuit are a ________________ of the population of the N.W.T.",
                    ["Unlike the rest of Canada, the Dene and the Inuit are a ", " of the population of the N.W.T."],
                ),
                q(3, "What are the Dene struggling for?"),
                q(4, "What does it mean when the Dene call for a just land settlement?", 4),
                q(5, "In both Haudenosaunee and English, what does the concept of nation mean?", 4),
                q(6, "What are some other ways of describing a nation?"),
                q(7, "What are two differences between immigrants and Aboriginal peoples?", 4),
                q(8, "What is the difference between a person and a people?", 4),
                q(9, "What are human rights?"),
                q(10, "Why and how do Aboriginal rights belong to a group of people?", 4),
                multiple_choice_q(
                    11,
                    "What do many people consider the birth of Metis nationalism?",
                    ["WWI", "Metis land settlements", "Battle of Seven Oaks", "Six Nations Confederacy"],
                ),
                q(12, "What is a land claim?"),
                q(13, "What is the difference between individual and collective rights?", 4),
                q(14, "Explain one example of how individual rights are seen as collective rights by Aboriginal Peoples.", 4),
                q(15, "Self-determination is ultimately the right of a nation to determine what?", 4),
                q(16, "What does holistic mean and how does it relate to First Nation worldviews?", 4),
                q(17, "Why does a group have to be interdependent?", 4),
                q(18, "What do most First Nations typically think about interacting with another group in another territory?", 4),
                q(19, "How and why did traditional early First Nations treaties work?", 5),
                q(20, "What is the Great Law of Peace, and how does it reflect Indigenous worldviews?", 5),
                q(21, "What does the Two Row Wampum treaty recognize?", 4),
                q(22, "What is the Agreement of 1844?", 4),
                q(23, "Instead of compensation, how did Cuthbert Grant's peace agreement work?", 4),
                q(24, "What was the common European worldview in the sixteenth century?", 4),
                q(25, "What are the two types of colonization described in the textbook?", 4),
                q(26, "Why is the Royal Proclamation still significant today?", 4),
                q(27, "Why is the British North America Act still significant today?", 4),
                q(28, "Why did the government negotiate numbered treaties?", 4),
                q(29, "What was the difference between First Nations views of treaty agreements and European views?", 5),
                q(30, "What did First Nations want in exchange for allowing settlers on their land?", 4),
                q(31, "What were some of the problems with the treaty interpreters?", 4),
                q(32, "What was one verbal promise of Treaty One that was not included in the written document?", 4),
                q(33, "What were two concessions won in Treaty Six?", 4),
                q(34, "Did the treaties include all groups living in one area when they were originally written? Why or why not?", 5),
            ],
        },
        {
            "id": "traditional-governance",
            "title": "Traditional Governance and Colonization",
            "instructions": "Complete questions 35-68 using the Theme 1 readings and the Road Allowance People video where it fits.",
            "sourceRef": "Textbook pages 36-73.",
            "images": [],
            "prompts": [
                q(35, "What do the sweetgrass, the stone, and the fire of the pipe symbolize?", 4),
                fill_blank_q(36, "Complete the statement: The responsibility of today is:", ["The responsibility of today is ", "."]),
                table_q(
                    37,
                    "For each geographical area, summarize the environmental challenges and resources that shaped lives and culture.",
                    ["Pacific Northwest", "Plateau", "Plains", "Eastern Woodlands", "Subarctic", "Arctic"],
                    ["Environmental challenges", "Resources"],
                ),
                q(38, "What are two traditional Alberta meeting places?"),
                q(39, "What were some uses of controlled burns?", 4),
                fill_blank_q(
                    40,
                    "The Blackfoot call their land ____________________, which means ________________.",
                    ["The Blackfoot call their land ", ", which means ", "."],
                ),
                q(41, "Governance always responded to whose needs?"),
                q(42, "How were people reminded of the Creator's laws?", 4),
                q(43, "Describe how The Societies worked as law libraries.", 4),
                q(44, "What were the Laws of St. Albert?", 4),
                q(45, "How does a person become an informal leader?", 4),
                q(46, "What often makes Aboriginal leaders distinct from non-Aboriginal leaders?", 4),
                q(47, "How do students benefit when Elders share?", 4),
                q(48, "Why are Metis and Inuit peoples' relationships with the federal government different than those of other First Nations?", 5),
                q(49, "For First Nations, what did a guarantee of hunting rights by treaty also mean?", 4),
                fill_blank_q(
                    50,
                    "First Nations believed they were maintaining their right to _______________ themselves while Europeans believed the opposite.",
                    ["First Nations believed they were maintaining their right to ", " themselves while Europeans believed the opposite."],
                ),
                q(51, "Why was the Federal Government not eager to negotiate treaties with the Inuit of the Arctic?", 4),
                q(52, "Describe the system the government used to sort Metis people in Canada's early days.", 5),
                q(53, "Who was the famous Red River Metis leader?"),
                q(54, "What are some reasons the SCRIP program was a disaster?", 5),
                q(55, "What did the 1991 Royal Commission on Aboriginal Peoples report about the Metis?", 4),
                q(
                    56,
                    "Was life along the Road Allowances always bad? Why or why not?",
                    5,
                    resources=[theme_resource("Road Allowance People", "video", road_allowance_url, "Watch Video")],
                ),
                q(57, "What is assimilation? What government legislation had as its goal assimilation of First Nations?", 5),
                q(58, "Under the Indian Act, how was it determined who was Indian and who was not?", 4),
                q(59, "What sort of authority did the Indian Agent exercise over the bands they oversaw?", 4),
                q(60, "Why did Band Councils have no real authority?", 4),
                q(61, "Why did few bands adopt the European style government systems?", 4),
                q(62, "What is one example of how bands resisted the Indian Act?", 4),
                q(63, "How did the Federal Government force bands to comply with the Indian Act?", 4),
                q(64, "List and explain three changes made to the Indian Act in 1951.", 6),
                q(65, "For 1969, 1973, and 1988, list one important event in changes to First Nations/Federal Government interaction.", 6),
                q(66, "What is devolution?", 4),
                q(67, "Is it possible to still use traditional First Nations procedures for selecting band leaders? Explain.", 5),
                q(68, "What are three roles of a Tribal Council?", 5),
            ],
        },
        {
            "id": "rights-self-government",
            "title": "Aboriginal Rights and Self-Government",
            "instructions": "Complete questions 69-87 using the Theme 1 readings and the Metis self-governance video where it fits.",
            "sourceRef": "Textbook pages 76-105.",
            "images": [],
            "prompts": [
                q(69, "What do you think is the most important aspect of the lifestyle of the bush described in the opening story by James Carpenter?", 5),
                q(70, "Do you think Aboriginal People in Canada supported the new constitution in 1982? Why or why not?", 5),
                q(71, "What is a concern some Aboriginal leaders had over the Charter of Rights and Freedoms?", 4),
                q(72, "What did the Constitution Act mean for Metis?", 4),
                fill_blank_q(
                    73,
                    "________________ was a ________________ MLA who decided it was better to kill the accord than to betray his principles by ignoring the concerns of Aboriginal Peoples across the country.",
                    ["", " was a ", " MLA who decided it was better to kill the accord than to betray his principles by ignoring the concerns of Aboriginal Peoples across the country."],
                    ["Name", "Role"],
                ),
                q(74, "What did the Charlottetown Accord say about Aboriginal self-government?", 4),
                q(75, "What are the five main national Aboriginal political organizations?", 6),
                q(76, "Define Aboriginal title.", 4),
                q(77, "Why was it impossible for First Nations to give the land to the Europeans?", 4),
                q(78, "What is the problem with bands not being able to sell or mortgage land?", 4),
                multiple_choice_q(79, "True or False: First Nations people can hunt, fish and trap year-round outside of their province.", ["True", "False"]),
                q(80, "For an Aboriginal group to sell or use the resources of the land in a modern way, what must the practice be part of?", 4),
                q(81, "What makes a one-size-fits-all self-government impossible for Canada's Aboriginal Peoples?", 5),
                q(82, "What did the government response to the Commission pledge to do for First Nations?", 4),
                q(83, "Who are the three partners in self-government negotiations?", 4),
                q(84, "What are some benefits of enshrining self-government in the Constitution?", 5),
                q(85, "What is a current example of an Aboriginal public government in Canada?", 4),
                q(86, "Which bill created the first Aboriginal municipal-style government in Canada?", 4),
                q(87, "Urban-living First Nations individuals without a land base might benefit best from which type of government?", 4),
            ],
        },
        {
            "id": "assignment-1-2",
            "title": "Assignment 1.2: Rebuilding Self-Government",
            "instructions": "Use the self-government section after completing the numbered questions.",
            "sourceRef": "Textbook page 98 and response criteria.",
            "prompts": [
                activity_prompt(
                    "assignment-1-2",
                    "Choose one of the four advantages of the self-government promises on page 98 and discuss why you think it might be good for a First Nation. You may do additional research. Respond in paragraph form using the response criteria.",
                    8,
                    resources=[theme_resource("M\u00e9tis Self-Governance", "video", metis_governance_url, "Watch Video")],
                ),
            ],
        },
    ]


def build_theme_activities(units: list[dict], audit: dict) -> list[dict]:
    theme_one = next((unit for unit in units if unit["id"] == "theme-1"), None)
    if not theme_one:
        return []

    unit_resources = {item["title"]: item for item in theme_one.get("items", [])}
    road_allowance_url = unit_resources.get("Road Allowance People", {}).get("url", "https://www.youtube.com/embed/OFpTIbuaAao")
    metis_governance_url = unit_resources.get("Métis Self Governance", {}).get("url", "https://www.youtube.com/embed/IBuu2AjtwTM")
    oral_tradition_url = copy_theme_one_oral_tradition_reading(audit)
    for item in theme_one.get("items", []):
        if item.get("title") == "Walking Together: The Oral Tradition":
            item["url"] = oral_tradition_url
            item["external"] = False
    theme_one_items = theme_one.get("items", [])
    chapter_one_resource = {
        "id": "theme-1-chapter-1",
        "title": "Chapter 1",
        "kind": "chapter",
        "url": "./pdf-viewer.html?file=./assets/library/chapter-1.pdf&title=Chapter%201",
        "external": False,
    }
    theme_one["items"] = [item for item in theme_one_items if item.get("title") != "Chapter 1"]
    theme_one["items"].insert(0, chapter_one_resource)

    activity = {
        "id": "theme-1-online-booklet",
        "unitId": "theme-1",
        "title": "Theme 1 Questions",
        "intro": "Work through the checkpoints below using the textbook, linked readings, and videos. Write in complete sentences and use examples from the resources.",
        "resources": [],
        "sections": [
            {
                "id": "oral-tradition",
                "title": "Oral Tradition",
                "instructions": "Use the oral tradition reading and video before responding.",
                "prompts": [
                    {
                        "id": "oral-tradition-response",
                        "label": "In 2-3 paragraphs, define oral tradition and explain its importance and purpose.",
                        "rows": 7,
                    },
                    {
                        "id": "storytelling-example",
                        "label": "Identify one way storytelling can pass on values, responsibilities, history, or community knowledge.",
                        "rows": 4,
                    },
                ],
            },
            {
                "id": "nations-rights",
                "title": "Nations, Peoples, and Inherent Rights",
                "instructions": "Use textbook pages 2-18 to consolidate the short-answer questions from this part of the booklet.",
                "prompts": [
                    {
                        "id": "dene-declaration",
                        "label": "Summarize what the Dene are struggling for and what a just land settlement means.",
                        "rows": 5,
                    },
                    {
                        "id": "nation-people-rights",
                        "label": "Explain the terms nation, people, inherent rights, and collective rights in your own words.",
                        "rows": 6,
                    },
                    {
                        "id": "individual-collective",
                        "label": "Compare individual rights and collective rights, then give one Aboriginal rights example.",
                        "rows": 5,
                    },
                ],
            },
            {
                "id": "treaties-colonization",
                "title": "Treaties and Colonization",
                "instructions": "Use textbook pages 19-33 to connect early treaty relationships with later numbered treaties.",
                "prompts": [
                    {
                        "id": "early-treaties",
                        "label": "How did early First Nations treaty relationships work, and what did the Great Law of Peace or Two Row Wampum show about worldview?",
                        "rows": 6,
                    },
                    {
                        "id": "colonial-worldview",
                        "label": "How did European colonial worldview affect Indigenous lands, governance, and rights?",
                        "rows": 5,
                    },
                    {
                        "id": "numbered-treaties",
                        "label": "Why did the government negotiate numbered treaties, and how did First Nations understand treaty agreements differently?",
                        "rows": 6,
                    },
                ],
            },
            {
                "id": "traditional-governance",
                "title": "Traditional Governance and the Indian Act",
                "instructions": "Use textbook pages 36-73 and the Road Allowance video where it fits.",
                "prompts": [
                    {
                        "id": "land-resources",
                        "label": "Choose one geographic region and explain how land, environmental challenges, and resources shaped community life.",
                        "rows": 5,
                    },
                    {
                        "id": "land-governance",
                        "label": "Explain how land and governance are connected. Include one example such as the Sundance, societies, or Métis governance.",
                        "rows": 6,
                    },
                    {
                        "id": "leadership-elders",
                        "label": "Describe qualities of Indigenous leadership and explain the role Elders can play in learning.",
                        "rows": 5,
                    },
                    {
                        "id": "indian-act",
                        "label": "Explain assimilation, the Indian Act, Indian Agents, Band Councils, and one way First Nations resisted or changed the system.",
                        "rows": 7,
                    },
                ],
            },
            {
                "id": "self-government",
                "title": "Rebuilding Self-Government",
                "instructions": "Use textbook pages 76-105 and the Métis self-governance video.",
                "prompts": [
                    {
                        "id": "rights-freedoms",
                        "label": "Identify the key constitutional or legal developments that matter for Aboriginal rights and self-government.",
                        "rows": 6,
                    },
                    {
                        "id": "models-self-government",
                        "label": "Explain why one model of self-government cannot fit all Aboriginal Peoples in Canada.",
                        "rows": 5,
                    },
                    {
                        "id": "rebuilding-self-government",
                        "label": "Choose one self-government promise or advantage and explain why it could benefit a First Nation.",
                        "rows": 7,
                    },
                ],
            },
        ],
    }
    activity["intro"] = "Work through the Theme 1 booklet questions below using the textbook, linked readings, and videos. Write in complete sentences and use examples from the resources."
    activity["resources"] = []
    page_images = extract_theme_one_images(audit)
    activity["sections"] = build_theme_one_sections(page_images, oral_tradition_url, road_allowance_url, metis_governance_url)
    activity["sourceQuestionCount"] = 87
    activity["responsePromptCount"] = sum(len(section["prompts"]) for section in activity["sections"])
    audit["themeActivities"].append(
        {
            "unitId": "theme-1",
            "sourcePath": rel(THEME_1_BOOKLET_PDF),
            "bookletPages": 32,
            "sourceQuestionCount": activity["sourceQuestionCount"],
            "promptCount": activity["responsePromptCount"],
            "studentFacingAnswerKey": False,
        }
    )
    return [activity]


def write_course_data(units: list[dict], library_items: list[dict], film_room_items: list[dict], assignments: list[dict], theme_activities: list[dict], audit: dict) -> None:
    data = {
        "course": {
            "title": "Aboriginal Studies 30",
            "shortTitle": "AS30",
            "subtitle": "Course Shell",
            "enableLibrary": True,
        },
        "units": units,
        "themeActivities": theme_activities,
        "libraryItems": library_items,
        "filmRoomItems": film_room_items,
        "assignments": assignments,
        "quizzes": [],
        "sourceAudit": {
            "sourceZip": rel(SOURCE_ZIP),
            "courseMaterialsDir": rel(COURSE_MATERIALS_DIR),
            "includedUnits": len(units),
            "libraryItems": len(library_items),
            "filmRoomItems": len(film_room_items),
            "assignments": len(assignments),
            "themeActivities": len(theme_activities),
            "excludedAnswerKeyCount": len(audit["excludedAnswerKeys"]),
        },
    }
    source = "window.ABORIGINAL_STUDIES_30_DATA = "
    source += json.dumps(data, indent=2, ensure_ascii=False)
    source += ";\n"
    (WORKSPACE_DIR / "course-data.js").write_text(source, encoding="utf-8")


INDEX_HTML = """<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
  <title>Aboriginal Studies 30</title>
  <link href="https://fonts.googleapis.com/css2?family=Rajdhani:wght@400;500;600;700;800&display=swap" rel="stylesheet" />
  <link href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css" rel="stylesheet" />
  <link rel="stylesheet" href="./styles.css" />
</head>
<body data-project-slug="aboriginal-studies-30">
  <div class="app-shell">
    <aside class="sidebar">
      <header class="brand-block">
        <div>
          <h1>Aboriginal<br />Studies 30</h1>
          <p class="mono">Course Shell</p>
        </div>
        <div class="brand-mark" aria-hidden="true">AS</div>
      </header>

      <nav class="sidebar-nav" aria-label="Course sections">
        <button id="nav-home" class="nav-item active" type="button">
          <span class="nav-left"><i class="fa-solid fa-layer-group nav-icon"></i><span>Units</span></span>
        </button>
        <button id="nav-quizzes" class="nav-item" type="button">
          <span class="nav-left"><i class="fa-solid fa-circle-question nav-icon"></i><span>Quizzes</span></span>
        </button>
        <button id="nav-assignments" class="nav-item" type="button">
          <span class="nav-left"><i class="fa-solid fa-clipboard-list nav-icon"></i><span>Assignments</span></span>
        </button>
        <button id="nav-library" class="nav-item" type="button">
          <span class="nav-left"><i class="fa-solid fa-book nav-icon"></i><span>Library</span></span>
        </button>
        <button id="nav-film" class="nav-item" type="button">
          <span class="nav-left"><i class="fa-solid fa-tv nav-icon"></i><span>Film Room</span></span>
        </button>
      </nav>

      <div class="sidebar-save-host" data-google-hosted-controls-host="true"></div>
    </aside>

    <main class="content">
      <div class="content-inner">
        <section class="progress-shell" aria-label="Course progress">
          <div class="progress-panel">
            <div class="progress-inner">
              <div class="progress-top">
                <span class="progress-label mono">Course progress</span>
                <span id="progress-count" class="progress-count mono">0/4 units</span>
              </div>
              <div class="progress-track" aria-hidden="true">
                <div class="progress-fill" id="progress-fill"></div>
              </div>
              <div class="progress-bottom">
                <span id="progress-percent" class="progress-percent">0% <span class="progress-complete">complete</span></span>
                <span class="progress-meta mono"><i class="fa-solid fa-layer-group"></i> Modules: <strong>0/4</strong></span>
              </div>
            </div>
          </div>
        </section>

        <section class="modules-shell">
          <h2 id="section-title">Units</h2>
          <div id="content-body" class="content-body"></div>
        </section>
      </div>
    </main>
  </div>

  <script src="./course-data.js"></script>
  <script src="./main.js"></script>
</body>
</html>
"""


MAIN_JS = r"""const DATA = window.ABORIGINAL_STUDIES_30_DATA || {};

const STORAGE_KEYS = {
  progress: 'aboriginal-studies-30.progress',
  ui: 'aboriginal-studies-30.ui',
  activityResponses: 'aboriginal-studies-30.activityResponses'
};

const refs = {
  sectionTitle: document.getElementById('section-title'),
  contentBody: document.getElementById('content-body'),
  progressFill: document.getElementById('progress-fill'),
  progressPercent: document.getElementById('progress-percent'),
  progressCount: document.getElementById('progress-count'),
  navHome: document.getElementById('nav-home'),
  navQuizzes: document.getElementById('nav-quizzes'),
  navAssignments: document.getElementById('nav-assignments'),
  navLibrary: document.getElementById('nav-library'),
  navFilm: document.getElementById('nav-film')
};

const units = DATA.units || [];
const libraryItems = DATA.libraryItems || [];
const assignments = DATA.assignments || [];
const quizzes = DATA.quizzes || [];
const filmRoomItems = DATA.filmRoomItems || [];
const themeActivities = DATA.themeActivities || [];
const reviewUnlockAll = true;
const routeableSections = new Set(['home', 'unit', 'quizzes', 'assignments', 'assignment', 'library', 'film']);

let progress = loadJson(STORAGE_KEYS.progress, { completedUnits: [], completedAssignments: [] });
let activityResponses = loadJson(STORAGE_KEYS.activityResponses, {});
let state = loadJson(STORAGE_KEYS.ui, {
  section: 'home',
  activeUnitId: units[0]?.id || null,
  activeLibraryId: null,
  activeAssignmentId: null,
  activeFilmId: null
});

const routeParams = new URLSearchParams(window.location.search);
const requestedSection = routeParams.get('section');
if (routeableSections.has(requestedSection)) {
  state.section = requestedSection;
}
const requestedUnit = routeParams.get('unit');
if (requestedUnit && units.some((unit) => unit.id === requestedUnit)) {
  state.section = 'unit';
  state.activeUnitId = requestedUnit;
}

function loadJson(key, fallback) {
  try {
    return JSON.parse(localStorage.getItem(key) || '') || fallback;
  } catch (_error) {
    return fallback;
  }
}

function saveJson(key, value) {
  localStorage.setItem(key, JSON.stringify(value));
}

function escapeHtml(value) {
  return String(value || '')
    .replace(/&/g, '&amp;')
    .replace(/</g, '&lt;')
    .replace(/>/g, '&gt;')
    .replace(/"/g, '&quot;');
}

function completedUnitSet() {
  return new Set(progress.completedUnits || []);
}

function completedAssignmentSet() {
  return new Set(progress.completedAssignments || []);
}

function isUnitComplete(unitId) {
  return completedUnitSet().has(unitId);
}

function getUnitIndex(unitId) {
  return units.findIndex((unit) => unit.id === unitId);
}

function isUnitUnlocked(unitId) {
  if (reviewUnlockAll) return true;
  const index = getUnitIndex(unitId);
  if (index <= 0) return true;
  return isUnitComplete(units[index - 1].id);
}

function isAssignmentUnlocked(assignment) {
  if (reviewUnlockAll) return Boolean(assignment);
  return assignment && (!assignment.unitId || isUnitComplete(assignment.unitId));
}

function setSection(section) {
  state.section = section;
  saveJson(STORAGE_KEYS.ui, state);
  render();
}

function setActiveUnit(unitId) {
  state.section = 'unit';
  state.activeUnitId = unitId;
  saveJson(STORAGE_KEYS.ui, state);
  render();
}

function setActiveLibrary(itemId) {
  state.section = 'library';
  state.activeLibraryId = itemId;
  saveJson(STORAGE_KEYS.ui, state);
  render();
}

function setActiveAssignment(itemId) {
  state.section = 'assignment';
  state.activeAssignmentId = itemId;
  saveJson(STORAGE_KEYS.ui, state);
  render();
}

function setActiveFilm(itemId) {
  state.section = 'film';
  state.activeFilmId = itemId;
  saveJson(STORAGE_KEYS.ui, state);
  render();
}

function markUnitComplete(unitId) {
  if (!progress.completedUnits.includes(unitId)) {
    progress.completedUnits.push(unitId);
  }
  saveJson(STORAGE_KEYS.progress, progress);
  render();
}

function markAssignmentComplete(assignmentId) {
  if (!progress.completedAssignments.includes(assignmentId)) {
    progress.completedAssignments.push(assignmentId);
  }
  saveJson(STORAGE_KEYS.progress, progress);
  render();
}

function getUnitActivity(unitId) {
  return themeActivities.find((activity) => activity.unitId === unitId) || null;
}

function activityResponseKey(activityId, promptId) {
  return `${activityId}::${promptId}`;
}

function saveActivityResponse(key, value) {
  activityResponses[key] = value;
  saveJson(STORAGE_KEYS.activityResponses, activityResponses);
}

function autoGrowActivityTextarea(field) {
  if (!(field instanceof HTMLTextAreaElement)) return;
  field.style.height = 'auto';
  const computed = window.getComputedStyle(field);
  const maxHeight = Number.parseFloat(computed.maxHeight);
  const hasMaxHeight = Number.isFinite(maxHeight);
  const nextHeight = hasMaxHeight ? Math.min(field.scrollHeight, maxHeight) : field.scrollHeight;
  field.style.height = `${nextHeight}px`;
  field.style.overflowY = hasMaxHeight && field.scrollHeight > maxHeight ? 'auto' : 'hidden';
}

function updateProgress() {
  const total = units.length || 1;
  const done = units.filter((unit) => isUnitComplete(unit.id)).length;
  const percent = Math.round((done / total) * 100);
  refs.progressFill.style.width = `${percent}%`;
  refs.progressPercent.innerHTML = `${percent}% <span class="progress-complete">complete</span>`;
  refs.progressCount.textContent = `${done}/${units.length} units`;
  document.querySelector('.progress-meta strong').textContent = `${done}/${units.length}`;
}

function setActiveNav() {
  const navMap = {
    home: refs.navHome,
    unit: refs.navHome,
    quizzes: refs.navQuizzes,
    assignments: refs.navAssignments,
    assignment: refs.navAssignments,
    library: refs.navLibrary,
    film: refs.navFilm
  };
  for (const button of [refs.navHome, refs.navQuizzes, refs.navAssignments, refs.navLibrary, refs.navFilm]) {
    button?.classList.remove('active');
  }
  navMap[state.section]?.classList.add('active');
}

function renderHome() {
  refs.sectionTitle.textContent = 'Units';
  refs.contentBody.innerHTML = `
    <div class="stack-list">
      ${units.map((unit) => {
        const locked = !isUnitUnlocked(unit.id);
        const complete = isUnitComplete(unit.id);
        return `
          <button type="button" class="stack-card stack-card-button unit-card${locked ? ' is-locked' : ''}${complete ? ' is-complete' : ''}" data-unit-id="${unit.id}" ${locked ? 'disabled' : ''}>
            <span class="card-code mono">${escapeHtml(unit.code)}</span>
            <span class="card-lock-content">
              <strong>${escapeHtml(unit.title)}</strong>
              <span>${escapeHtml(locked ? 'Complete the previous unit to unlock this one.' : unit.description)}</span>
            </span>
          </button>
        `;
      }).join('')}
    </div>
  `;
  refs.contentBody.querySelectorAll('[data-unit-id]').forEach((button) => {
    button.addEventListener('click', () => setActiveUnit(button.dataset.unitId));
  });
}

function renderUnit() {
  const unit = units.find((item) => item.id === state.activeUnitId) || units[0];
  if (!unit) {
    renderHome();
    return;
  }
  const locked = !isUnitUnlocked(unit.id);
  const complete = isUnitComplete(unit.id);
  refs.sectionTitle.textContent = unit.title;
  refs.contentBody.innerHTML = `
    <div class="unit-view">
      <article class="detail-card">
        <div class="detail-head">
          <span class="card-code mono">${escapeHtml(unit.code)}</span>
          <h3>Resources</h3>
        </div>
        <p>${escapeHtml(locked ? 'This unit unlocks after the previous unit is marked complete.' : unit.description)}</p>
        <div class="resource-list">
          ${(unit.items || []).map((item) => renderResourceRow(item)).join('')}
        </div>
      </article>
      ${renderUnitActivity(unit, locked)}
      <div class="unit-completion-panel">
        <button type="button" id="back-to-units" class="secondary-button">Back to Units</button>
        <button type="button" id="mark-unit-complete" class="primary-button" ${locked || complete ? 'disabled' : ''}>${complete ? 'Completed' : 'Mark Complete'}</button>
      </div>
    </div>
  `;
  document.getElementById('back-to-units')?.addEventListener('click', () => setSection('home'));
  document.getElementById('mark-unit-complete')?.addEventListener('click', () => markUnitComplete(unit.id));
  bindActivityControls();
}

function renderResourceRow(item) {
  const url = item.url || '#';
  const label = item.kind === 'chapter' ? 'Open Chapter' : (item.kind === 'video' || item.kind === 'film' ? 'Watch' : 'Open');
  return `
    <div class="resource-row">
      <div>
        <span class="resource-kind mono">${escapeHtml(item.kind || 'resource')}</span>
        <strong>${escapeHtml(item.title)}</strong>
      </div>
      <a href="${escapeHtml(url)}" target="_blank" rel="noopener noreferrer">${label}</a>
    </div>
  `;
}

function renderUnitActivity(unit, locked) {
  const activity = getUnitActivity(unit.id);
  if (!activity) return '';
  if (locked) {
    return `
      <section class="activity-shell is-locked">
        <h3>${escapeHtml(activity.title)}</h3>
        <p>This online booklet unlocks with the unit.</p>
      </section>
    `;
  }
  const activityResources = Array.isArray(activity.resources) && activity.resources.length
    ? `
      <div class="activity-resources">
        ${activity.resources.map((resource) => renderActivityResource(resource)).join('')}
      </div>
    `
    : '';
  return `
    <section class="activity-shell" data-activity-id="${escapeHtml(activity.id)}">
      <div class="activity-head">
        <div>
          <span class="card-code mono">Online Work</span>
          <h3>${escapeHtml(activity.title)}</h3>
        </div>
        <span class="activity-status" data-activity-save-status>Responses save automatically.</span>
      </div>
      <p class="activity-intro">${escapeHtml(activity.intro)}</p>
      ${activityResources}
      <div class="activity-sections">
        ${(activity.sections || []).map((section, index) => renderActivitySection(activity, section, index)).join('')}
      </div>
      <div class="activity-actions">
        <button type="button" id="copy-activity-responses" class="secondary-button" data-copy-activity="${escapeHtml(activity.id)}">Copy Responses</button>
      </div>
    </section>
  `;
}

function renderActivityResource(resource) {
  const embedUrl = resource.kind === 'video' ? toEmbedUrl(resource.url) : '';
  return `
    <div class="activity-resource">
      <div class="activity-resource-text">
        <span class="resource-kind mono">${escapeHtml(resource.kind || 'resource')}</span>
        <strong>${escapeHtml(resource.title)}</strong>
      </div>
      ${embedUrl ? `
        <div class="activity-video">
          <iframe src="${escapeHtml(embedUrl)}" title="${escapeHtml(resource.title)}" allow="accelerometer; autoplay; clipboard-write; encrypted-media; gyroscope; picture-in-picture; web-share" referrerpolicy="strict-origin-when-cross-origin" allowfullscreen loading="lazy"></iframe>
        </div>
      ` : `
        <a href="${escapeHtml(resource.url)}" target="_blank" rel="noopener noreferrer">${escapeHtml(resource.actionLabel || 'Open')}</a>
      `}
    </div>
  `;
}

function renderActivitySection(activity, section, index) {
  const sourceRef = section.sourceRef ? `<p class="activity-source-ref">${escapeHtml(section.sourceRef)}</p>` : '';
  return `
    <section class="activity-section">
      <div class="activity-section-head">
        <span class="mono">${String(index + 1).padStart(2, '0')}</span>
        <div>
          <h4>${escapeHtml(section.title)}</h4>
          <p>${escapeHtml(section.instructions || '')}</p>
          ${sourceRef}
        </div>
      </div>
      ${renderActivitySectionImages(section.images || [])}
      <div class="activity-prompts">
        ${(section.prompts || []).map((prompt) => renderActivityPrompt(activity, prompt)).join('')}
      </div>
    </section>
  `;
}

function renderActivitySectionImages(images) {
  if (!images.length) return '';
  return `
    <div class="activity-section-images">
      ${images.map((image) => `
        <div class="activity-section-image">
          <img src="${escapeHtml(image.src)}" alt="${escapeHtml(image.alt || 'Theme booklet image')}" loading="lazy" />
        </div>
      `).join('')}
    </div>
  `;
}

function renderActivityPromptLabel(prompt) {
  return prompt.number
    ? `<span class="activity-question-label"><span class="activity-question-number">Q${escapeHtml(prompt.number)}</span><span>${escapeHtml(prompt.label)}</span></span>`
    : `<span>${escapeHtml(prompt.label)}</span>`;
}

function renderActivityPrompt(activity, prompt) {
  const key = activityResponseKey(activity.id, prompt.id);
  const inputId = `activity-${activity.id}-${prompt.id}`;
  const promptText = renderActivityPromptLabel(prompt);
  if (prompt.kind === 'fillBlank') return renderFillBlankPrompt(activity, prompt, promptText);
  if (prompt.kind === 'multipleChoice') return renderMultipleChoicePrompt(activity, prompt, promptText);
  if (prompt.kind === 'table') return renderTablePrompt(activity, prompt, promptText);
  return `
    <label class="activity-prompt" for="${escapeHtml(inputId)}">
      ${promptText}
      ${renderActivityPromptResources(prompt.resources || [])}
      <textarea id="${escapeHtml(inputId)}" class="activity-response" rows="${Number(prompt.rows) || 5}" data-activity-response="${escapeHtml(key)}">${escapeHtml(activityResponses[key] || '')}</textarea>
    </label>
  `;
}

function renderActivityPromptResources(resources) {
  if (!Array.isArray(resources) || !resources.length) return '';
  return `
    <div class="activity-prompt-resources">
      ${resources.map((resource) => {
        const embedUrl = resource.kind === 'video' ? toEmbedUrl(resource.url) : '';
        return `
          <div class="activity-prompt-resource">
            <div>
              <span class="resource-kind mono">${escapeHtml(resource.kind || 'resource')}</span>
              <strong>${escapeHtml(resource.title)}</strong>
            </div>
            ${embedUrl ? `
              <div class="activity-prompt-resource-video">
                <iframe src="${escapeHtml(embedUrl)}" title="${escapeHtml(resource.title)}" allow="accelerometer; autoplay; clipboard-write; encrypted-media; gyroscope; picture-in-picture; web-share" referrerpolicy="strict-origin-when-cross-origin" allowfullscreen loading="lazy"></iframe>
              </div>
            ` : `
              <a href="${escapeHtml(resource.url)}" target="_blank" rel="noopener noreferrer">${escapeHtml(resource.actionLabel || 'Open')}</a>
            `}
          </div>
        `;
      }).join('')}
    </div>
  `;
}

function promptFieldKey(activity, prompt, suffix) {
  const base = activityResponseKey(activity.id, prompt.id);
  return suffix ? `${base}.${suffix}` : base;
}

function fieldToken(value) {
  return String(value || 'item').toLowerCase().replace(/[^a-z0-9]+/g, '-').replace(/^-|-$/g, '') || 'item';
}

function renderFillBlankPrompt(activity, prompt, promptText) {
  const parts = Array.isArray(prompt.textParts) && prompt.textParts.length ? prompt.textParts : [prompt.label, ''];
  const blanks = Array.isArray(prompt.blanks) && prompt.blanks.length ? prompt.blanks : [{ id: 'blank-1', label: 'Blank' }];
  const heading = prompt.number
    ? `<span class="activity-fill-heading"><span class="activity-question-number">Q${escapeHtml(prompt.number)}</span></span>`
    : promptText;
  const line = blanks.map((blank, index) => {
    const key = promptFieldKey(activity, prompt, blank.id || `blank-${index + 1}`);
    const inputId = `activity-${activity.id}-${prompt.id}-${blank.id || index}`;
    return `
      ${escapeHtml(parts[index] || '')}
      <input id="${escapeHtml(inputId)}" type="text" class="activity-blank-input" data-activity-response="${escapeHtml(key)}" value="${escapeHtml(activityResponses[key] || '')}" aria-label="${escapeHtml(`Q${prompt.number || ''} ${blank.label || 'blank'}`)}" />
    `;
  }).join('') + escapeHtml(parts[blanks.length] || '');
  return `
    <fieldset class="activity-prompt activity-fill-blank">
      <legend>${heading}</legend>
      ${renderActivityPromptResources(prompt.resources || [])}
      <div class="activity-fill-line">${line}</div>
    </fieldset>
  `;
}

function renderMultipleChoicePrompt(activity, prompt, promptText) {
  const key = promptFieldKey(activity, prompt);
  const choices = Array.isArray(prompt.choices) ? prompt.choices : [];
  return `
    <fieldset class="activity-prompt activity-choice-prompt">
      <legend>${promptText}</legend>
      ${renderActivityPromptResources(prompt.resources || [])}
      <div class="activity-choice-list">
        ${choices.map((choice, index) => {
          const inputId = `activity-${activity.id}-${prompt.id}-choice-${index}`;
          const checked = activityResponses[key] === choice ? 'checked' : '';
          return `
            <label class="activity-choice" for="${escapeHtml(inputId)}">
              <input id="${escapeHtml(inputId)}" type="radio" name="${escapeHtml(key)}" value="${escapeHtml(choice)}" data-activity-response="${escapeHtml(key)}" ${checked} />
              <span>${escapeHtml(choice)}</span>
            </label>
          `;
        }).join('')}
      </div>
    </fieldset>
  `;
}

function renderTablePrompt(activity, prompt, promptText) {
  const rows = Array.isArray(prompt.rows) ? prompt.rows : [];
  const columns = Array.isArray(prompt.columns) ? prompt.columns : [];
  return `
    <div class="activity-prompt activity-table-prompt">
      <div>${promptText}</div>
      ${renderActivityPromptResources(prompt.resources || [])}
      <div class="activity-table-wrap">
        <table class="activity-table">
          <thead>
            <tr>
              <th scope="col">Area</th>
              ${columns.map((column) => `<th scope="col">${escapeHtml(column)}</th>`).join('')}
            </tr>
          </thead>
          <tbody>
            ${rows.map((row) => `
              <tr>
                <th scope="row">${escapeHtml(row)}</th>
                ${columns.map((column) => {
                  const key = promptFieldKey(activity, prompt, `${fieldToken(row)}.${fieldToken(column)}`);
                  return `<td><textarea class="activity-table-response" rows="3" data-activity-response="${escapeHtml(key)}" aria-label="${escapeHtml(`Q${prompt.number || ''} ${row} ${column}`)}">${escapeHtml(activityResponses[key] || '')}</textarea></td>`;
                }).join('')}
              </tr>
            `).join('')}
          </tbody>
        </table>
      </div>
    </div>
  `;
}

function bindActivityControls() {
  refs.contentBody.querySelectorAll('textarea[data-activity-response]').forEach((field) => {
    autoGrowActivityTextarea(field);
  });
  refs.contentBody.querySelectorAll('[data-activity-response]').forEach((field) => {
    const saveField = () => {
      autoGrowActivityTextarea(field);
      if (field.type === 'radio' && !field.checked) return;
      saveActivityResponse(field.dataset.activityResponse, field.value);
      const status = refs.contentBody.querySelector('[data-activity-save-status]');
      if (status) status.textContent = 'Saved just now.';
    };
    field.addEventListener(field.type === 'radio' ? 'change' : 'input', saveField);
  });
  refs.contentBody.querySelector('[data-copy-activity]')?.addEventListener('click', async (event) => {
    const activity = themeActivities.find((item) => item.id === event.currentTarget.dataset.copyActivity);
    if (!activity) return;
    const lines = [activity.title];
    for (const section of activity.sections || []) {
      lines.push('', section.title);
      for (const prompt of section.prompts || []) {
        lines.push(...activityPromptResponseLines(activity, prompt));
      }
    }
    const status = refs.contentBody.querySelector('[data-activity-save-status]');
    try {
      await navigator.clipboard.writeText(lines.join('\n'));
      if (status) status.textContent = 'Responses copied.';
    } catch (_error) {
      if (status) status.textContent = 'Copy unavailable. Responses are still saved here.';
    }
  });
}

function activityPromptResponseLines(activity, prompt) {
  const promptLabel = prompt.number ? `Q${prompt.number}. ${prompt.label}` : prompt.label;
  if (prompt.kind === 'fillBlank') {
    const blanks = Array.isArray(prompt.blanks) && prompt.blanks.length ? prompt.blanks : [{ id: 'blank-1', label: 'Blank' }];
    return [
      promptLabel,
      blanks.map((blank, index) => {
        const key = promptFieldKey(activity, prompt, blank.id || `blank-${index + 1}`);
        return `${blank.label || `Blank ${index + 1}`}: ${activityResponses[key] || ''}`;
      }).join('\n')
    ];
  }
  if (prompt.kind === 'multipleChoice') {
    const key = promptFieldKey(activity, prompt);
    return [promptLabel, activityResponses[key] || ''];
  }
  if (prompt.kind === 'table') {
    const rows = Array.isArray(prompt.rows) ? prompt.rows : [];
    const columns = Array.isArray(prompt.columns) ? prompt.columns : [];
    const tableLines = [promptLabel];
    for (const row of rows) {
      tableLines.push(row);
      for (const column of columns) {
        const key = promptFieldKey(activity, prompt, `${fieldToken(row)}.${fieldToken(column)}`);
        tableLines.push(`${column}: ${activityResponses[key] || ''}`);
      }
    }
    return tableLines;
  }
  const key = promptFieldKey(activity, prompt);
  return [promptLabel, activityResponses[key] || ''];
}

function renderQuizzes() {
  refs.sectionTitle.textContent = 'Quizzes';
  if (!quizzes.length) {
    refs.contentBody.innerHTML = `
      <div class="empty-card">
        <h3>No quizzes loaded yet</h3>
        <p>Quiz materials can be added here without changing the course shell.</p>
      </div>
    `;
    return;
  }
}

function renderAssignments() {
  refs.sectionTitle.textContent = 'Assignments';
  refs.contentBody.innerHTML = `
    <div class="stack-list">
      ${assignments.map((assignment) => {
        const locked = !isAssignmentUnlocked(assignment);
        const complete = completedAssignmentSet().has(assignment.id);
        return `
          <article class="stack-card assignment-card${locked ? ' is-locked' : ''}${complete ? ' is-complete' : ''}">
            <span class="card-code mono">${escapeHtml(assignment.unitTitle)}</span>
            <h3>${escapeHtml(assignment.title)}</h3>
            <p>${escapeHtml(locked ? 'Complete the related unit to unlock this assignment.' : assignment.summary)}</p>
            <div class="card-actions">
              <button type="button" data-assignment-id="${assignment.id}" ${locked ? 'disabled' : ''}>View Assignment</button>
              <a href="${escapeHtml(assignment.docxPath)}" target="_blank" rel="noopener noreferrer">Download DOCX</a>
            </div>
          </article>
        `;
      }).join('')}
    </div>
  `;
  refs.contentBody.querySelectorAll('[data-assignment-id]').forEach((button) => {
    button.addEventListener('click', () => setActiveAssignment(button.dataset.assignmentId));
  });
}

function renderAssignmentDetail() {
  const assignment = assignments.find((item) => item.id === state.activeAssignmentId) || assignments[0];
  if (!assignment) {
    renderAssignments();
    return;
  }
  const locked = !isAssignmentUnlocked(assignment);
  const complete = completedAssignmentSet().has(assignment.id);
  refs.sectionTitle.textContent = assignment.title;
  refs.contentBody.innerHTML = `
    <article class="detail-card assignment-detail">
      <div class="detail-head">
        <span class="card-code mono">${escapeHtml(assignment.unitTitle)}</span>
        <h3>${escapeHtml(assignment.title)}</h3>
      </div>
      ${locked ? '<p>Complete the related unit to unlock this assignment.</p>' : assignment.instructionsHtml}
      ${renderAssignmentLinks(assignment)}
      <div class="detail-actions">
        <button type="button" id="back-to-assignments" class="secondary-button">Back to Assignments</button>
        <a class="secondary-link" href="${escapeHtml(assignment.docxPath)}" target="_blank" rel="noopener noreferrer">Download DOCX</a>
        <button type="button" id="mark-assignment-complete" class="primary-button" ${locked || complete ? 'disabled' : ''}>${complete ? 'Completed' : 'Mark Complete'}</button>
      </div>
    </article>
  `;
  document.getElementById('back-to-assignments')?.addEventListener('click', () => setSection('assignments'));
  document.getElementById('mark-assignment-complete')?.addEventListener('click', () => markAssignmentComplete(assignment.id));
}

function renderAssignmentLinks(assignment) {
  if (!assignment.links?.length || !isAssignmentUnlocked(assignment)) return '';
  return `
    <div class="link-list">
      ${assignment.links.map((link) => `
        <a href="${escapeHtml(link.url)}" target="_blank" rel="noopener noreferrer">${escapeHtml(link.name)}</a>
      `).join('')}
    </div>
  `;
}

function renderLibrary() {
  const active = libraryItems.find((item) => item.id === state.activeLibraryId);
  const viewerSrc = active
    ? `./pdf-viewer.html?file=${encodeURIComponent(active.file)}&title=${encodeURIComponent(active.title)}`
    : '';
  refs.sectionTitle.textContent = 'Library';
  refs.contentBody.innerHTML = `
    <div class="stack-list">
      ${active ? `
        <article class="viewer-card">
          <div class="viewer-card-head">
            <div>
              <span class="card-code mono">${escapeHtml(active.code)}</span>
              <h3>${escapeHtml(active.title)}</h3>
            </div>
            <div class="card-actions">
              <a href="${escapeHtml(active.file)}" target="_blank" rel="noopener noreferrer">Download PDF</a>
              <button type="button" id="close-library-viewer">Close Viewer</button>
            </div>
          </div>
          <iframe src="${viewerSrc}" title="${escapeHtml(active.title)}"></iframe>
        </article>
      ` : ''}
      ${libraryItems.map((item) => `
        <article class="stack-card">
          <span class="card-code mono">${escapeHtml(item.code)}</span>
          <h3>${escapeHtml(item.title)}</h3>
          <p>${escapeHtml(item.description)}</p>
          <div class="card-actions">
            <button type="button" data-library-id="${item.id}">${item.kind === 'chapter' ? 'View Chapter' : 'View Resource'}</button>
            <a href="${escapeHtml(item.file)}" target="_blank" rel="noopener noreferrer">Download PDF</a>
          </div>
        </article>
      `).join('')}
    </div>
  `;
  refs.contentBody.querySelectorAll('[data-library-id]').forEach((button) => {
    button.addEventListener('click', () => setActiveLibrary(button.dataset.libraryId));
  });
  document.getElementById('close-library-viewer')?.addEventListener('click', () => {
    state.activeLibraryId = null;
    saveJson(STORAGE_KEYS.ui, state);
    render();
  });
}

function renderFilmRoom() {
  const active = filmRoomItems.find((item) => item.id === state.activeFilmId) || filmRoomItems[0] || null;
  if (!active) {
    refs.sectionTitle.textContent = 'Film Room';
    refs.contentBody.innerHTML = `
      <div class="empty-card">
        <h3>No films loaded yet</h3>
        <p>Video resources can be added here without changing the course shell.</p>
      </div>
    `;
    return;
  }
  const activeVideoNumber = Math.max(1, filmRoomItems.findIndex((item) => item.id === active.id) + 1);
  const activeType = toEmbedUrl(active.url) ? 'Embedded source' : 'Source link';
  const activeModuleLabel = moduleLabelFor(active);
  const activeModuleCode = moduleCodeFor(active);
  refs.sectionTitle.textContent = 'Film Room';
  refs.contentBody.innerHTML = `
    <section class="film-room-shell">
      <div class="film-room-stage">
        <div class="film-room-sign">
          <div>
            <p class="mono film-room-kicker">Aboriginal Studies Archive</p>
            <h4>Film Room</h4>
          </div>
          <div class="mono film-room-count">${filmRoomItems.length} videos loaded</div>
        </div>
        <div class="film-room-tv-wrap">
          <div class="film-room-antenna" aria-hidden="true">
            <span></span>
            <span></span>
          </div>
          <div class="film-room-tv">
            <div class="film-room-screen-shell">
              <div class="film-room-screen">
                ${renderMediaFrame(active)}
              </div>
            </div>
            <div class="film-room-console">
              <div class="film-room-slot" aria-hidden="true"></div>
              <div class="film-room-led mono">${escapeHtml(activeModuleCode)}</div>
            </div>
          </div>
        </div>
      </div>
      <aside class="film-room-sidebar">
        <article class="film-room-panel">
          <p class="mono film-room-kicker">Video catalog</p>
          <h4>Load a video</h4>
          <p>Use the playlist to switch videos without leaving the course shell.</p>
          <label class="film-room-label" for="film-room-select">Playlist</label>
          <select id="film-room-select" class="film-room-select" data-film-room-select>
            ${filmRoomItems.map((item) => `
              <option value="${item.id}"${item.id === active.id ? ' selected' : ''}>${escapeHtml(moduleLabelFor(item))} - ${escapeHtml(item.title)}</option>
            `).join('')}
          </select>
        </article>
        <article class="film-room-panel film-room-now-playing">
          <p class="mono film-room-kicker">Now loaded</p>
          <h4>${escapeHtml(activeModuleLabel)}</h4>
          <p class="film-room-title">${escapeHtml(active.title)}</p>
          <p>${escapeHtml(active.description)}</p>
          <div class="film-room-meta mono">
            <span>${escapeHtml(activeType)}</span>
            <span>${activeVideoNumber} / ${filmRoomItems.length}</span>
          </div>
          <a class="film-room-source" href="${escapeHtml(active.url)}" target="_blank" rel="noopener noreferrer">Open Source</a>
        </article>
      </aside>
    </section>
  `;
  refs.contentBody.querySelector('[data-film-room-select]')?.addEventListener('change', (event) => {
    setActiveFilm(event.target.value);
  });
}

function moduleLabelFor(item) {
  if (item.moduleLabel) return item.moduleLabel;
  const unit = units.find((candidate) => candidate.id === item.unitId);
  return unit?.title || 'Course media';
}

function moduleCodeFor(item) {
  if (item.moduleCode) return item.moduleCode;
  const unit = units.find((candidate) => candidate.id === item.unitId);
  return unit?.code || 'AS30';
}

function renderMediaFrame(item) {
  const embedUrl = toEmbedUrl(item.url);
  if (!embedUrl) {
    return `
      <div class="film-room-external">
        <span class="mono">${escapeHtml(item.kind || 'resource')}</span>
        <strong>${escapeHtml(item.title)}</strong>
        <a href="${escapeHtml(item.url)}" target="_blank" rel="noopener noreferrer">Open Source</a>
      </div>
    `;
  }
  return `<iframe src="${escapeHtml(embedUrl)}" title="${escapeHtml(item.title)}" allow="accelerometer; autoplay; clipboard-write; encrypted-media; gyroscope; picture-in-picture; web-share" referrerpolicy="strict-origin-when-cross-origin" allowfullscreen loading="lazy"></iframe>`;
}

function toEmbedUrl(url) {
  if (!url) return '';
  if (/youtube\.com\/embed\//i.test(url)) return url;
  const watchMatch = url.match(/[?&]v=([^&]+)/);
  if (/youtube\.com/i.test(url) && watchMatch) return `https://www.youtube.com/embed/${watchMatch[1]}`;
  const shortMatch = url.match(/youtu\.be\/([^?&#]+)/i);
  if (shortMatch) return `https://www.youtube.com/embed/${shortMatch[1]}`;
  if (/archive\.org\/embed\//i.test(url)) return url;
  const archiveMatch = url.match(/archive\.org\/details\/([^?&#/]+)/i);
  if (archiveMatch) return `https://archive.org/embed/${archiveMatch[1]}`;
  return '';
}

function render() {
  updateProgress();
  setActiveNav();
  if (state.section === 'unit') return renderUnit();
  if (state.section === 'quizzes') return renderQuizzes();
  if (state.section === 'assignments') return renderAssignments();
  if (state.section === 'assignment') return renderAssignmentDetail();
  if (state.section === 'library') return renderLibrary();
  if (state.section === 'film') return renderFilmRoom();
  return renderHome();
}

refs.navHome?.addEventListener('click', () => setSection('home'));
refs.navQuizzes?.addEventListener('click', () => setSection('quizzes'));
refs.navAssignments?.addEventListener('click', () => setSection('assignments'));
refs.navLibrary?.addEventListener('click', () => setSection('library'));
refs.navFilm?.addEventListener('click', () => setSection('film'));

render();
"""


STYLES_CSS = r""":root {
  color-scheme: dark;
  --bg: #0b111a;
  --bg-2: #111822;
  --panel: #151b25;
  --panel-2: #0f131a;
  --line: #2a3748;
  --line-soft: #1d2734;
  --line-strong: #3a4658;
  --text: #f4f7fb;
  --muted: #9aa6b6;
  --quiet: #78869a;
  --green: #00ffca;
  --green-dim: #00a676;
  --accent: #00ffca;
  --accent-2: #fbbf24;
  --danger: #fb7185;
}

* {
  box-sizing: border-box;
}

body {
  margin: 0;
  min-height: 100vh;
  background:
    radial-gradient(circle at top left, rgba(34, 54, 86, 0.9), transparent 30%),
    radial-gradient(circle at top right, rgba(21, 14, 58, 0.85), transparent 22%),
    linear-gradient(180deg, #0a1018 0%, #091018 55%, #071019 100%);
  color: var(--text);
  font-family: "Rajdhani", sans-serif;
  overflow: hidden;
}

button,
a {
  font: inherit;
}

a {
  color: var(--accent);
  font-weight: 700;
  text-underline-offset: 3px;
}

.app-shell {
  display: grid;
  grid-template-columns: 320px minmax(0, 1fr);
  min-height: 100vh;
}

.sidebar {
  background: linear-gradient(180deg, rgba(15, 20, 30, 0.96), rgba(14, 18, 26, 0.96));
  border-right: 1px solid var(--line);
  display: flex;
  flex-direction: column;
  min-width: 0;
}

.brand-block {
  padding: 24px 22px 20px;
  border-bottom: 1px solid var(--line-soft);
  display: flex;
  align-items: flex-start;
  justify-content: space-between;
  gap: 16px;
}

.brand-block h1 {
  margin: 0;
  font-size: 1.75rem;
  line-height: 0.9;
  letter-spacing: 0.08em;
  text-transform: uppercase;
  font-weight: 800;
  color: #ffffff;
}

.mono {
  font-family: "Rajdhani", sans-serif;
  font-size: 12px;
  font-weight: 800;
  letter-spacing: 0.18em;
  text-transform: uppercase;
}

.brand-block .mono {
  margin: 14px 0 0;
  color: #8b96a9;
  font-size: 0.65rem;
}

.brand-mark {
  width: 76px;
  height: 76px;
  display: grid;
  place-items: center;
  border: 1px solid #3a4658;
  background:
    radial-gradient(circle at 30% 28%, rgba(0, 255, 202, 0.14), transparent 42%),
    linear-gradient(180deg, #1c2531 0%, #121a23 100%);
  color: var(--accent);
  font-weight: 900;
  flex: none;
}

.sidebar-nav {
  padding: 16px 0 20px;
  display: grid;
  gap: 2px;
  overflow-y: auto;
}

.nav-item {
  width: 100%;
  border: 0;
  border-left: 4px solid transparent;
  display: flex;
  align-items: center;
  justify-content: space-between;
  gap: 12px;
  background: transparent;
  color: #8a94a3;
  padding: 16px 24px;
  cursor: pointer;
  text-align: left;
  text-transform: uppercase;
  letter-spacing: 0.14em;
  font-weight: 700;
  font-size: 0.82rem;
  transition: background 0.18s ease, color 0.18s ease, border-color 0.18s ease;
}

.nav-item:hover {
  background: rgba(255, 255, 255, 0.03);
  color: #eef2f8;
}

.nav-item.active {
  background: rgba(0, 166, 118, 0.92);
  color: #ffffff;
  border-left-color: #e6fff7;
  box-shadow: inset 0 0 0 1px rgba(255, 255, 255, 0.06), 0 0 18px rgba(0, 166, 118, 0.28);
}

.nav-left {
  display: inline-flex;
  align-items: center;
  gap: 10px;
}

.nav-icon {
  width: 24px;
  text-align: center;
  font-size: 1.05rem;
  color: currentColor;
  opacity: 1;
}

.sidebar-save-host {
  margin-top: auto;
  padding: 12px 12px 18px;
}

.sidebar-save-host:empty {
  display: none;
}

.content {
  min-width: 0;
  overflow: hidden;
}

.content-inner {
  min-width: 0;
  height: 100vh;
  overflow-y: auto;
  padding: 38px 42px 48px;
  background:
    radial-gradient(circle at 10% 0%, rgba(21, 31, 53, 0.9), transparent 40%),
    radial-gradient(circle at 88% 12%, rgba(31, 19, 60, 0.65), transparent 32%),
    linear-gradient(180deg, rgba(12, 18, 27, 0.8), rgba(8, 14, 22, 0.96));
}

.progress-shell {
  margin-bottom: 34px;
}

.progress-panel {
  max-width: 1200px;
  padding: 2px;
  background: #243347;
  clip-path: polygon(22px 0, 100% 0, 100% calc(100% - 22px), calc(100% - 22px) 100%, 0 100%, 0 22px);
  box-shadow: inset 0 0 18px rgba(0, 0, 0, 0.75);
}

.progress-inner {
  clip-path: polygon(18px 0, 100% 0, 100% calc(100% - 18px), calc(100% - 18px) 100%, 0 100%, 0 18px);
  background: rgba(17, 24, 34, 0.98);
  padding: 24px 28px 20px;
}

.progress-top {
  display: flex;
  justify-content: space-between;
  gap: 16px;
  align-items: flex-start;
  margin-bottom: 14px;
}

.progress-label {
  color: #b2bdcb;
  font-size: 0.72rem;
  font-weight: 700;
}

.progress-count {
  color: #93a4ba;
  font-size: 0.72rem;
  font-weight: 700;
}

.progress-track {
  position: relative;
  overflow: hidden;
  height: 22px;
  background: linear-gradient(180deg, #101722, #0a1018);
  border: 1px solid #394558;
  padding: 2px;
}

.progress-track::before {
  content: "";
  position: absolute;
  inset: 0;
  background: repeating-linear-gradient(90deg, rgba(255, 255, 255, 0.06) 0 4px, transparent 4px 8px);
  pointer-events: none;
}

.progress-fill {
  position: relative;
  z-index: 1;
  width: 0%;
  height: 100%;
  background: linear-gradient(90deg, var(--green-dim), var(--green));
  box-shadow: 0 0 12px rgba(0, 255, 202, 0.45);
}

.progress-bottom {
  margin-top: 16px;
  display: flex;
  justify-content: space-between;
  gap: 16px;
  align-items: flex-end;
  flex-wrap: wrap;
}

.progress-percent {
  font-size: clamp(3rem, 5vw, 4.8rem);
  font-weight: 800;
  line-height: 0.9;
  letter-spacing: 0.02em;
}

.progress-percent .progress-complete {
  font-size: 0.42em;
  font-weight: 500;
  color: #d2d8e0;
}

.progress-meta {
  display: flex;
  gap: 18px;
  flex-wrap: wrap;
  font-size: 0.88rem;
  color: #aab4c1;
}

.progress-meta i {
  color: var(--green);
  margin-right: 6px;
}

.progress-meta strong {
  color: #ffffff;
}

.modules-shell {
  max-width: 1200px;
}

.modules-shell h2 {
  margin: 0 0 18px;
  font-size: clamp(1.8rem, 2.6vw, 2.5rem);
  line-height: 1;
  text-transform: uppercase;
  letter-spacing: 0.08em;
  font-weight: 800;
}

.stack-list {
  display: grid;
  gap: 14px;
}

.stack-card,
.detail-card,
.viewer-card,
.empty-card {
  border: 1px solid var(--line);
  border-left: 4px solid var(--accent);
  border-radius: 10px;
  background: var(--panel);
  color: var(--text);
  padding: 18px;
}

.stack-card-button {
  display: grid;
  grid-template-columns: 70px minmax(0, 1fr);
  gap: 16px;
  text-align: left;
  cursor: pointer;
}

.stack-card-button:disabled {
  cursor: not-allowed;
}

.unit-card.is-locked .card-lock-content {
  filter: blur(2px);
  opacity: 0.72;
}

.stack-card.is-complete,
.unit-card.is-complete,
.assignment-card.is-complete {
  border-left-color: var(--accent-2);
}

.card-code,
.resource-kind {
  color: var(--accent);
}

.stack-card strong,
.stack-card h3,
.detail-card h3,
.viewer-card h3,
.empty-card h3 {
  display: block;
  margin: 0 0 8px;
  font-size: 22px;
  line-height: 1.15;
}

.stack-card p,
.detail-card p,
.empty-card p,
.assignment-detail li {
  margin: 0;
  color: var(--muted);
  font-size: 16px;
  line-height: 1.6;
}

.detail-head,
.viewer-card-head {
  display: flex;
  align-items: flex-start;
  justify-content: space-between;
  gap: 16px;
  margin-bottom: 12px;
}

.resource-list {
  display: grid;
  gap: 10px;
  margin: 18px 0;
}

.resource-row {
  display: flex;
  align-items: center;
  justify-content: space-between;
  gap: 14px;
  border: 1px solid var(--line);
  border-radius: 8px;
  background: var(--panel-2);
  padding: 12px;
}

.resource-row strong {
  display: block;
  margin-top: 4px;
}

.unit-view {
  display: grid;
  gap: 16px;
}

.unit-completion-panel {
  display: flex;
  gap: 10px;
  flex-wrap: wrap;
}

.activity-shell {
  border: 1px solid var(--line);
  border-left: 4px solid var(--accent);
  border-radius: 10px;
  background: rgba(15, 19, 26, 0.92);
  padding: 20px;
}

.activity-shell.is-locked {
  filter: blur(2px);
  opacity: 0.72;
}

.activity-head,
.activity-section-head {
  display: flex;
  align-items: flex-start;
  justify-content: space-between;
  gap: 18px;
}

.activity-head h3,
.activity-section h4 {
  margin: 4px 0 0;
  color: var(--text);
  line-height: 1.15;
}

.activity-head h3 {
  font-size: 24px;
}

.activity-status {
  color: var(--muted);
  font-size: 0.9rem;
  line-height: 1.4;
}

.activity-intro {
  margin: 14px 0 18px;
  color: #c7d0db;
  font-size: 16px;
  line-height: 1.6;
}

.activity-resources {
  display: grid;
  grid-template-columns: repeat(2, minmax(0, 1fr));
  gap: 12px;
  margin-bottom: 18px;
}

.activity-resource {
  border: 1px solid var(--line);
  border-radius: 8px;
  background: var(--panel-2);
  padding: 12px;
}

.activity-resource strong {
  display: block;
  margin-top: 4px;
  font-size: 18px;
  line-height: 1.2;
}

.activity-resource a {
  display: inline-flex;
  margin-top: 12px;
  min-height: 36px;
  align-items: center;
  border: 1px solid transparent;
  border-radius: 8px;
  background: var(--accent);
  color: #061017;
  padding: 8px 12px;
  text-decoration: none;
  font-weight: 900;
}

.activity-video {
  margin-top: 12px;
  aspect-ratio: 16 / 9;
  border-radius: 8px;
  overflow: hidden;
  background: #05080d;
}

.activity-video iframe {
  display: block;
  width: 100%;
  height: 100%;
  border: 0;
}

.activity-sections {
  display: grid;
}

.activity-section {
  border-top: 1px solid var(--line);
  padding: 18px 0;
}

.activity-section:last-child {
  padding-bottom: 0;
}

.activity-section-head {
  justify-content: flex-start;
}

.activity-section-head .mono {
  color: var(--accent);
}

.activity-section h4 {
  font-size: 21px;
}

.activity-section p {
  margin: 6px 0 0;
  color: var(--muted);
  font-size: 15px;
  line-height: 1.5;
}

.activity-source-ref {
  color: var(--accent) !important;
  font-weight: 800;
}

.activity-section-images {
  display: grid;
  grid-template-columns: repeat(auto-fit, minmax(160px, 1fr));
  gap: 12px;
  margin-top: 16px;
}

.activity-section-image {
  margin: 0;
  border: 1px solid var(--line);
  border-radius: 8px;
  background: var(--panel-2);
  overflow: hidden;
}

.activity-section-image img {
  display: block;
  width: 100%;
  height: 150px;
  object-fit: cover;
}

.activity-prompts {
  display: grid;
  gap: 12px;
  margin-top: 14px;
}

.activity-prompt {
  display: grid;
  gap: 8px;
  color: #d9e2ee;
  font-size: 16px;
  line-height: 1.45;
  border: 1px solid transparent;
  margin: 0;
  padding: 0;
}

.activity-prompt legend {
  display: block;
  width: 100%;
  padding: 0;
}

.activity-fill-heading {
  display: inline-flex;
  align-items: center;
}

.activity-prompt-resources {
  display: grid;
  gap: 10px;
  margin: 4px 0 6px;
}

.activity-prompt-resource {
  display: grid;
  gap: 10px;
  border: 1px solid var(--line);
  border-radius: 8px;
  background: rgba(255, 255, 255, 0.035);
  padding: 12px;
}

.activity-prompt-resource strong {
  display: block;
  margin-top: 3px;
  color: var(--text);
}

.activity-prompt-resource a {
  justify-self: start;
  color: var(--accent);
  font-weight: 800;
  text-decoration: underline;
  text-underline-offset: 3px;
}

.activity-prompt-resource-video {
  aspect-ratio: 16 / 9;
  overflow: hidden;
  border: 1px solid var(--line-strong);
  border-radius: 8px;
  background: #050a10;
}

.activity-prompt-resource-video iframe {
  display: block;
  width: 100%;
  height: 100%;
  border: 0;
}

.activity-question-label {
  display: grid;
  grid-template-columns: auto minmax(0, 1fr);
  align-items: start;
  gap: 10px;
}

.activity-question-number {
  min-width: 42px;
  color: var(--accent);
  font-weight: 900;
}

.activity-fill-line {
  display: flex;
  align-items: center;
  flex-wrap: wrap;
  gap: 8px;
}

.activity-blank-input {
  min-width: min(260px, 100%);
  flex: 1 1 220px;
  border: 0;
  border-bottom: 2px solid var(--line-strong);
  background: rgba(255, 255, 255, 0.03);
  color: var(--text);
  padding: 8px 10px;
  font: inherit;
}

.activity-blank-input:focus {
  outline: 2px solid rgba(0, 255, 202, 0.28);
  outline-offset: 2px;
  border-bottom-color: var(--accent);
}

.activity-choice-list {
  display: grid;
  grid-template-columns: repeat(auto-fit, minmax(220px, 1fr));
  gap: 8px;
}

.activity-choice {
  display: flex;
  align-items: center;
  gap: 10px;
  border: 1px solid var(--line);
  border-radius: 8px;
  background: var(--panel-2);
  padding: 10px 12px;
  cursor: pointer;
}

.activity-choice input {
  width: 18px;
  height: 18px;
  accent-color: var(--accent);
}

.activity-table-wrap {
  overflow-x: auto;
  border: 1px solid var(--line);
  border-radius: 8px;
}

.activity-table {
  width: 100%;
  min-width: 720px;
  border-collapse: collapse;
  background: var(--panel-2);
}

.activity-table th,
.activity-table td {
  border: 1px solid var(--line);
  padding: 10px;
  vertical-align: top;
}

.activity-table th {
  color: var(--text);
  text-align: left;
  font-weight: 900;
}

.activity-table-response {
  width: 100%;
  min-height: 86px;
  max-height: 260px;
  resize: none;
  overflow-y: auto;
  box-sizing: border-box;
  border: 1px solid var(--line-strong);
  border-radius: 8px;
  background: #0b111a;
  color: var(--text);
  padding: 10px;
  line-height: 1.45;
}

.activity-table-response:focus {
  outline: 2px solid rgba(0, 255, 202, 0.28);
  outline-offset: 2px;
}

.activity-response {
  width: 100%;
  min-height: 120px;
  max-height: 360px;
  resize: none;
  overflow-y: auto;
  box-sizing: border-box;
  border: 1px solid var(--line-strong);
  border-radius: 8px;
  background: #0b111a;
  color: var(--text);
  padding: 12px;
  line-height: 1.5;
}

.activity-response:focus {
  outline: 2px solid rgba(0, 255, 202, 0.28);
  outline-offset: 2px;
}

.activity-actions {
  display: flex;
  gap: 10px;
  flex-wrap: wrap;
  margin-top: 18px;
}

.card-actions,
.detail-actions,
.link-list {
  display: flex;
  align-items: center;
  gap: 10px;
  flex-wrap: wrap;
  margin-top: 16px;
}

.card-actions button,
.card-actions a,
.detail-actions button,
.detail-actions a,
.primary-button,
.secondary-button,
.secondary-link {
  display: inline-flex;
  align-items: center;
  justify-content: center;
  min-height: 38px;
  border-radius: 8px;
  border: 1px solid var(--line);
  background: var(--panel-2);
  color: var(--text);
  padding: 8px 12px;
  text-decoration: none;
  font-weight: 800;
  cursor: pointer;
}

.card-actions a,
.primary-button {
  border-color: transparent;
  background: var(--accent);
  color: #061017;
}

button:disabled {
  opacity: 0.5;
}

.viewer-card {
  padding: 0;
  overflow: hidden;
}

.viewer-card-head {
  margin: 0;
  padding: 16px 18px;
  border-bottom: 1px solid var(--line);
  background: rgba(255, 255, 255, 0.03);
}

.viewer-card iframe {
  display: block;
  width: 100%;
  min-height: 740px;
  border: 0;
  background: #0b1220;
}

.viewer-fallback {
  padding: 18px;
}

.film-room-shell {
  display: grid;
  grid-template-columns: minmax(0, 1.35fr) minmax(280px, 360px);
  gap: 18px;
  align-items: start;
}

.film-room-stage,
.film-room-panel {
  border: 1px solid var(--line);
  border-radius: 10px;
}

.film-room-stage {
  position: relative;
  overflow: hidden;
  padding: 24px;
  background: linear-gradient(180deg, rgba(15, 24, 38, 0.96), rgba(7, 13, 22, 0.96));
}

.film-room-stage::before {
  content: "";
  position: absolute;
  inset: 0;
  background:
    linear-gradient(rgba(65, 85, 110, 0.22) 1px, transparent 1px) 0 0 / 100% 52px,
    linear-gradient(90deg, rgba(65, 85, 110, 0.18) 1px, transparent 1px) 0 0 / 52px 100%;
  opacity: 0.58;
  pointer-events: none;
}

.film-room-sign,
.film-room-tv-wrap,
.film-room-sidebar {
  position: relative;
  z-index: 1;
}

.film-room-sign {
  display: flex;
  justify-content: space-between;
  align-items: end;
  gap: 16px;
  margin-bottom: 18px;
}

.film-room-sign h4,
.film-room-panel h4 {
  margin: 0;
  font-size: 1.08rem;
  text-transform: uppercase;
  letter-spacing: 0.08em;
}

.film-room-sign h4 {
  color: var(--text);
  font-size: clamp(1.16rem, 2vw, 1.8rem);
}

.film-room-kicker {
  margin: 0 0 8px;
  font-size: 0.68rem;
  color: var(--accent);
}

.film-room-count {
  color: #93a4ba;
  font-size: 0.74rem;
  white-space: nowrap;
}

.film-room-tv-wrap {
  display: grid;
  justify-items: center;
  padding-top: 10px;
}

.film-room-antenna {
  position: relative;
  width: 142px;
  height: 42px;
  margin-bottom: -8px;
}

.film-room-antenna span {
  position: absolute;
  bottom: 0;
  left: 50%;
  width: 2px;
  height: 62px;
  background: rgba(190, 205, 225, 0.72);
  transform-origin: bottom center;
  box-shadow: 0 0 0 1px rgba(0, 0, 0, 0.12);
}

.film-room-antenna span:first-child {
  transform: translateX(-24px) rotate(-36deg);
}

.film-room-antenna span:last-child {
  transform: translateX(24px) rotate(36deg);
}

.film-room-tv {
  width: min(100%, 900px);
  border-radius: 28px;
  padding: 22px 22px 18px;
  background: linear-gradient(180deg, #1b2431, #101722);
  border: 1px solid #334155;
  box-shadow:
    0 20px 40px rgba(0, 0, 0, 0.38),
    inset 0 1px 0 rgba(255, 255, 255, 0.05);
}

.film-room-screen-shell {
  padding: 20px;
  border-radius: 24px;
  background: linear-gradient(180deg, #101722, #06090d);
  box-shadow: inset 0 0 0 1px rgba(255, 255, 255, 0.03);
}

.film-room-screen {
  position: relative;
  overflow: hidden;
  border-radius: 18px;
  aspect-ratio: 4 / 3;
  background: #040506;
  box-shadow:
    inset 0 0 60px rgba(0, 0, 0, 0.78),
    inset 0 0 0 1px rgba(255, 255, 255, 0.03);
}

.film-room-screen::before {
  content: "";
  position: absolute;
  inset: 0;
  background:
    radial-gradient(circle at 50% 16%, rgba(255, 255, 255, 0.12), transparent 34%),
    radial-gradient(circle at 50% 68%, rgba(0, 0, 0, 0), rgba(0, 0, 0, 0.34) 78%);
  pointer-events: none;
  z-index: 2;
}

.film-room-screen::after {
  content: "";
  position: absolute;
  inset: 0;
  background: repeating-linear-gradient(
    180deg,
    rgba(255, 255, 255, 0.07) 0,
    rgba(255, 255, 255, 0.07) 1px,
    transparent 1px,
    transparent 4px
  );
  opacity: 0.28;
  mix-blend-mode: screen;
  pointer-events: none;
  z-index: 3;
}

.film-room-screen iframe {
  display: block;
  width: 100%;
  height: 100%;
  border: 0;
  background: #000;
}

.film-room-external {
  position: relative;
  z-index: 1;
  min-height: 100%;
  display: grid;
  align-content: center;
  justify-items: center;
  gap: 14px;
  padding: 28px;
  text-align: center;
}

.film-room-external strong {
  font-size: clamp(1.1rem, 2vw, 1.65rem);
  line-height: 1.15;
}

.film-room-external a,
.film-room-source {
  border: 1px solid var(--line-strong);
  border-radius: 8px;
  color: #061017;
  background: var(--accent);
  padding: 9px 12px;
  text-decoration: none;
  font-weight: 900;
}

.film-room-console {
  display: grid;
  grid-template-columns: minmax(0, 1fr) auto;
  gap: 18px;
  align-items: center;
  margin-top: 18px;
}

.film-room-slot {
  height: 12px;
  border-radius: 999px;
  background: #090a0c;
  box-shadow: inset 0 0 0 1px rgba(255, 255, 255, 0.04);
}

.film-room-led {
  min-width: 88px;
  padding: 10px 12px;
  border-radius: 8px;
  background: #101722;
  color: var(--accent);
  text-align: center;
  letter-spacing: 0.16em;
  box-shadow: inset 0 0 0 1px rgba(52, 211, 153, 0.22);
}

.film-room-sidebar {
  display: grid;
  gap: 14px;
}

.film-room-panel {
  padding: 18px;
  background: rgba(11, 17, 26, 0.88);
}

.film-room-panel p:last-of-type {
  color: #c1c9d4;
  line-height: 1.55;
}

.film-room-label {
  display: block;
  margin: 16px 0 8px;
  font-size: 0.72rem;
  font-weight: 700;
  text-transform: uppercase;
  letter-spacing: 0.08em;
  color: var(--accent);
}

.film-room-select {
  width: 100%;
  min-height: 46px;
  border: 1px solid rgba(52, 64, 82, 0.96);
  border-radius: 8px;
  background: rgba(13, 20, 30, 0.92);
  color: var(--text);
  padding: 0 14px;
  font: inherit;
}

.film-room-select:focus {
  outline: 2px solid rgba(52, 211, 153, 0.3);
  outline-offset: 2px;
}

.film-room-title {
  margin: 12px 0 0;
  color: #f5f8fd;
  font-size: 0.98rem;
  line-height: 1.55;
}

.film-room-meta {
  margin-top: 14px;
  display: flex;
  justify-content: space-between;
  gap: 10px;
  color: var(--accent);
  font-size: 0.72rem;
}

.film-room-source {
  display: inline-block;
  margin-top: 16px;
}

.assignment-detail :where(p, li) {
  margin: 0 0 12px;
}

.assignment-detail ul {
  color: var(--muted);
}

.link-list {
  align-items: stretch;
}

.link-list a {
  border: 1px solid var(--line);
  border-radius: 8px;
  background: var(--panel-2);
  padding: 10px 12px;
  text-decoration: none;
}

@media (max-width: 860px) {
  .app-shell {
    grid-template-columns: 1fr;
  }

  .sidebar {
    position: static;
  }

  .content-inner {
    width: 100%;
    padding: 22px 16px 42px;
  }

  .film-room-shell {
    grid-template-columns: 1fr;
  }

  .film-room-stage {
    padding: 18px;
  }

  .activity-resources {
    grid-template-columns: 1fr;
  }

  .activity-head {
    display: grid;
  }

  .film-room-sign {
    flex-direction: column;
    align-items: flex-start;
  }

  .stack-card-button {
    grid-template-columns: 1fr;
  }
}

@media (max-width: 640px) {
  .film-room-tv {
    border-radius: 20px;
    padding: 16px 16px 14px;
  }

  .film-room-screen-shell {
    padding: 14px;
    border-radius: 18px;
  }

  .film-room-console {
    grid-template-columns: 1fr;
  }

  .film-room-led {
    width: 100%;
  }

  .film-room-meta {
    flex-direction: column;
  }
}
"""


PDF_VIEWER_HTML = r"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Chapter Viewer</title>
  <style>
    :root {
      color-scheme: dark;
      --bg: #08111c;
      --panel: #0f1826;
      --line: rgba(95, 120, 155, 0.24);
      --text: #e8eef6;
      --muted: #8a96a8;
      --accent: #34d399;
    }

    * { box-sizing: border-box; }
    body {
      margin: 0;
      min-height: 100vh;
      background: var(--bg);
      color: var(--text);
      font-family: Arial, sans-serif;
    }
    .viewer-shell { min-height: 100vh; padding: 18px; }
    .viewer-toolbar {
      position: sticky;
      top: 0;
      z-index: 2;
      display: flex;
      justify-content: space-between;
      gap: 12px;
      margin-bottom: 16px;
      padding: 12px 14px;
      border: 1px solid var(--line);
      border-radius: 10px;
      background: #0a111c;
    }
    .viewer-title {
      font-size: 12px;
      font-weight: 800;
      letter-spacing: 0.12em;
      text-transform: uppercase;
      color: var(--muted);
    }
    .viewer-meta {
      font-size: 12px;
      color: var(--accent);
      font-weight: 700;
    }
    .viewer-status {
      padding: 16px 18px;
      border: 1px solid var(--line);
      border-radius: 10px;
      background: var(--panel);
      color: var(--muted);
      font-size: 14px;
    }
    .viewer-status.error { color: #ff9db0; }
    .viewer-pages { display: grid; gap: 18px; }
    .page-card {
      padding: 14px;
      border: 1px solid var(--line);
      border-radius: 10px;
      background: var(--panel);
    }
    .page-label {
      margin: 0 0 10px;
      font-size: 12px;
      font-weight: 800;
      color: var(--muted);
    }
    canvas {
      display: block;
      width: 100%;
      height: auto;
      border-radius: 8px;
      background: white;
    }
  </style>
</head>
<body>
  <div class="viewer-shell">
    <div class="viewer-toolbar">
      <div>
        <div class="viewer-title" id="viewer-title">Chapter Viewer</div>
        <div class="viewer-meta" id="viewer-meta">Loading document</div>
      </div>
    </div>
    <div id="viewer-status" class="viewer-status">Loading PDF...</div>
    <div id="viewer-pages" class="viewer-pages" hidden></div>
  </div>

  <script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.min.js"></script>
  <script>
    const params = new URLSearchParams(window.location.search);
    const file = params.get('file');
    const title = params.get('title') || 'Chapter Viewer';

    const titleNode = document.getElementById('viewer-title');
    const metaNode = document.getElementById('viewer-meta');
    const statusNode = document.getElementById('viewer-status');
    const pagesNode = document.getElementById('viewer-pages');

    titleNode.textContent = title;
    metaNode.textContent = file ? 'Rendering in app' : 'Missing file';
    pdfjsLib.GlobalWorkerOptions.workerSrc = 'https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.worker.min.js';

    function setError(message) {
      statusNode.textContent = message;
      statusNode.classList.add('error');
      pagesNode.hidden = true;
    }

    function getScale(viewport) {
      const maxWidth = Math.min(window.innerWidth - 80, 1100);
      return Math.max(1, maxWidth / viewport.width);
    }

    async function renderPdf() {
      if (!file) {
        setError('No PDF file was provided to the viewer.');
        return;
      }
      try {
        const pdf = await pdfjsLib.getDocument(file).promise;
        metaNode.textContent = pdf.numPages === 1 ? '1 page' : `${pdf.numPages} pages`;
        pagesNode.hidden = false;
        statusNode.hidden = true;
        for (let pageNumber = 1; pageNumber <= pdf.numPages; pageNumber += 1) {
          const page = await pdf.getPage(pageNumber);
          const baseViewport = page.getViewport({ scale: 1 });
          const viewport = page.getViewport({ scale: getScale(baseViewport) });
          const card = document.createElement('section');
          card.className = 'page-card';
          const label = document.createElement('p');
          label.className = 'page-label';
          label.textContent = `Page ${pageNumber}`;
          const canvas = document.createElement('canvas');
          const context = canvas.getContext('2d');
          canvas.width = viewport.width;
          canvas.height = viewport.height;
          card.append(label, canvas);
          pagesNode.appendChild(card);
          await page.render({ canvasContext: context, viewport }).promise;
        }
      } catch (error) {
        setError(error && error.message ? error.message : 'Unable to render this PDF in the app.');
      }
    }

    renderPdf();
  </script>
</body>
</html>
"""


def write_project_manifest() -> None:
    now = datetime.now().isoformat(timespec="seconds")
    manifest = {
        "id": "aboriginal-studies-30",
        "slug": "aboriginal-studies-30",
        "sourcePath": rel(SOURCE_ZIP),
        "inputKind": "brightspace-zip",
        "brightspaceTarget": "course-page",
        "previewModes": ["workspace"],
        "workspaceEntrypoint": rel(WORKSPACE_DIR / "index.html"),
        "rawEntrypoint": rel(RAW_DIR / "original.html"),
        "migrationState": "migrated",
        "projectType": "conversion",
        "preferredWorkflows": ["conversion"],
        "canonicalEntry": rel(WORKSPACE_DIR / "index.html"),
        "canonicalSources": [
            rel(WORKSPACE_DIR / "index.html"),
            rel(WORKSPACE_DIR / "course-data.js"),
            rel(WORKSPACE_DIR / "main.js"),
            rel(WORKSPACE_DIR / "styles.css"),
            rel(WORKSPACE_DIR / "pdf-viewer.html"),
        ],
        "importedFirstPassOrigin": {
            "sourceSystem": "brightspace",
            "sourcePath": rel(SOURCE_ZIP),
            "importedAt": now,
            "notes": "Sports Wellness-style shell generated from the Aboriginal Studies 30 Brightspace export and chapter PDF folder.",
        },
        "exportTargets": [
            {"target": "brightspace", "enabled": True, "notes": "Workspace can be exported after review."},
            {"target": "google-hosted", "enabled": False, "notes": "Ready for metadata/config when a Firebase Hosting site is provided."},
        ],
        "authoringStatus": "active",
        "generatedOutputs": [],
        "regenerateCommand": "python projects/aboriginal-studies-30/meta/build_sports_style_course.py",
        "injectedComponents": [],
        "referenceOnly": [],
        "sourceOfTruthNotes": "Edit workspace files for course shell changes. Regenerate from the source Brightspace ZIP and Course Materials folder when the source mapping changes.",
        "googleHosted": {
            "trackedStorageKeys": [
                "aboriginal-studies-30.progress",
                "aboriginal-studies-30.ui",
                "aboriginal-studies-30.activityResponses",
            ],
            "authMode": "google",
        },
        "learningSource": "other",
        "learningTrust": "auto",
        "learningUpdatedAt": now,
        "updatedAt": now,
        "workspaceApprovedAt": now,
    }
    (META_DIR / "project.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")


def write_e2e_contract() -> None:
    contract = {
        "$schema": "../../../e2e/project-e2e-contract.schema.json",
        "projectSlug": "aboriginal-studies-30",
        "requiredTestIds": [
            "studio-shell",
            "course-studio-tab",
            "workspace-project-select",
            "project-root",
            "workspace-preview-frame",
        ],
    }
    (META_DIR / "e2e-contract.json").write_text(json.dumps(contract, indent=2), encoding="utf-8")


def build() -> None:
    audit = {
        "sourceZip": rel(SOURCE_ZIP),
        "courseMaterialsDir": rel(COURSE_MATERIALS_DIR),
        "includedUnits": [],
        "unitItems": [],
        "assignments": [],
        "copiedCourseMaterialPdfs": [],
        "copiedZipPdfs": [],
        "copiedSourcePages": [],
        "themeActivities": [],
        "excludedAnswerKeys": [
            "AB-Studies-30-Theme-1-Key.pdf",
            "AB-Studies-30-Theme-2-Key.pdf",
            "AB-Studies-30-Theme-3-Key.pdf",
            "AB-Studies-30-Theme-4-Key.pdf",
        ],
    }
    if WORKSPACE_DIR.exists():
        shutil.rmtree(WORKSPACE_DIR)
    WORKSPACE_DIR.mkdir(parents=True, exist_ok=True)
    RAW_DIR.mkdir(parents=True, exist_ok=True)
    META_DIR.mkdir(parents=True, exist_ok=True)

    with zipfile.ZipFile(SOURCE_ZIP) as zip_file:
        manifest_root = read_manifest(zip_file)
        resources = read_resources(manifest_root)
        library_items = build_library(zip_file, audit)
        units, film_room_items, dropbox_unit_by_resource_code = build_units_and_film_room(zip_file, manifest_root, resources, audit)
        assignments = build_assignments(dropbox_unit_by_resource_code, units, audit, zip_file)
        theme_activities = build_theme_activities(units, audit)
        replace_oral_tradition_assignment_links(assignments)

    write_course_data(units, library_items, film_room_items, assignments, theme_activities, audit)
    (WORKSPACE_DIR / "index.html").write_text(INDEX_HTML, encoding="utf-8")
    (WORKSPACE_DIR / "main.js").write_text(MAIN_JS, encoding="utf-8")
    (WORKSPACE_DIR / "styles.css").write_text(STYLES_CSS, encoding="utf-8")
    (WORKSPACE_DIR / "pdf-viewer.html").write_text(PDF_VIEWER_HTML, encoding="utf-8")
    (RAW_DIR / "original.html").write_text(
        textwrap.dedent(
            f"""\
            <!doctype html>
            <html lang="en">
            <head><meta charset="utf-8"><title>Aboriginal Studies 30 Source</title></head>
            <body>
              <h1>Aboriginal Studies 30 Source Package</h1>
              <p>Generated from Brightspace ZIP: {html.escape(str(SOURCE_ZIP))}</p>
              <p>Generated from course materials folder: {html.escape(str(COURSE_MATERIALS_DIR))}</p>
            </body>
            </html>
            """
        ),
        encoding="utf-8",
    )
    (META_DIR / "source-zip-audit.json").write_text(json.dumps(audit, indent=2), encoding="utf-8")
    write_project_manifest()
    write_e2e_contract()


if __name__ == "__main__":
    build()
