from __future__ import annotations

import importlib.util
import io
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from lxml import html as lxml_html
from PIL import Image


REPO_ROOT = Path(__file__).resolve().parents[2]
BUILDER_PATH = (
    REPO_ROOT
    / "projects"
    / "social-studies-10-1-docx-export"
    / "meta"
    / "build_unit1_docx_export.py"
)

spec = importlib.util.spec_from_file_location("social_docx_builder", BUILDER_PATH)
builder = importlib.util.module_from_spec(spec)
assert spec and spec.loader
sys.modules[spec.name] = builder
spec.loader.exec_module(builder)


def make_zip(entries: dict[str, bytes]) -> zipfile.ZipFile:
    data = io.BytesIO()
    with zipfile.ZipFile(data, "w") as zip_out:
        for name, payload in entries.items():
            zip_out.writestr(name, payload)
    data.seek(0)
    return zipfile.ZipFile(data)


def png_bytes(width: int = 64, height: int = 48) -> bytes:
    image = Image.new("RGB", (width, height), "#87b9d6")
    output = io.BytesIO()
    image.save(output, format="PNG")
    return output.getvalue()


def make_exporter(zip_file: zipfile.ZipFile):
    exporter = builder.UnitOneDocxExporter.__new__(builder.UnitOneDocxExporter)
    exporter.zip_file = zip_file
    exporter.audit = {
        "embeddedImages": 0,
        "mediaReferences": [],
        "unresolvedAssets": [],
        "linkedLocalResources": [],
    }
    return exporter


class SocialStudiesDocxExportTests(unittest.TestCase):
    def test_first_existing_path_prefers_environment_override(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            existing = Path(temp_dir) / "source.zip"
            existing.write_bytes(b"zip")

            resolved = builder.first_existing_path("unused-env-name", [Path(temp_dir) / "missing.zip", existing])

        self.assertEqual(resolved, existing)

    def test_output_path_falls_back_when_canonical_docx_is_locked(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            canonical = Path(temp_dir) / "lesson.docx"
            canonical.write_bytes(b"open in word")
            exporter = builder.UnitOneDocxExporter.__new__(builder.UnitOneDocxExporter)
            exporter.audit = {}
            exporter.output_path_is_locked = lambda path: path == canonical

            resolved = exporter.available_output_path(canonical)

        self.assertEqual(resolved, Path(temp_dir) / "lesson - refreshed.docx")

    def test_missing_images_are_audited_without_visible_source_noise(self) -> None:
        with make_zip({}) as zip_file:
            exporter = make_exporter(zip_file)
            document = Document()
            paragraph = document.add_paragraph()
            image_node = lxml_html.fragment_fromstring('<img src="assets/missing.png">')
            context = builder.RenderContext(
                base_href="content/unit/lesson.html",
                heading_base=2,
                source_title="Lesson",
            )

            embedded = exporter.add_image_to_paragraph(paragraph, image_node, context, max_width=6.3)

        self.assertFalse(embedded)
        self.assertEqual(paragraph.text, "")
        self.assertEqual(
            exporter.audit["unresolvedAssets"],
            [{"sourceHtml": "content/unit/lesson.html", "src": "assets/missing.png", "kind": "image"}],
        )

    def test_image_source_links_wrapping_images_do_not_render_body_text(self) -> None:
        with make_zip({"content/unit/assets/photo.png": png_bytes()}) as zip_file:
            exporter = make_exporter(zip_file)
            document = Document()
            paragraph = document.add_paragraph()
            link_node = lxml_html.fragment_fromstring(
                '<a href="https://example.test/image-source"><img src="assets/photo.png" alt="Photo"></a>'
            )
            context = builder.RenderContext(
                base_href="content/unit/lesson.html",
                heading_base=2,
                source_title="Lesson",
            )

            exporter.add_link_or_local_resource(paragraph, link_node, context)

        self.assertNotIn("Image source", paragraph._p.xml)
        self.assertNotIn("https://example.test/image-source", paragraph._p.xml)
        self.assertEqual(exporter.audit["embeddedImages"], 1)

    def test_ambiguous_image_basename_resolution_prefers_nearest_package_path(self) -> None:
        with make_zip(
            {
                "content/unit/shared/assets/soccer.png": b"unit image",
                "content/other/assets/soccer.png": b"other image",
            }
        ) as zip_file:
            resolved = builder.resolve_package_href(
                "content/unit/lesson/page.html",
                "assets/soccer.png",
                zip_file,
            )

        self.assertEqual(resolved, "content/unit/shared/assets/soccer.png")

    def test_media_cards_hide_raw_embed_urls_from_document_body(self) -> None:
        with make_zip({}) as zip_file:
            exporter = make_exporter(zip_file)
            document = exporter.new_document()
            media_node = lxml_html.fragment_fromstring(
                '<iframe title="How technology evolves" src="https://embed.ted.com/talks/kevin_kelly_how_technology_evolves"></iframe>'
            )
            context = builder.RenderContext(
                base_href="content/unit/lesson.html",
                heading_base=2,
                source_title="Lesson",
            )

            exporter.add_media_card(document, media_node, context)

        body_xml = document._body._element.xml
        self.assertIn("How technology evolves", body_xml)
        self.assertNotIn("https://embed.ted.com", body_xml)
        self.assertNotIn("Iframe preserved from Brightspace", body_xml)

    def test_youtube_embeds_use_provider_thumbnail_when_possible(self) -> None:
        with make_zip({}) as zip_file:
            exporter = make_exporter(zip_file)

            thumbnail = exporter.media_thumbnail_url(
                "https://www.youtube.com/embed/onD5UOP5z_c?si=Ed4OK6Ssz_cyMSiQ&rel=0"
            )

        self.assertEqual(thumbnail, "https://img.youtube.com/vi/onD5UOP5z_c/hqdefault.jpg")

    def test_site_header_uses_word_text_not_scaled_header_image(self) -> None:
        with make_zip({}) as zip_file:
            exporter = make_exporter(zip_file)
            document = exporter.new_document()

            exporter.add_site_header(document, "SS10-1 U1 Lesson 2")

        body_xml = document._body._element.xml
        self.assertIn("SS10-1 U1 Lesson 2", body_xml)
        self.assertNotIn("<w:drawing>", body_xml)

    def test_youtube_preview_picture_is_hyperlinked(self) -> None:
        with make_zip({}) as zip_file:
            exporter = make_exporter(zip_file)
            exporter.fetch_remote_image_bytes = lambda _url: png_bytes(480, 270)
            document = exporter.new_document()
            media_node = lxml_html.fragment_fromstring(
                '<iframe title="YouTube video player" src="https://www.youtube.com/embed/onD5UOP5z_c?rel=0"></iframe>'
            )
            context = builder.RenderContext(
                base_href="content/unit/lesson.html",
                heading_base=2,
                source_title="Lesson",
            )

            exporter.add_media_card(document, media_node, context)

        body_xml = document._body._element.xml
        self.assertIn("<w:hyperlink", body_xml)
        self.assertIn("<w:drawing>", body_xml)
        self.assertLess(body_xml.index("<w:hyperlink"), body_xml.index("<w:drawing>"))


if __name__ == "__main__":
    unittest.main()
